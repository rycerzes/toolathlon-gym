#!/usr/bin/env python3
"""Evaluation script for supplier-contract-compliance-audit task validation"""

from argparse import ArgumentParser
import json
import os
import sys

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try: return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError): return False


def str_match(a, b):
    if a is None or b is None: return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_xlsx_content(workspace, groundtruth_workspace="."):
    print("\n=== Check: XLSX files ===")
    import openpyxl
    for fname in ["compliance_assessment.xlsx", "corrective_actions.xlsx"]:
        xlsx_path = os.path.join(workspace, fname)
        if not os.path.isfile(xlsx_path):
            record(f"xlsx {fname} exists", False, "Not found")
            continue
        record(f"xlsx {fname} exists", True)
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                record(f"xlsx {fname} '{ws.title}' has data", len(rows) >= 2, f"{len(rows)} rows")
            # --- Groundtruth XLSX value comparison ---
            gt_path = os.path.join(groundtruth_workspace, fname)
            if os.path.isfile(gt_path):
                gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
                for gt_sname in gt_wb.sheetnames:
                    gt_ws = gt_wb[gt_sname]
                    a_ws = None
                    for asn in wb.sheetnames:
                        if asn.strip().lower() == gt_sname.strip().lower():
                            a_ws = wb[asn]; break
                    if a_ws is None:
                        record(f"GT sheet '{gt_sname}' in {fname}", False, f"Available: {wb.sheetnames}")
                        continue
                    gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                    a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                    record(f"GT '{gt_sname}' row count in {fname}", len(a_rows) == len(gt_rows),
                           f"Expected {len(gt_rows)}, got {len(a_rows)}")
                    for ri in range(min(3, len(gt_rows))):
                        if ri >= len(a_rows): break
                        ok = True
                        for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                            gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                            if gv is None: continue
                            if isinstance(gv, (int, float)):
                                if not num_close(av, gv, max(abs(gv)*0.1, 1.0)): ok = False; break
                            else:
                                if not str_match(av, gv): ok = False; break
                        record(f"GT '{gt_sname}' row {ri+1} values in {fname}", ok,
                               f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'missing'}")
                gt_wb.close()
            wb.close()
        except Exception as e:
            record(f"xlsx {fname} readable", False, str(e))


def check_docx_content(workspace):
    print("\n=== Check: DOCX files ===")
    from docx import Document
    for fname in ["audit_findings.docx", "audit_summary.docx"]:
        path = os.path.join(workspace, fname)
        if not os.path.isfile(path):
            record(f"docx {fname} exists", False, "Not found")
            continue
        record(f"docx {fname} exists", True)
        try:
            doc = Document(path)
            record(f"docx {fname} has content", len(doc.paragraphs) > 0, f"{len(doc.paragraphs)} paragraphs")
        except Exception as e:
            record(f"docx {fname} readable", False, str(e))


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
