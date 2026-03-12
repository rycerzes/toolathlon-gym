"""Evaluation for scholarly-literature-review-gcal-word-email.

Checks:
1. Literature_Review.docx exists with title, at least 6 paper sections, and Summary section
2. GCal event "Literature Review Presentation" on 2026-03-21
3. Email to research-team@university.example.com with "Literature Review" in subject
"""
import argparse
import json
import os
import sys

import psycopg2

DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "toolathlon_gym", "user": "eigent", "password": "camel"}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = (detail[:300] + "...") if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def check_word(agent_ws):
    print("\n=== Check 1: Literature_Review.docx ===")
    path = os.path.join(agent_ws, "Literature_Review.docx")
    check("File Literature_Review.docx exists", os.path.isfile(path))
    if not os.path.isfile(path):
        return

    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        check("Word doc is readable", False, str(e))
        return

    full_text = " ".join(p.text for p in doc.paragraphs)
    full_text_lower = full_text.lower()

    check("Doc contains 'Literature' and 'Review' in title/heading",
          "literature" in full_text_lower and "review" in full_text_lower)

    # Count paper entries - look for citation-related keywords
    citation_indicators = ["citation", "citations", "cite", "year:", "authors:", "pub_year", "2017", "2019", "2020", "2022", "2023"]
    has_citations = sum(1 for kw in citation_indicators if kw in full_text_lower)
    check("Doc contains citation data (years/authors)", has_citations >= 3,
          f"Citation indicators found: {has_citations}")

    # Check key paper titles appear
    key_papers = ["attention", "bert", "gpt", "llama", "transformer"]
    found_papers = sum(1 for p in key_papers if p in full_text_lower)
    check("Doc mentions at least 3 key LLM papers", found_papers >= 3,
          f"Found {found_papers}/5 key paper keywords")

    # Check for Summary/Conclusion section
    check("Doc has Summary or Conclusion section",
          "summary" in full_text_lower or "conclusion" in full_text_lower)

    # Count headings to estimate number of paper sections
    heading_count = sum(1 for p in doc.paragraphs if p.style.name.startswith("Heading"))
    check("Doc has at least 7 headings (title + 6+ paper sections + summary)",
          heading_count >= 7,
          f"Found {heading_count} headings")


def check_gcal():
    print("\n=== Check 2: Google Calendar Event ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, summary, start_datetime, description FROM gcal.events
        WHERE summary ILIKE '%literature%review%'
           OR summary ILIKE '%literature review%'
           OR (summary ILIKE '%literature%' AND summary ILIKE '%review%')
        LIMIT 10
    """)
    events = cur.fetchall()
    check("GCal event with 'Literature Review' in title found",
          len(events) > 0, "No matching events in gcal.events")

    if events:
        event = events[0]
        evt_start = event[2]
        check("Event scheduled on 2026-03-21",
              evt_start is not None and "2026-03-21" in str(evt_start),
              f"Event start: {evt_start}")
        desc = str(event[3] or "").lower()
        check("Event description mentions LLM or large language models",
              "llm" in desc or "large language" in desc or "literature" in desc,
              f"Description: {str(event[3])[:100]}")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, to_addr, body_text FROM email.messages
        WHERE subject ILIKE '%literature%review%'
           OR to_addr::text ILIKE '%research-team%'
        LIMIT 10
    """)
    rows = cur.fetchall()
    check("Email with 'Literature Review' in subject found",
          len(rows) > 0, "No matching email found")

    if rows:
        to_addrs = [str(r[1]) for r in rows]
        check("Email sent to research-team@university.example.com",
              any("research-team" in addr for addr in to_addrs),
              f"To addresses: {to_addrs}")
        bodies = [str(r[2] or "").lower() for r in rows]
        check("Email body mentions paper titles or authors",
              any("attention" in b or "bert" in b or "gpt" in b or "llama" in b or "vaswani" in b
                  or "devlin" in b or "brown" in b or "touvron" in b for b in bodies),
              f"Body excerpt: {bodies[0][:200] if bodies else ''}")

    cur.close()
    conn.close()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=== Evaluation: scholarly-literature-review-gcal-word-email ===")

    check_word(args.agent_workspace)
    check_gcal()
    check_email()

    print(f"\n=== SUMMARY: {PASS_COUNT} passed, {FAIL_COUNT} failed ===")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"pass": PASS_COUNT, "fail": FAIL_COUNT}, f)

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
