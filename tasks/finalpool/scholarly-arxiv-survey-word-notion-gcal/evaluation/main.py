"""
Evaluation for scholarly-arxiv-survey-word-notion-gcal task.

Checks:
1. LLM_Reasoning_Survey.docx exists and has required sections
2. Word doc mentions all 5 reasoning methods and papers
3. Word doc has comparative analysis section
4. Notion database "Reasoning Papers" exists
5. Notion has at least 5 paper entries
6. GCal has 5 reading group sessions in the week of March 16-20, 2026
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document LLM_Reasoning_Survey.docx ===")

    docx_path = os.path.join(agent_workspace, "LLM_Reasoning_Survey.docx")
    if not os.path.exists(docx_path):
        record("LLM_Reasoning_Survey.docx exists", False, f"Not found at {docx_path}")
        return
    record("LLM_Reasoning_Survey.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Check for table
    has_table = len(doc.tables) >= 1
    table_text = ""
    if has_table:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_text += cell.text.lower() + " "
    combined_text = all_text + " " + table_text

    # Check required sections
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith("Heading") and p.text.strip()]
    headings_lower = [h.lower() for h in headings]

    has_abstract = any("abstract" in h for h in headings_lower) or "abstract" in all_text[:500]
    has_intro = any("introduction" in h for h in headings_lower)
    has_background = any("background" in h for h in headings_lower)
    has_taxonomy = any("taxonomy" in h or "method" in h for h in headings_lower)
    has_comparative = any("comparative" in h or "comparison" in h for h in headings_lower)
    has_challenges = any("challenge" in h or "open" in h for h in headings_lower)
    has_conclusion = any("conclusion" in h for h in headings_lower)

    section_count = sum([has_abstract, has_intro, has_background, has_taxonomy,
                         has_comparative, has_challenges, has_conclusion])
    record("Has at least 5 required sections", section_count >= 5,
           f"Found {section_count}/7: abstract={has_abstract}, intro={has_intro}, "
           f"background={has_background}, taxonomy={has_taxonomy}, "
           f"comparative={has_comparative}, challenges={has_challenges}, conclusion={has_conclusion}")

    # Check for 5 reasoning methods
    has_cot = "chain-of-thought" in combined_text or "chain of thought" in combined_text
    has_tot = "tree of thought" in combined_text or "tree-of-thought" in combined_text
    has_sc = "self-consistency" in combined_text or "self consistency" in combined_text
    has_verify = ("step by step" in combined_text and "verif" in combined_text) or "process supervision" in combined_text
    has_auto = "auto-cot" in combined_text or "automatic chain" in combined_text or "automatic cot" in combined_text

    method_count = sum([has_cot, has_tot, has_sc, has_verify, has_auto])
    record("Mentions at least 4 of 5 reasoning methods", method_count >= 4,
           f"Found {method_count}/5: CoT={has_cot}, ToT={has_tot}, SC={has_sc}, "
           f"Verify={has_verify}, AutoCoT={has_auto}")

    # Check for paper titles/authors
    has_wei = "wei" in combined_text
    has_yao = "yao" in combined_text
    has_wang = "wang" in combined_text
    has_lightman = "lightman" in combined_text or "let's verify" in combined_text or "lets verify" in combined_text
    has_zhang = "zhang" in combined_text and "auto" in combined_text

    author_count = sum([has_wei, has_yao, has_wang, has_lightman, has_zhang])
    record("References at least 3 paper authors", author_count >= 3,
           f"Found {author_count}/5: Wei={has_wei}, Yao={has_yao}, Wang={has_wang}, "
           f"Lightman={has_lightman}, Zhang(Auto)={has_zhang}")

    # Check comparative table or analysis
    has_comparison_content = (
        has_table or
        ("comparative" in combined_text and any(m in combined_text for m in ["accuracy", "cost", "performance"]))
    )
    record("Has comparative analysis content (table or structured comparison)",
           has_comparison_content,
           "No table found and no comparative analysis keywords")


def check_notion_database():
    print("\n=== Check 2: Notion Database ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    # Check for "Reasoning Papers" database (title is jsonb)
    cur.execute("SELECT id, title, properties FROM notion.databases")
    databases = cur.fetchall()

    reasoning_db = None
    for db_id, title_json, props in databases:
        # title is jsonb: could be a list like [{"text": {"content": "..."}}] or a string
        title_str = ""
        if isinstance(title_json, list):
            for item in title_json:
                if isinstance(item, dict):
                    title_str += item.get("text", {}).get("content", "") if isinstance(item.get("text"), dict) else ""
                    title_str += item.get("plain_text", "") if "plain_text" in item else ""
        elif isinstance(title_json, str):
            title_str = title_json
        else:
            title_str = json.dumps(title_json) if title_json else ""

        if "reasoning" in title_str.lower() and "paper" in title_str.lower():
            reasoning_db = (db_id, title_str, props)
            break

    db_titles = []
    for _, t, _ in databases:
        db_titles.append(str(t)[:60] if t else "None")
    record("Notion 'Reasoning Papers' database exists", reasoning_db is not None,
           f"Databases found: {db_titles}")

    # Check for paper pages
    cur.execute("SELECT properties FROM notion.pages")
    pages = cur.fetchall()

    # Count pages that look like paper entries (have reasoning-related keywords)
    paper_pages = 0
    method_keywords = ["chain", "tree", "self-consistency", "verify", "step by step",
                       "auto-cot", "automatic", "reasoning", "thought"]
    for (props,) in pages:
        props_str = json.dumps(props).lower() if isinstance(props, dict) else str(props).lower()
        if any(kw in props_str for kw in method_keywords):
            paper_pages += 1

    record("At least 5 paper entries in Notion", paper_pages >= 5,
           f"Found {paper_pages} paper-like pages out of {len(pages)} total")

    # Check for database properties (if database exists)
    if reasoning_db:
        db_id, db_title, db_props = reasoning_db
        if db_props:
            props = db_props if isinstance(db_props, dict) else {}
            props_str = json.dumps(props).lower()
            has_method = "method" in props_str
            has_year = "year" in props_str
            has_key = "key" in props_str or "contribution" in props_str
            prop_count = sum([has_method, has_year, has_key])
            record("Database has key properties (Method, Year, Key_Contribution)",
                   prop_count >= 2,
                   f"Properties: {list(props.keys()) if isinstance(props, dict) else 'N/A'}")

    cur.close()
    conn.close()


def check_gcal():
    print("\n=== Check 3: Google Calendar Reading Group Sessions ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    # Check for reading group events in the week of March 16-20, 2026
    cur.execute("""
        SELECT summary, start_datetime, end_datetime
        FROM gcal.events
        WHERE start_datetime >= '2026-03-16' AND start_datetime < '2026-03-21'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()

    reading_events = [
        e for e in events
        if "reading group" in (e[0] or "").lower()
    ]

    record("At least 5 reading group events in March 16-20", len(reading_events) >= 5,
           f"Found {len(reading_events)} reading group events in target week")

    if reading_events:
        # Check duration (should be ~1 hour)
        summary, start_dt, end_dt = reading_events[0]
        if start_dt and end_dt:
            duration_hours = (end_dt - start_dt).total_seconds() / 3600
            record("Reading sessions are ~1 hour", 0.5 <= duration_hours <= 1.5,
                   f"Duration: {duration_hours:.1f} hours")

        # Check they are on different days
        dates = set(e[1].date() for e in reading_events if e[1])
        record("Sessions on different days (at least 4 distinct dates)", len(dates) >= 4,
               f"Found {len(dates)} distinct dates: {sorted(dates)}")

        # Check summaries contain topic descriptors
        summaries_text = " ".join(e[0].lower() for e in reading_events)
        topic_keywords = ["chain", "tree", "self", "verif", "step", "auto", "thought", "consistency"]
        topic_matches = sum(1 for kw in topic_keywords if kw in summaries_text)
        record("Session titles contain paper topic descriptors", topic_matches >= 3,
               f"Matched {topic_matches} topic keywords in summaries: {[e[0] for e in reading_events]}")

    cur.close()
    conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion_database()
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
