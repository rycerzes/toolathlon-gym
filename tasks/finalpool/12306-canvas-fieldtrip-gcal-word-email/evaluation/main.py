"""
Evaluation for 12306-canvas-fieldtrip-gcal-word-email.

Checks:
1. Field_Trip_Notice.docx exists in agent_workspace
2. Word doc has at least 4 headings
3. Word doc contains G235 and G236 train numbers (or departure times 17:30 and 15:00)
4. Word doc contains Qufu and Beijing
5. Word doc contains 1106 (round trip cost per person)
6. GCal has at least 2 new field trip events
7. Canvas announcement created with field trip content
8. Email sent to students@university.edu
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Field_Trip_Notice.docx ===")
    import glob

    pattern = os.path.join(agent_workspace, "*.docx")
    all_docx = glob.glob(pattern)
    trip_docs = [f for f in all_docx if any(
        kw in os.path.basename(f).lower()
        for kw in ["field", "trip", "notice", "qufu", "fieldtrip"]
    )]

    if not trip_docs:
        record("Field trip notice docx exists", False,
               f"No matching docx found in {agent_workspace}")
        return
    record("Field trip notice docx exists", True)

    doc_path = trip_docs[0]
    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Word doc has at least 4 headings", len(headings) >= 4,
           f"Found {len(headings)} headings")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    has_g235 = "g235" in full_text
    has_g236 = "g236" in full_text
    has_times = "17:30" in full_text and "15:00" in full_text
    record("Contains G235 and G236 train numbers (or departure times)",
           (has_g235 and has_g236) or has_times,
           f"G235:{has_g235}, G236:{has_g236}, times:{has_times}")

    has_qufu = "qufu" in full_text
    has_beijing = "beijing" in full_text
    record("Contains Qufu and Beijing", has_qufu and has_beijing,
           f"Qufu:{has_qufu}, Beijing:{has_beijing}")

    has_cost = "1106" in full_text
    record("Contains round-trip cost 1106", has_cost,
           f"Text snippet: {full_text[:300]}")


def check_gcal():
    print("\n=== Check 2: Calendar field trip events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE (
            (start_datetime >= '2026-03-12' AND start_datetime < '2026-03-13')
            OR (start_datetime >= '2026-03-15' AND start_datetime < '2026-03-16')
        )
        AND summary NOT ILIKE '%regular class%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()
    record("At least 2 new field trip calendar events", len(events) >= 2,
           f"Found {len(events)} events: {[e[0] for e in events]}")


def check_canvas_announcement():
    print("\n=== Check 3: Canvas announcement ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        cur.execute("""
            SELECT COUNT(*) FROM canvas.announcements
            WHERE title ILIKE '%field%'
               OR title ILIKE '%qufu%'
               OR title ILIKE '%trip%'
               OR title ILIKE '%travel%'
               OR message ILIKE '%qufu%'
               OR message ILIKE '%g235%'
        """)
        cnt = cur.fetchone()[0]
        record("Canvas announcement for field trip created", cnt >= 1,
               f"Found {cnt} matching announcements")
    except Exception as e:
        # Canvas may not have this table in test env
        record("Canvas announcement check (canvas table query)", False, str(e))
    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 4: Email to students@university.edu ===")
    cnt = 0
    sent = 0
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            "SELECT COUNT(*) FROM email.messages"
            " WHERE to_addr::text ILIKE '%students@university.edu%'"
            "   AND from_addr NOT ILIKE '%students@university.edu%'"
        )
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
    except Exception:
        pass
    try:
        conn2 = psycopg2.connect(**DB_CONFIG)
        cur2 = conn2.cursor()
        cur2.execute(
            "SELECT COUNT(*) FROM email.sent_log"
            " WHERE to_addr::text ILIKE '%students@university.edu%'"
        )
        sent = cur2.fetchone()[0]
        cur2.close()
        conn2.close()
    except Exception:
        pass
    record("Email sent to students@university.edu", cnt >= 1 or sent >= 1,
           f"messages: {cnt}, sent_log: {sent}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_canvas_announcement()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
