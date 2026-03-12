"""
Evaluation for yt-transcript-afrobeat-song-list-excel-email task.

Checks:
1. Afrobeat_Tracklist.xlsx exists in agent_workspace
2. "Tracklist" sheet has >= 8 data rows
3. "Tracklist" sheet has Track_Number, Song_Title, Artist columns
4. "Artist_Summary" sheet exists with >= 4 rows
5. Curator_Notes.docx exists with >= 3 headings
6. Curator_Notes.docx contains music-related keywords
7. Notion page exists with "Afrobeat" or "Mix" in title
8. Email sent to music@label.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1-4: Afrobeat_Tracklist.xlsx ===")
    xlsx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".xlsx") and ("afrobeat" in fname.lower() or "tracklist" in fname.lower()):
            xlsx_path = os.path.join(agent_workspace, fname)
            break

    record("Afrobeat_Tracklist.xlsx exists", xlsx_path is not None,
           f"No matching xlsx found in {agent_workspace}")

    if not xlsx_path:
        record("Tracklist sheet has >= 8 data rows", False, "xlsx not found")
        record("Tracklist sheet has required columns", False, "xlsx not found")
        record("Artist_Summary sheet has >= 4 rows", False, "xlsx not found")
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        # Check Tracklist sheet
        tracklist_sheet = None
        for name in wb.sheetnames:
            if "tracklist" in name.lower() or "track" in name.lower():
                tracklist_sheet = wb[name]
                break
        if tracklist_sheet is None and wb.sheetnames:
            tracklist_sheet = wb[wb.sheetnames[0]]

        if tracklist_sheet:
            data_rows = [r for r in tracklist_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Tracklist sheet has >= 8 data rows", len(data_rows) >= 8,
                   f"Found {len(data_rows)} data rows")

            headers = [str(c.value).strip() if c.value else "" for c in next(tracklist_sheet.iter_rows(max_row=1))]
            headers_lower = [h.lower() for h in headers]
            has_track_num = any("track" in h and "num" in h or h == "track_number" for h in headers_lower)
            has_song = any("song" in h or "title" in h for h in headers_lower)
            has_artist = any("artist" in h for h in headers_lower)
            record("Tracklist sheet has Track_Number, Song_Title, Artist columns",
                   has_track_num and has_song and has_artist,
                   f"Headers found: {headers}")
        else:
            record("Tracklist sheet has >= 8 data rows", False, "No Tracklist sheet found")
            record("Tracklist sheet has required columns", False, "No Tracklist sheet found")

        # Check Artist_Summary sheet
        summary_sheet = None
        for name in wb.sheetnames:
            if "artist" in name.lower() or "summary" in name.lower():
                summary_sheet = wb[name]
                break

        record("Artist_Summary sheet exists", summary_sheet is not None,
               f"Sheets: {wb.sheetnames}")

        if summary_sheet:
            data_rows = [r for r in summary_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Artist_Summary sheet has >= 4 data rows", len(data_rows) >= 4,
                   f"Found {len(data_rows)} rows")
        else:
            record("Artist_Summary sheet has >= 4 data rows", False, "Sheet not found")

        # --- Groundtruth XLSX value comparison ---
        gt_path = os.path.join(groundtruth_workspace, "Afrobeat_Tracklist.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                       f"Expected {len(gt_rows)}, got {len(a_rows)}")
                for ri in range(min(3, len(gt_rows))):
                    if ri >= len(a_rows):
                        break
                    ok = True
                    for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                        gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                        if gv is None:
                            continue
                        if isinstance(gv, (int, float)):
                            if not num_close(av, gv, max(abs(gv) * 0.1, 1.0)):
                                ok = False
                                break
                        else:
                            if not str_match(av, gv):
                                ok = False
                                break
                    record(f"GT '{gt_sname}' row {ri+1} values", ok,
                           f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'missing'}")
            gt_wb.close()

    except Exception as e:
        record("Tracklist sheet has >= 8 data rows", False, str(e))
        record("Tracklist sheet has required columns", False, str(e))
        record("Artist_Summary sheet exists", False, str(e))
        record("Artist_Summary sheet has >= 4 data rows", False, str(e))


def check_word(agent_workspace):
    print("\n=== Check 5-6: Curator_Notes.docx ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("curator" in fname.lower() or "notes" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        # Try any docx
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Curator_Notes.docx exists", docx_path is not None,
           f"No matching docx in {agent_workspace}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains music-related keywords", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        keywords = ["track", "song", "artist", "afrobeat", "mix", "music"]
        found = [k for k in keywords if k in full_text]
        record("Word doc contains music-related keywords", len(found) >= 3,
               f"Found keywords: {found}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains music-related keywords", False, str(e))


def check_notion():
    print("\n=== Check 7: Notion page with Afrobeat/Mix in title ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT id, properties FROM notion.pages
        ORDER BY created_time DESC
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    found = False
    for row_id, props in rows:
        props_str = json.dumps(props).lower() if props else ""
        if any(k in props_str for k in ["afrobeat", "mix", "tracklist"]):
            found = True
            break
    record("Notion page exists with Afrobeat/Mix/Tracklist in title",
           found, f"Total notion pages: {len(rows)}")


def check_email():
    print("\n=== Check 8: Email sent to music@label.com ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else to_addr.lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "music@label.com" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to music@label.com", matching is not None,
           f"Total messages found: {len(messages)}")
    if matching:
        subj = matching[0] or ""
        body = matching[3] or ""
        has_content = any(k in (subj + " " + body).lower()
                          for k in ["afrobeat", "tracklist", "track", "artist", "mix"])
        record("Email mentions Afrobeat/tracklist content", has_content,
               f"Subject: {subj}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    agent_ws = args.agent_workspace

    check_excel(agent_ws, args.groundtruth_workspace)
    check_word(agent_ws)
    check_notion()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
