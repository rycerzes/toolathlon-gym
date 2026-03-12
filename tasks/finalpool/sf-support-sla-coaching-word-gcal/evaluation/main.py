"""Evaluation for sf-support-sla-coaching-word-gcal."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    agent_ws = args.agent_workspace or task_root

    all_errors = []

    # --- Check 1: Word document ---
    print("Checking Word document...")
    docx_path = os.path.join(agent_ws, "SLA_Coaching_Plan.docx")
    if not os.path.exists(docx_path):
        all_errors.append("SLA_Coaching_Plan.docx not found in agent workspace")
    else:
        from docx import Document
        doc = Document(docx_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Also get table text
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text

        if "SLA Performance Coaching Plan" not in full_text and "sla performance coaching" not in full_text.lower():
            all_errors.append("Word doc missing title 'SLA Performance Coaching Plan'")

        for priority in ["High", "Medium", "Low"]:
            if priority not in full_text:
                all_errors.append(f"Word doc missing priority: {priority}")

        for ticket_count in ["6466", "15774", "9348"]:
            if ticket_count not in full_text:
                all_errors.append(f"Word doc missing ticket count: {ticket_count}")

        if "recommendation" not in full_text.lower() and "coaching" not in full_text.lower():
            all_errors.append("Word doc missing coaching recommendations section")

        print(f"    Word doc checks done")

    # --- Check 2: GCal events ---
    print("Checking Google Calendar events...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        quarters = [
            ("Q1", "2026-03-15"),
            ("Q2", "2026-06-15"),
            ("Q3", "2026-09-15"),
            ("Q4", "2026-12-15"),
        ]
        for q, date in quarters:
            cur.execute("""
                SELECT COUNT(*) FROM gcal.events
                WHERE LOWER(summary) LIKE %s
                AND start_datetime::date = %s
            """, (f"%{q.lower()}%", date))
            count = cur.fetchone()[0]
            if count == 0:
                all_errors.append(f"GCal event for {q} SLA Coaching Review on {date} not found")
            else:
                print(f"    {q} event found")

        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking GCal events: {e}")

    # --- Check 3: Email sent ---
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(*) FROM email.messages
            WHERE to_addr::text ILIKE '%support_manager@company.com%'
        """)
        count = cur.fetchone()[0]
        cur.close()
        conn.close()
        if count == 0:
            all_errors.append("No email sent to support_manager@company.com")
        else:
            print(f"    Email found ({count} messages)")
    except Exception as e:
        all_errors.append(f"Error checking email: {e}")

    # --- Final result ---
    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
