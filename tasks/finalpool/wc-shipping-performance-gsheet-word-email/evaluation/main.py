"""Evaluation for wc-shipping-performance-gsheet-word-email.

Checks:
1. Google Sheet "Shipping Performance Dashboard" with zone data (at least 3 rows)
2. Word doc Shipping_Performance_Report.docx with heading, table, and paragraphs
3. Email to logistics@store.example.com from operations@store.example.com
   Subject: "Monthly Shipping Performance Report"
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected shipping data (from actual WC DB query)
EXPECTED_ZONES = [
    {"name": "standard shipping", "orders": 64, "revenue": 25369.91},
    {"name": "free shipping", "orders": 34, "revenue": 12734.76},
    {"name": "express shipping", "orders": 30, "revenue": 12673.21},
]


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=5.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%shipping%'
    """)
    sheets = cur.fetchall()
    check("Shipping Performance Dashboard spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if sheets:
        ss_id = sheets[0][0]
        cur.execute("""
            SELECT c.value FROM gsheet.cells c
            JOIN gsheet.sheets s ON c.spreadsheet_id = s.spreadsheet_id AND c.sheet_id = s.id
            WHERE c.spreadsheet_id = %s
        """, (ss_id,))
        cells = cur.fetchall()
        all_values = " ".join(str(c[0]) for c in cells if c[0])
        all_lower = all_values.lower()

        # Check at least 3 zone rows
        zone_names_found = sum(
            1 for z in EXPECTED_ZONES
            if any(part in all_lower for part in z["name"].split())
        )
        check("GSheet has data for at least 3 shipping methods",
              zone_names_found >= 3,
              f"Found {zone_names_found}/3 shipping method names")
        check("GSheet contains Standard Shipping data",
              "standard" in all_lower, "Standard not found")
        check("GSheet contains Express Shipping data",
              "express" in all_lower, "Express not found")
        check("GSheet contains revenue data",
              "25369" in all_values or "12734" in all_values or "12673" in all_values,
              "Revenue values not found")

    cur.close()
    conn.close()


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Shipping_Performance_Report.docx")
    check("Shipping_Performance_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}")
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Shipping Performance Report' heading",
          any("shipping performance report" in h for h in headings) or
          "shipping performance report" in full_text,
          "Heading not found")

    para_count = len([p for p in doc.paragraphs if p.text.strip()])
    check("Word doc has at least 3 paragraphs", para_count >= 3,
          f"Found {para_count} non-empty paragraphs")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        table = doc.tables[0]
        table_text = " ".join(
            cell.text.lower()
            for row in table.rows
            for cell in row.cells
        )
        check("Table contains Zone_Name column",
              "zone" in table_text or "name" in table_text, "Column not found")
        check("Table contains shipping method data",
              "standard" in table_text or "express" in table_text,
              "Shipping data not in table")

    check("Word doc mentions Standard Shipping",
          "standard" in full_text, "Standard shipping not mentioned")
    check("Word doc mentions Express Shipping",
          "express" in full_text, "Express shipping not mentioned")


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "logistics@store.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to logistics@store.example.com", found is not None)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from operations@store.example.com",
              "operations@store.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        check("Subject contains 'Monthly Shipping Performance Report'",
              "shipping" in (subj or "").lower() and "performance" in (subj or "").lower(),
              f"Subject: {subj}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("WC SHIPPING PERFORMANCE GSHEET WORD EMAIL - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
