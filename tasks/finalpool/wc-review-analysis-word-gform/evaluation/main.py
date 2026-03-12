"""Evaluation for wc-review-analysis-word-gform."""
import argparse
import json
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

FILE_PASS = 0
FILE_FAIL = 0
DB_PASS = 0
DB_FAIL = 0


def check(name, condition, detail="", db=False):
    global FILE_PASS, FILE_FAIL, DB_PASS, DB_FAIL
    if condition:
        if db:
            DB_PASS += 1
        else:
            FILE_PASS += 1
        print(f"  [PASS] {name}")
    else:
        if db:
            DB_FAIL += 1
        else:
            FILE_FAIL += 1
        d = (detail[:300]) if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def num_close(a, b, tol=0.5):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_word_doc(agent_workspace, groundtruth_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Review_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "review" in p.text.lower() and "analysis" in p.text.lower():
            has_heading = True
            break
    check("Document has review analysis heading", has_heading)

    # Check tables
    check("Document has at least 3 tables", len(doc.tables) >= 3,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 3:
        return False

    # Load groundtruth
    gt_file = os.path.join(groundtruth_workspace, "Review_Data.xlsx")
    if not os.path.isfile(gt_file):
        check("Groundtruth file exists", False)
        return False

    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check category table (table 0)
    gt_cats = list(gt_wb["By Category"].iter_rows(min_row=2, values_only=True))
    table1 = doc.tables[0]
    cat_rows = []
    for row in table1.rows[1:]:
        cells = [cell.text.strip() for cell in row.cells]
        cat_rows.append(cells)

    check("Category table has 8 rows", len(cat_rows) == 8, f"Got {len(cat_rows)} rows")

    for gt_row in gt_cats:
        cat_name, review_count, avg_rating, pos_count, neg_count = gt_row
        matched = None
        for r in cat_rows:
            if r and cat_name.lower() in r[0].lower():
                matched = r
                break
        if matched:
            # Check avg rating
            found_rating = False
            for cell in matched[1:]:
                try:
                    val = float(cell.replace(",", ""))
                    if num_close(val, avg_rating, 0.1):
                        found_rating = True
                        break
                except (ValueError, AttributeError):
                    continue
            check(f"Category {cat_name} avg rating", found_rating,
                  f"Expected ~{avg_rating}")
        else:
            check(f"Category {cat_name} found in table", False)

    # Check summary
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    has_summary = "total" in full_text or "overall" in full_text or "categories" in full_text
    check("Document has summary text", has_summary)

    # Check top/worst products sections
    has_top = False
    has_worst = False
    for p in doc.paragraphs:
        if "top" in p.text.lower() and "rated" in p.text.lower():
            has_top = True
        if "lowest" in p.text.lower() or "worst" in p.text.lower():
            has_worst = True
    check("Has top rated products section", has_top)
    check("Has lowest rated products section", has_worst)

    return True


def check_gform():
    """Check Google Form (non-blocking)."""
    print("\n=== Checking Google Form (non-blocking) ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()
    cur.execute("SELECT form_id, title, question_type FROM gform.questions")
    questions = cur.fetchall()
    cur.close()
    conn.close()

    check("At least 1 form created", len(forms) >= 1,
          f"Found {len(forms)}", db=True)

    if forms:
        found_feedback = any("feedback" in (f[1] or "").lower() or "improvement" in (f[1] or "").lower()
                            for f in forms)
        check("Form title mentions feedback or improvement", found_feedback,
              f"Forms: {[f[1] for f in forms]}", db=True)

    check("At least 3 questions created", len(questions) >= 3,
          f"Found {len(questions)}", db=True)

    if questions:
        q_texts = " ".join((q[1] or "") for q in questions).lower()
        has_category_q = "category" in q_texts
        has_rating_q = "experience" in q_texts or "rate" in q_texts
        has_suggestion_q = "improve" in q_texts or "suggest" in q_texts
        check("Has category question", has_category_q, db=True)
        check("Has rating question", has_rating_q, db=True)
        check("Has suggestion question", has_suggestion_q, db=True)


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("WC REVIEW ANALYSIS WORD GFORM - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace, gt_dir)
    check_gform()

    total_pass = FILE_PASS + DB_PASS
    total_fail = FILE_FAIL + DB_FAIL
    file_ok = FILE_FAIL == 0

    print(f"\n=== SUMMARY ===")
    print(f"  File checks - Passed: {FILE_PASS}, Failed: {FILE_FAIL}")
    print(f"  DB checks   - Passed: {DB_PASS}, Failed: {DB_FAIL}")
    if DB_FAIL > 0:
        print(f"  WARNING: {DB_FAIL} DB checks failed (not blocking)")
    print(f"  Overall: {'PASS' if file_ok else 'FAIL'}")

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump({"passed": total_pass, "failed": total_fail, "success": file_ok}, f, indent=2)

    sys.exit(0 if file_ok else 1)


if __name__ == "__main__":
    main()
