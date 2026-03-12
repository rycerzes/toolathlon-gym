#!/usr/bin/env python3
"""Evaluation script for customer-satisfaction-analytics-framework task validation"""

from argparse import ArgumentParser
import json
import os
import sys

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_xlsx_content(workspace, groundtruth_workspace="."):
    print("\n=== Check: XLSX files ===")
    import openpyxl
    for fname in ["satisfaction_analysis.xlsx", "action_plans.xlsx"]:
        xlsx_path = os.path.join(workspace, fname)
        if not os.path.isfile(xlsx_path):
            record(f"xlsx {fname} exists", False, "Not found")
            continue
        record(f"xlsx {fname} exists", True)
        try:
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                record(f"xlsx {fname} '{ws.title}' has data", len(rows) >= 2, f"{len(rows)} rows")
        except Exception as e:
            record(f"xlsx {fname} readable", False, str(e))
            continue

        # --- Groundtruth value comparison for this file ---
        gt_path = os.path.join(groundtruth_workspace, fname)
        if not os.path.isfile(gt_path):
            wb.close()
            continue

        gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
        for gt_sheet_name in gt_wb.sheetnames:
            gt_ws = gt_wb[gt_sheet_name]
            agent_ws = None
            for asn in wb.sheetnames:
                if asn.strip().lower() == gt_sheet_name.strip().lower():
                    agent_ws = wb[asn]
                    break
            if agent_ws is None:
                record(f"GT {fname} sheet '{gt_sheet_name}' exists in agent", False,
                       f"Available: {wb.sheetnames}")
                continue

            gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
            agent_rows = [r for r in agent_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]

            record(f"GT {fname} '{gt_sheet_name}' row count", len(agent_rows) == len(gt_rows),
                   f"Expected {len(gt_rows)}, got {len(agent_rows)}")

            check_indices = list(range(min(3, len(gt_rows))))
            if len(gt_rows) > 3:
                check_indices.append(len(gt_rows) - 1)
            for idx in check_indices:
                gt_row = gt_rows[idx]
                if idx < len(agent_rows):
                    a_row = agent_rows[idx]
                    row_ok = True
                    for col_idx in range(min(len(gt_row), len(a_row) if a_row else 0)):
                        gt_val = gt_row[col_idx]
                        a_val = a_row[col_idx]
                        if gt_val is None:
                            continue
                        if isinstance(gt_val, (int, float)):
                            ok = num_close(a_val, gt_val, max(abs(gt_val) * 0.1, 1.0))
                        else:
                            ok = str_match(a_val, gt_val)
                        if not ok:
                            record(f"GT {fname} '{gt_sheet_name}' row {idx+1} col {col_idx+1}",
                                   False, f"Expected {gt_val}, got {a_val}")
                            row_ok = False
                            break
                    if row_ok:
                        record(f"GT {fname} '{gt_sheet_name}' row {idx+1} values match", True)
                else:
                    record(f"GT {fname} '{gt_sheet_name}' row {idx+1} exists", False, "Row missing in agent")
        gt_wb.close()
        wb.close()


def check_docx_content(workspace):
    print("\n=== Check: DOCX files ===")
    from docx import Document
    for fname in ["satisfaction_report.docx", "executive_summary.docx"]:
        path = os.path.join(workspace, fname)
        if not os.path.isfile(path):
            record(f"docx {fname} exists", False, "Not found")
            continue
        record(f"docx {fname} exists", True)
        try:
            doc = Document(path)
            record(f"docx {fname} has content", len(doc.paragraphs) > 0, f"{len(doc.paragraphs)} paragraphs")
        except Exception as e:
            record(f"docx {fname} readable", False, str(e))


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False)
    args = parser.parse_args()

    ws = args.agent_workspace
    if not os.path.isdir(ws):
        print(f"Agent workspace not found: {ws}")
        sys.exit(1)

    check_xlsx_content(ws, args.groundtruth_workspace)
    check_docx_content(ws)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
