"""
Evaluation for yt-transcript-music-schedule-gcal-gsheet-word task.

Checks:
1. Radio_Show_Script.docx (or similar) exists in agent_workspace
2. Word doc has >= 3 heading-level sections
3. Word doc text contains Afrobeat and >= 2 artist names
4. GCal has >= 3 new events with Afrobeat/Show in summary in April 2026
5. GSheet has spreadsheet with Playlist sheet containing >= 8 rows
6. Email sent to station@radioafrica.fm
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0
ARTIST_NAMES = ["burna boy", "wizkid", "rema", "davido", "ckay", "asake", "ayra starr",
                "fireboy", "tems", "omah lay", "kizz daniel"]


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Check 1-3: Radio Show Script Word Doc ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("script" in fname.lower() or "show" in fname.lower() or "radio" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Radio show script Word doc exists", docx_path is not None,
           f"No script/show/radio docx found in {agent_workspace}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains Afrobeat and artist names", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        has_afrobeat = "afrobeat" in full_text
        found_artists = [a for a in ARTIST_NAMES if a in full_text]
        record("Word doc contains Afrobeat and >= 2 artist names",
               has_afrobeat and len(found_artists) >= 2,
               f"Afrobeat: {has_afrobeat}, artists found: {found_artists[:5]}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains Afrobeat and artist names", False, str(e))


def check_gcal():
    print("\n=== Check 4: GCal Afrobeat Show Events in April 2026 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        AND (summary ILIKE '%afrobeat%' OR summary ILIKE '%show%')
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    record("GCal has >= 3 Afrobeat/Show events in April 2026",
           len(events) >= 3, f"Found {len(events)} matching events")

    if events:
        # Check Sunday evening (20:00 UTC or within evening hours)
        sunday_events = [e for e in events if e[1] and e[1].weekday() == 6]
        record("Events are on Sundays", len(sunday_events) >= 3,
               f"Sunday events: {len(sunday_events)}")


def check_gsheet():
    print("\n=== Check 5: GSheet Radio_Broadcast_Schedule ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE title ILIKE '%radio%' OR title ILIKE '%broadcast%' OR title ILIKE '%schedule%'
        ORDER BY created_at DESC LIMIT 1
    """)
    spreadsheet = cur.fetchone()

    if not spreadsheet:
        cur.execute("SELECT id, title FROM gsheet.spreadsheets ORDER BY created_at DESC LIMIT 1")
        spreadsheet = cur.fetchone()

    record("Radio_Broadcast_Schedule spreadsheet exists", spreadsheet is not None,
           "No matching spreadsheet found")

    if spreadsheet:
        spreadsheet_id, title = spreadsheet
        cur.execute("""
            SELECT id, title FROM gsheet.sheets
            WHERE spreadsheet_id = %s AND title ILIKE '%%playlist%%'
            LIMIT 1
        """, (spreadsheet_id,))
        playlist_sheet = cur.fetchone()

        record("Playlist sheet exists in spreadsheet", playlist_sheet is not None,
               f"Sheets in spreadsheet: {title}")

        if playlist_sheet:
            sheet_id = playlist_sheet[0]
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                AND value IS NOT NULL AND value != ''
            """, (spreadsheet_id, sheet_id))
            data_rows = cur.fetchone()[0]
            record("Playlist sheet has >= 8 data rows", data_rows >= 8,
                   f"Found {data_rows} data rows")
        else:
            record("Playlist sheet has >= 8 data rows", False, "No Playlist sheet")

        # Check Show_Schedule sheet
        cur.execute("""
            SELECT id, title FROM gsheet.sheets
            WHERE spreadsheet_id = %s AND (title ILIKE '%%show%%' OR title ILIKE '%%schedule%%')
            LIMIT 1
        """, (spreadsheet_id,))
        show_sheet = cur.fetchone()
        record("Show_Schedule sheet exists", show_sheet is not None,
               "No Show_Schedule sheet found")

        if show_sheet:
            sheet_id = show_sheet[0]
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
                AND value IS NOT NULL AND value != ''
            """, (spreadsheet_id, sheet_id))
            data_rows = cur.fetchone()[0]
            record("Show_Schedule sheet has >= 4 rows (April Sundays)", data_rows >= 4,
                   f"Found {data_rows} rows")
        else:
            record("Show_Schedule sheet has >= 4 rows", False, "Sheet not found")
    else:
        for chk in ["Playlist sheet exists", "Playlist sheet has >= 8 data rows",
                    "Show_Schedule sheet exists", "Show_Schedule sheet has >= 4 rows"]:
            record(chk, False, "No spreadsheet found")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 6: Email to station@radioafrica.fm ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else to_addr.lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "station@radioafrica.fm" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to station@radioafrica.fm", matching is not None,
           f"Total messages: {len(messages)}")
    if matching:
        body = (matching[0] or "") + " " + (matching[3] or "")
        has_content = any(k in body.lower() for k in ["schedule", "broadcast", "afrobeat", "sunday", "show"])
        record("Email mentions show schedule content", has_content,
               f"Subject: {matching[0]}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_xlsx_content(workspace, groundtruth_workspace="."):
    """Check Radio_Broadcast_Schedule_local.xlsx has valid content."""
    print("\n=== Checking XLSX Content ===")
    try:
        import openpyxl
    except ImportError:
        record("openpyxl available", False, "Cannot import openpyxl")
        return False

    xlsx_path = os.path.join(workspace, "Radio_Broadcast_Schedule_local.xlsx")
    if not os.path.isfile(xlsx_path):
        record("Radio_Broadcast_Schedule_local.xlsx exists", False, f"Not found: {xlsx_path}")
        return False
    record("Radio_Broadcast_Schedule_local.xlsx exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        record("XLSX has at least one sheet", len(wb.worksheets) >= 1,
               f"Found {len(wb.worksheets)} sheets")
        all_ok = True
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            has_data = len(rows) >= 2
            record(f"XLSX sheet '{ws.title}' has data rows", has_data,
                   f"Only {len(rows)} rows")
            if not has_data:
                all_ok = False
        # --- Groundtruth XLSX value comparison ---
        gt_path = os.path.join(groundtruth_workspace, "Radio_Broadcast_Schedule_local.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                       f"Expected {len(gt_rows)}, got {len(a_rows)}")
                for ri in range(min(3, len(gt_rows))):
                    if ri >= len(a_rows):
                        break
                    ok = True
                    for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                        gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                        if gv is None:
                            continue
                        if isinstance(gv, (int, float)):
                            if not num_close(av, gv, max(abs(gv) * 0.1, 1.0)):
                                ok = False
                                break
                        else:
                            if not str_match(av, gv):
                                ok = False
                                break
                    record(f"GT '{gt_sname}' row {ri+1} values", ok,
                           f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'missing'}")
            gt_wb.close()

        wb.close()
        return all_ok
    except Exception as e:
        record("XLSX readable", False, str(e))
        return False


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gcal()
    check_gsheet()
    check_email()
    check_xlsx_content(args.agent_workspace, args.groundtruth_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
