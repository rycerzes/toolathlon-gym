"""
Evaluation for canvas-assignment-word-notion task.

Checks:
1. Word document Assignment_Guide.docx with correct assignment data
2. Notion page created (non-blocking)
"""

import argparse
import json
import os
import sys

import psycopg2

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)


DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected assignments for course_id=7 (Creative Computing & Culture Fall 2014)
EXPECTED_ASSIGNMENT_COUNT = 10
EXPECTED_TOTAL_POINTS = 300.0
EXPECTED_ASSIGNMENT_NAMES = [
    "CMA 24295", "CMA 24296", "CMA 24297", "CMA 24298",
    "TMA 24291", "TMA 24292", "TMA 24293", "TMA 24294",
    "Final Exam 24299", "Final Exam 40088"
]


def check(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Assignment_Guide.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Assignment_Guide.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "creative computing" in p.text.lower() and "assignment guide" in p.text.lower():
            has_heading = True
            break
    check("Document has correct heading", has_heading)

    # Check course code
    full_text = " ".join(p.text for p in doc.paragraphs)
    check("Document mentions CCC-2014J", "CCC-2014J" in full_text)

    # Check table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    check("Table has 10 assignment rows", len(data_rows) == EXPECTED_ASSIGNMENT_COUNT,
          f"Got {len(data_rows)} rows")

    # Check assignment names appear
    found_names = 0
    row_texts = " ".join(str(cell) for row in data_rows for cell in row)
    for name in EXPECTED_ASSIGNMENT_NAMES:
        if name in row_texts:
            found_names += 1
    check("At least 8/10 assignment names found", found_names >= 8,
          f"Found {found_names}/10")

    # Check some due dates
    has_dates = sum(1 for r in data_rows if "2014" in r[1] or "2015" in r[1])
    has_no_date = sum(1 for r in data_rows if "no due date" in r[1].lower())
    check("Table has assignments with dates", has_dates >= 8, f"Found {has_dates}")
    check("Table has assignments without dates", has_no_date >= 1, f"Found {has_no_date}")

    # Check total assignments and points in text
    check("Document mentions total assignments count",
          "10" in full_text and "total assignments" in full_text.lower(),
          "Expected 'Total Assignments: 10'")

    check("Document mentions total points",
          "300" in full_text and "total points" in full_text.lower(),
          "Expected 'Total Points: 300.0'")

    return True


def check_notion():
    """Check Notion page - NON-BLOCKING."""
    print("\n=== Checking Notion (non-blocking) ===")

    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()

        cur.execute("SELECT id, properties FROM notion.pages")
        pages = cur.fetchall()
        cur.close()
        conn.close()

        if len(pages) == 0:
            print("  [WARN] No Notion pages found (non-blocking)")
            return

        found = False
        for page_id, props in pages:
            props_str = json.dumps(props) if isinstance(props, dict) else str(props)
            if "CCC-2014J" in props_str or "assignment" in props_str.lower():
                found = True
                print(f"  [INFO] Found relevant Notion page: {page_id}")
                break

        if found:
            print("  [INFO] Notion page with assignment overview exists")
        else:
            print("  [WARN] No matching Notion page found (non-blocking)")

    except Exception as e:
        print(f"  [WARN] Notion check error (non-blocking): {e}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("CANVAS ASSIGNMENT WORD NOTION - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_notion()  # Non-blocking

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
