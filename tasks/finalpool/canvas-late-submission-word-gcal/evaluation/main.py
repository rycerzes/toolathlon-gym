"""Evaluation for canvas-late-submission-word-gcal."""
import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

# Top 3 courses by late submission rate (from actual DB data)
TOP3_COURSES = [
    "Creative Computing & Culture (Spring 2014)",
    "Creative Computing & Culture (Fall 2014)",
    "Data-Driven Design (Spring 2013)",
]

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Late_Submission_Report.docx")
    if not os.path.isfile(docx_path):
        check("Late_Submission_Report.docx exists", False, f"Not found: {docx_path}")
        return
    check("Late_Submission_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        check("Word doc has meaningful content (>= 100 chars)",
              len(all_text.strip()) >= 100,
              f"Content length: {len(all_text)}")
        check("Word doc contains 'late submission'",
              "late submission" in all_text or "late_submission" in all_text,
              f"Sample: {all_text[:300]}")
        check("Word doc contains course names",
              "creative computing" in all_text or "data-driven design" in all_text,
              f"Sample: {all_text[:300]}")
        check("Word doc contains late rate percentages",
              "60" in all_text or "58" in all_text or "56" in all_text,
              f"Sample: {all_text[:300]}")
        check("Word doc has recommendations section",
              "recommend" in all_text or "policy" in all_text or "intervention" in all_text,
              f"Sample: {all_text[:300]}")
    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, description, start_datetime
            FROM gcal.events
            WHERE LOWER(summary) LIKE '%%late submission%%'
               OR LOWER(summary) LIKE '%%late submission review%%'
        """)
        events = cur.fetchall()
        check("At least 3 'Late Submission Review' events scheduled",
              len(events) >= 3,
              f"Found {len(events)} events")

        if events:
            # Check events are on 2026-04-01
            on_target_date = [e for e in events if e[2] and '2026-04-01' in str(e[2])]
            check("Events scheduled on 2026-04-01",
                  len(on_target_date) >= 3,
                  f"{len(on_target_date)} events on 2026-04-01 out of {len(events)} total")

            # Check events contain course names or relevant context
            event_texts = " ".join(
                (str(e[0]) + " " + str(e[1] or "")).lower() for e in events
            )
            check("Events mention relevant courses",
                  "creative computing" in event_texts or "data-driven" in event_texts or "late" in event_texts,
                  f"Event summary/desc sample: {event_texts[:300]}")

        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%registrar@university.edu%%'
               OR subject ILIKE '%%late submission%%'
        """)
        emails = cur.fetchall()
        check("Email sent to registrar@university.edu", len(emails) >= 1,
              "No matching email found")
        if emails:
            email = emails[0]
            subject = str(email[1]).lower() if email[1] else ""
            check("Email subject contains 'late' or 'submission'",
                  "late" in subject or "submission" in subject,
                  f"Subject: {email[1]}")
            body = str(email[3]) if email[3] else ""
            check("Email body has content", len(body) > 30,
                  f"Body length: {len(body)}")
        cur.close()
        conn.close()
    except Exception as e:
        check("Email check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_calendar()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
