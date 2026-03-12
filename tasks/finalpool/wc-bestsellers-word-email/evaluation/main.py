"""
Evaluation for wc-bestsellers-word-email task.

Checks:
1. Word document Bestsellers_Report.docx with correct top 10 products
2. Email sent to sales-leads@store.com (non-blocking DB check)
"""
import argparse
import sys
import os
from pathlib import Path

FILE_PASS = 0
FILE_FAIL = 0
DB_PASS = 0
DB_FAIL = 0


def check(name, condition, detail="", db=False):
    global FILE_PASS, FILE_FAIL, DB_PASS, DB_FAIL
    if condition:
        if db:
            DB_PASS += 1
        else:
            FILE_PASS += 1
        print(f"  [PASS] {name}")
    else:
        if db:
            DB_FAIL += 1
        else:
            FILE_FAIL += 1
        d = detail[:300] if len(detail) > 300 else detail
        print(f"  [FAIL] {name}: {d}")


def get_expected_top10():
    """Query PostgreSQL for top 10 products by total_sales."""
    import psycopg2

    conn = psycopg2.connect(
        host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym",
        user="eigent", password="camel"
    )
    cur = conn.cursor()
    cur.execute("""
        SELECT name, total_sales::int as total_sold, price::float as price
        FROM wc.products
        ORDER BY total_sales::int DESC
        LIMIT 10
    """)
    products = cur.fetchall()
    conn.close()
    return products


def check_word(workspace, expected):
    """Check Bestsellers_Report.docx for correctness."""
    from docx import Document

    print("\n=== Checking Word Document ===")
    doc_path = Path(workspace) / "Bestsellers_Report.docx"

    if not doc_path.exists():
        check("Word file exists", False, f"Not found: {doc_path}")
        return
    check("Word file exists", True)

    doc = Document(str(doc_path))

    # Check for title heading
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    check("Has heading with 'best' or 'bestsell'",
          any("best" in h.lower() for h in headings),
          f"Headings: {[h[:50] for h in headings]}")

    # Check tables
    if len(doc.tables) == 0:
        check("Has a table", False, "No tables found")
        return
    check("Has a table", True)

    table = doc.tables[0]
    rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        rows.append(cells)

    check("Table has 10 data rows", len(rows) == 10, f"Got {len(rows)} rows")

    # Check top product matches
    if len(rows) >= 1:
        top_name_in_doc = rows[0][1] if len(rows[0]) > 1 else ""
        exp_top_name = expected[0][0]
        # Compare first 30 chars
        check("Top product name matches",
              top_name_in_doc[:30].lower() == exp_top_name[:30].lower(),
              f"Expected '{exp_top_name[:50]}', got '{top_name_in_doc[:50]}'")

    # Check total_sold values
    for i, (exp_name, exp_sold, exp_price) in enumerate(expected):
        if i >= len(rows):
            break
        row = rows[i]
        if len(row) >= 3:
            try:
                actual_sold = int(row[2])
                check(f"Rank {i+1} Total_Sold",
                      actual_sold == exp_sold,
                      f"Expected {exp_sold}, got {actual_sold}")
            except ValueError:
                check(f"Rank {i+1} Total_Sold", False, f"Cannot parse: {row[2]}")

    # Check summary section
    total_sold_all = sum(p[1] for p in expected)
    check("Summary mentions total units",
          str(total_sold_all) in full_text,
          f"Expected {total_sold_all} in text")


def check_email(expected):
    """Check email was sent (non-blocking DB check)."""
    import psycopg2

    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym",
            user="eigent", password="camel"
        )
        cur = conn.cursor()
    except Exception as e:
        check("DB connection", False, str(e), db=True)
        return

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    target = "sales-leads@store.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        to_str = str(to_addr or "").lower()
        if target in to_str:
            found = (subj, from_addr, to_addr, body)
            break

    check(f"Email sent to {target}", found is not None,
          f"Found {len(all_emails)} total emails", db=True)

    if found:
        subj, _, _, body = found
        body_lower = (body or "").lower()
        subj_lower = (subj or "").lower()

        check("Email subject mentions bestsellers/top",
              "bestseller" in subj_lower or "top" in subj_lower or "best" in subj_lower,
              f"Subject: {(subj or '')[:100]}", db=True)

        top_name_short = expected[0][0][:20].lower()
        check("Email body mentions top product",
              top_name_short in body_lower,
              f"Expected '{top_name_short}' in body", db=True)

        check("Email body mentions top product units",
              str(expected[0][1]) in (body or ""),
              f"Expected '{expected[0][1]}' in body", db=True)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("Fetching expected data...")
    expected = get_expected_top10()
    print(f"  Top 10 products loaded")

    check_word(args.agent_workspace, expected)
    check_email(expected)

    file_ok = FILE_FAIL == 0

    print(f"\n=== SUMMARY ===")
    print(f"  File checks - Passed: {FILE_PASS}, Failed: {FILE_FAIL}")
    print(f"  DB checks   - Passed: {DB_PASS}, Failed: {DB_FAIL}")
    if DB_FAIL > 0:
        print(f"  WARNING: {DB_FAIL} DB checks failed (not blocking)")
    print(f"  Overall: {'PASS' if file_ok else 'FAIL'}")

    if file_ok:
        print("\nPass all tests!")
        sys.exit(0)
    else:
        print("\nSome checks failed.")
        sys.exit(1)
