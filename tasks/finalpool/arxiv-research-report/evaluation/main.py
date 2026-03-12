"""
Evaluation script for arxiv-research-report task.
Checks that LLM_Reasoning_Survey.docx and .pdf exist with expected content.

Usage:
  python -m evaluation.main --agent_workspace <path> --groundtruth_workspace <path> --launch_time <time>
"""
import argparse
import os
import re
import sys


PASS_COUNT = 0
FAIL_COUNT = 0


def check(name: str, condition: bool, detail: str = ""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_truncated = (detail[:200] + "...") if len(detail) > 200 else detail
        print(f"  [FAIL] {name}: {detail_truncated}")


def normalize(text: str) -> str:
    """Normalize text for comparison: lowercase, collapse whitespace."""
    return re.sub(r'\s+', ' ', text.lower().strip())


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", type=str, required=True)
    parser.add_argument("--groundtruth_workspace", type=str, required=True)
    parser.add_argument("--launch_time", type=str, required=False)
    parser.add_argument("--res_log_file", type=str, required=False)
    args = parser.parse_args()

    docx_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.docx")
    pdf_path = os.path.join(args.agent_workspace, "LLM_Reasoning_Survey.pdf")

    # ── Check 1: Word document exists ────────────────────────────────────────
    check("LLM_Reasoning_Survey.docx exists", os.path.exists(docx_path),
          f"File not found at {docx_path}")

    if not os.path.exists(docx_path):
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    # Read the Word document
    try:
        from docx import Document
        doc = Document(docx_path)
        full_text = "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        check("Word document readable", False, str(e))
        print(f"\nResults: {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} passed, {FAIL_COUNT} failed")
        sys.exit(1)

    normalized = normalize(full_text)

    # ── Check 2: Minimum content length ──────────────────────────────────────
    check("Document has at least 500 characters",
          len(full_text.strip()) >= 500,
          f"Document has {len(full_text.strip())} characters")

    # ── Check 3: Title contains "Survey" or "Reasoning" ─────────────────────
    check("Document title mentions survey or reasoning",
          "survey" in normalized or "reasoning" in normalized,
          "Neither 'survey' nor 'reasoning' found")

    # ── Check 4: Has required section headings ───────────────────────────────
    headings = []
    for para in doc.paragraphs:
        if para.style and para.style.name and "heading" in para.style.name.lower():
            headings.append(para.text.lower())

    # Fall back to text-based heading detection if no styled headings
    heading_text = " ".join(headings) if headings else normalized

    has_intro = "introduction" in heading_text or "introduction" in normalized
    has_lit_review = ("literature review" in heading_text or "literature review" in normalized
                      or "literature" in heading_text)
    has_methodology = ("methodology" in heading_text or "methodology comparison" in normalized
                       or "methodology" in normalized)
    has_conclusion = ("conclusion" in heading_text or "conclusion" in normalized
                      or "summary" in normalized)

    check("Has Introduction section", has_intro, "No 'Introduction' heading or text found")
    check("Has Literature Review section", has_lit_review, "No 'Literature Review' heading or text found")
    check("Has Methodology Comparison section", has_methodology, "No 'Methodology' heading or text found")
    check("Has Conclusion section", has_conclusion, "No 'Conclusion' heading or text found")

    # ── Check 5: Has at least 4 headings ─────────────────────────────────────
    # Count headings from styles or text patterns
    heading_count = len(headings)
    if heading_count < 4:
        # Fall back: count lines that look like section headers
        for line in full_text.split("\n"):
            stripped = line.strip()
            if stripped and len(stripped) < 80 and not stripped.endswith("."):
                for kw in ["introduction", "literature", "methodology", "conclusion",
                           "survey", "review", "comparison"]:
                    if kw in stripped.lower():
                        heading_count += 1
                        break
    check("Has at least 4 headings/sections",
          heading_count >= 4,
          f"Found only {heading_count} headings")

    # ── Check 6: All 5 target paper titles present ───────────────────────────
    paper_titles = [
        "Chain-of-Thought Prompting Elicits Reasoning in Large Language Models",
        "Tree of Thoughts: Deliberate Problem Solving with Large Language Models",
        "Self-Consistency Improves Chain of Thought Reasoning in Language Models",
        "Process Supervision for Mathematical Reasoning",
        "Scaling LLM Reasoning with Reinforcement Learning",
    ]
    for title in paper_titles:
        check(f"Paper title present: {title[:60]}",
              title.lower() in normalized,
              "Title not found in document text")

    # ── Check 7: Key author names present ────────────────────────────────────
    key_authors = ["Jason Wei", "Shunyu Yao", "Xuezhi Wang"]
    for author in key_authors:
        check(f"Author present: {author}",
              author.lower() in normalized,
              f"Author '{author}' not found")

    # ── Check 8: Key domain terms present ────────────────────────────────────
    key_terms = ["chain-of-thought", "tree of thoughts", "self-consistency", "process supervision"]
    for term in key_terms:
        check(f"Key term present: {term}",
              term.lower() in normalized,
              f"Term '{term}' not found")

    # ── Check 9: Noise topics NOT prominently featured ───────────────────────
    noise_topics = ["image classification", "federated learning", "protein structure"]
    for topic in noise_topics:
        check(f"Noise topic NOT prominent: {topic}",
              topic.lower() not in normalized,
              f"Noise topic '{topic}' found in document -- should not be included")

    # ── Check 10: PDF file exists and has reasonable size ────────────────────
    check("LLM_Reasoning_Survey.pdf exists", os.path.exists(pdf_path),
          f"PDF not found at {pdf_path}")
    if os.path.exists(pdf_path):
        pdf_size = os.path.getsize(pdf_path)
        check("PDF file size > 5KB",
              pdf_size > 5000,
              f"PDF is only {pdf_size} bytes")

    # ── Summary ──────────────────────────────────────────────────────────────
    total = PASS_COUNT + FAIL_COUNT
    print(f"\nResults: {PASS_COUNT}/{total} passed, {FAIL_COUNT} failed")

    sys.exit(0 if FAIL_COUNT == 0 else 1)


if __name__ == "__main__":
    main()
