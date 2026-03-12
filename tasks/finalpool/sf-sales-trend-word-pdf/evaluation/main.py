"""
Evaluation for sf-sales-trend-word-pdf task.

Checks:
1. Word document Sales_Trend_Analysis.docx exists with correct data
2. PDF file Sales_Trend_Analysis.pdf exists with correct data
"""

import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=500.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def get_expected_data():
    """Fetch expected monthly revenue data from the database."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("""
        SELECT "MONTH_KEY", "TOTAL_REVENUE", "ORDER_COUNT", "UNIQUE_CUSTOMERS", "AVG_ORDER_VALUE"
        FROM sf_data."SALES_DW__ANALYTICS__MONTHLY_REVENUE"
        WHERE "MONTH_KEY" >= '2025-01-01' AND "MONTH_KEY" <= '2025-12-31'
        ORDER BY "MONTH_KEY"
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    return rows


def check_word_doc(agent_workspace):
    """Check the Word document structure and content."""
    print("\n=== Checking Sales_Trend_Analysis.docx ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Not found: {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)
    expected = get_expected_data()

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        if "sales trend" in p.text.lower() and "2025" in p.text:
            has_heading = True
            break
    check("Document has Sales Trend 2025 heading", has_heading)

    # Check table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if len(doc.tables) < 1:
        return False

    table = doc.tables[0]
    data_rows = []
    for row in table.rows[1:]:  # skip header
        cells = [cell.text.strip() for cell in row.cells]
        data_rows.append(cells)

    check("Table has 12 month rows", len(data_rows) == 12, f"Got {len(data_rows)} rows")

    # Verify revenue values for each month
    total_revenue = sum(float(r[1]) for r in expected)
    matched_months = 0
    for exp_row in expected:
        month_str = exp_row[0].strftime("%Y-%m")
        exp_revenue = float(exp_row[1])
        found = False
        for dr in data_rows:
            if month_str in dr[0]:
                # Try to extract revenue value
                rev_text = dr[1].replace("$", "").replace(",", "").strip()
                try:
                    if num_close(float(rev_text), exp_revenue, 100.0):
                        found = True
                except ValueError:
                    pass
                break
        if found:
            matched_months += 1

    check("At least 10/12 months have correct revenue", matched_months >= 10,
          f"Matched {matched_months}/12")

    # Check summary section
    full_text = " ".join(p.text for p in doc.paragraphs)
    check("Document mentions total revenue",
          "total" in full_text.lower() and ("revenue" in full_text.lower() or "annual" in full_text.lower()),
          "Expected total revenue in summary")

    # Check H1/H2 analysis
    has_h1h2 = "h1" in full_text.lower() or "first half" in full_text.lower()
    check("Document has H1/H2 trend analysis", has_h1h2)

    return True


def check_pdf(agent_workspace):
    """Check the PDF file exists and has content."""
    print("\n=== Checking Sales_Trend_Analysis.pdf ===")

    pdf_path = os.path.join(agent_workspace, "Sales_Trend_Analysis.pdf")
    check("PDF file exists", os.path.isfile(pdf_path), f"Not found: {pdf_path}")
    if not os.path.isfile(pdf_path):
        return False

    size = os.path.getsize(pdf_path)
    check("PDF file size reasonable", size > 1024, f"Size: {size} bytes")

    # Try to read PDF text
    try:
        import PyPDF2
        with open(pdf_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
    except ImportError:
        try:
            import pdfplumber
            with pdfplumber.open(pdf_path) as p:
                text = ""
                for page in p.pages:
                    text += page.extract_text() or ""
        except ImportError:
            print("  [WARN] No PDF reader available. Checking file existence only.")
            return True

    text_lower = text.lower()
    check("PDF contains Sales Trend title", "sales trend" in text_lower or "sales" in text_lower)

    # Check some month data appears
    months_found = sum(1 for m in ["2025-01", "2025-06", "2025-12"] if m in text)
    check("PDF contains month data", months_found >= 2,
          f"Found {months_found}/3 sample months")

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES TREND WORD PDF - EVALUATION")
    print("=" * 70)

    check_word_doc(args.agent_workspace)
    check_pdf(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": PASS_COUNT, "failed": FAIL_COUNT, "success": overall}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
