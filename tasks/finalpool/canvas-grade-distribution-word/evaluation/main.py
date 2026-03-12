"""
Evaluation for canvas-grade-distribution-word.
Checks:
1. Excel file Grade_Data.xlsx with correct grade distribution and assignment averages
2. Word document Grade_Report.docx with required content
"""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

COURSE_ID = 16  # FFF-2013J

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def get_expected_data():
    """Compute expected grade distribution from Canvas DB."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT s.score, a.points_possible
        FROM canvas.submissions s
        JOIN canvas.assignments a ON s.assignment_id = a.id
        WHERE a.course_id = %s AND s.score IS NOT NULL
    """, (COURSE_ID,))
    submissions = cur.fetchall()

    bands = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'F': 0}
    for score, points in submissions:
        if points and float(points) > 0:
            pct = (float(score) / float(points)) * 100
        else:
            pct = float(score)
        if pct >= 90:
            bands['A'] += 1
        elif pct >= 80:
            bands['B'] += 1
        elif pct >= 70:
            bands['C'] += 1
        elif pct >= 60:
            bands['D'] += 1
        else:
            bands['F'] += 1

    total = sum(bands.values())

    cur.execute("""
        SELECT a.name,
               ROUND(AVG(s.score)::numeric, 2) as avg_score,
               COUNT(s.id) as submission_count
        FROM canvas.assignments a
        JOIN canvas.submissions s ON s.assignment_id = a.id
        WHERE a.course_id = %s AND s.score IS NOT NULL
        GROUP BY a.id, a.name
        ORDER BY a.name
    """, (COURSE_ID,))
    assignments = cur.fetchall()

    cur.close()
    conn.close()

    return bands, total, assignments


def check_excel(agent_workspace, gt_workspace, expected_bands, expected_total, expected_assignments):
    """Check Grade_Data.xlsx."""
    print("\n=== Checking Excel ===")
    agent_file = os.path.join(agent_workspace, "Grade_Data.xlsx")
    gt_file = os.path.join(gt_workspace, "Grade_Data.xlsx")

    if not os.path.exists(agent_file):
        record("Excel file exists", False, f"Not found: {agent_file}")
        return False
    record("Excel file exists", True)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Distribution sheet
    a_rows = load_sheet_rows(agent_wb, "Distribution")
    g_rows = load_sheet_rows(gt_wb, "Distribution")

    if a_rows is None:
        record("Sheet 'Distribution' exists", False, f"Sheets: {agent_wb.sheetnames}")
        return False
    record("Sheet 'Distribution' exists", True)

    a_data = a_rows[1:] if len(a_rows) > 1 else []
    g_data = g_rows[1:] if len(g_rows) > 1 else []

    record("Distribution has 5 rows", len(a_data) == 5, f"Got {len(a_data)}")

    a_lookup = {}
    for row in a_data:
        if row and row[0]:
            a_lookup[str(row[0]).strip().upper()] = row

    for g_row in g_data:
        if not g_row or not g_row[0]:
            continue
        grade = str(g_row[0]).strip().upper()
        a_row = a_lookup.get(grade)
        if a_row is None:
            record(f"Grade band '{grade}' found", False, "Missing")
            continue
        record(f"Grade band '{grade}' found", True)
        if len(a_row) > 1 and len(g_row) > 1:
            record(f"  {grade} Count",
                   num_close(a_row[1], g_row[1], 10),
                   f"Agent={a_row[1]}, GT={g_row[1]}")
        if len(a_row) > 2 and len(g_row) > 2:
            record(f"  {grade} Percentage",
                   num_close(a_row[2], g_row[2], 0.5),
                   f"Agent={a_row[2]}, GT={g_row[2]}")

    # Check Assignment Averages sheet
    a_rows2 = load_sheet_rows(agent_wb, "Assignment Averages")
    if a_rows2 is None:
        record("Sheet 'Assignment Averages' exists", False, f"Sheets: {agent_wb.sheetnames}")
        return False
    record("Sheet 'Assignment Averages' exists", True)

    g_rows2 = load_sheet_rows(gt_wb, "Assignment Averages")
    a_data2 = a_rows2[1:] if len(a_rows2) > 1 else []
    g_data2 = g_rows2[1:] if len(g_rows2) > 1 else []

    record("Assignment Averages row count",
           abs(len(a_data2) - len(g_data2)) <= 2,
           f"Expected {len(g_data2)}, got {len(a_data2)}")

    # Check a few assignment averages
    a_assign_lookup = {}
    for row in a_data2:
        if row and row[0]:
            a_assign_lookup[str(row[0]).strip().lower()] = row

    for g_row in g_data2[:5]:
        if not g_row or not g_row[0]:
            continue
        name = str(g_row[0]).strip()
        key = name.lower()
        a_row = a_assign_lookup.get(key)
        if a_row is None:
            record(f"Assignment '{name}' found", False, "Missing")
            continue
        record(f"Assignment '{name}' found", True)
        if len(a_row) > 1 and len(g_row) > 1:
            record(f"  Avg_Score",
                   num_close(a_row[1], g_row[1], 2.0),
                   f"Agent={a_row[1]}, GT={g_row[1]}")

    return True


def check_word(agent_workspace, expected_total):
    """Check Grade_Report.docx exists and has required content."""
    print("\n=== Checking Word Document ===")

    docx_path = os.path.join(agent_workspace, "Grade_Report.docx")
    if not os.path.exists(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)

        # Check for heading
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        record("Document mentions FFF-2013J",
               "fff-2013j" in all_text or "fff 2013j" in all_text,
               "No reference to course code")

        record("Document mentions grade distribution",
               "grade" in all_text and "distribution" in all_text,
               "Missing 'grade distribution' text")

        # Check for tables
        record("Document has at least one table",
               len(doc.tables) >= 1,
               f"Found {len(doc.tables)} tables")

        # Check total submissions mentioned
        record("Document mentions total submissions",
               str(expected_total) in all_text or
               str(expected_total) in " ".join(
                   " ".join(cell.text for cell in row.cells)
                   for t in doc.tables for row in t.rows
               ),
               f"Expected {expected_total} to be mentioned")

    except ImportError:
        record("python-docx available", False, "Cannot import docx module")
        return False
    except Exception as e:
        record("Word document readable", False, str(e))
        return False

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    bands, total, assignments = get_expected_data()
    print(f"[eval] Total submissions: {total}")
    print(f"[eval] Grade bands: {bands}")

    excel_ok = check_excel(args.agent_workspace, gt_dir, bands, total, assignments)
    word_ok = check_word(args.agent_workspace, total)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = PASS_COUNT > 0 and FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")

    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
