"""
Evaluation for 12306-beijing-shanghai-trip-notion-gcal-word.

Checks:
1. Travel_Plan.docx exists in agent_workspace
2. Word doc has at least 3 headings
3. Word doc contains G11 and G2 train numbers (or key times 07:00 and 09:05)
4. Word doc mentions Beijing and Shanghai
5. Word doc contains total cost 1106 (two second-class tickets at 553 each)
6. Notion page exists with Beijing/Shanghai/Trip in title
7. GCal has at least 1 new travel event on 2026-03-10 (beyond Client Meeting)
8. Email sent to travel@consulting.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Travel_Plan.docx ===")
    import glob

    # Find any docx with Travel/Plan/Trip in the name
    pattern = os.path.join(agent_workspace, "*.docx")
    all_docx = glob.glob(pattern)
    travel_docs = [f for f in all_docx if any(
        kw in os.path.basename(f).lower()
        for kw in ["travel", "plan", "trip", "beijing", "shanghai"]
    )]

    if not travel_docs:
        record("Travel Plan docx exists", False, f"No matching docx found in {agent_workspace}")
        return
    record("Travel Plan docx exists", True)

    doc_path = travel_docs[0]
    try:
        import docx
        doc = docx.Document(doc_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Word doc has at least 3 headings", len(headings) >= 3,
           f"Found {len(headings)} headings")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()

    has_g11 = "g11" in full_text
    has_g2 = "g2" in full_text or "g 2" in full_text
    has_times = "07:00" in full_text and "09:05" in full_text
    record("Contains train numbers G11 and G2 (or departure times)",
           (has_g11 and has_g2) or has_times,
           f"G11:{has_g11}, G2:{has_g2}, times:{has_times}")

    has_beijing = "beijing" in full_text
    has_shanghai = "shanghai" in full_text
    record("Contains Beijing and Shanghai", has_beijing and has_shanghai,
           f"Beijing:{has_beijing}, Shanghai:{has_shanghai}")

    has_cost = "1106" in full_text or "553" in full_text
    record("Contains ticket cost (553 or 1106)", has_cost, f"Text snippet: {full_text[:200]}")


def check_notion():
    print("\n=== Check 2: Notion page ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT id, properties FROM notion.pages
        WHERE properties::text ILIKE '%beijing%'
           OR properties::text ILIKE '%shanghai%'
           OR properties::text ILIKE '%trip%'
           OR properties::text ILIKE '%travel%'
           OR properties::text ILIKE '%business%'
    """)
    rows = cur.fetchall()
    cur.close()
    conn.close()
    record("Notion page with Beijing/Shanghai/Trip exists", len(rows) >= 1,
           f"Found {len(rows)} matching pages")


def check_gcal():
    print("\n=== Check 3: Google Calendar events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-03-10'
          AND start_datetime < '2026-03-11'
          AND summary NOT ILIKE '%client meeting%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()
    record("At least 1 new travel calendar event on 2026-03-10", len(events) >= 1,
           f"Found {len(events)} non-meeting events: {[e[0] for e in events]}")


def check_email():
    print("\n=== Check 4: Email to travel@consulting.com ===")
    cnt = 0
    sent = 0
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute(
            "SELECT COUNT(*) FROM email.messages"
            " WHERE to_addr::text ILIKE '%travel@consulting.com%'"
            "   AND from_addr NOT ILIKE '%travel@consulting.com%'"
        )
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
    except Exception:
        pass
    try:
        conn2 = psycopg2.connect(**DB_CONFIG)
        cur2 = conn2.cursor()
        cur2.execute(
            "SELECT COUNT(*) FROM email.sent_log"
            " WHERE to_addr::text ILIKE '%travel@consulting.com%'"
        )
        sent = cur2.fetchone()[0]
        cur2.close()
        conn2.close()
    except Exception:
        pass
    record("Email sent to travel@consulting.com", cnt >= 1 or sent >= 1,
           f"messages: {cnt}, sent_log: {sent}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion()
    check_gcal()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    accuracy = PASS_COUNT / total * 100 if total > 0 else 0
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
