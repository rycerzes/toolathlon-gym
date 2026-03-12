"""
Generate groundtruth files:
  - Q4_2025_Sales_Report.xlsx (two sheets)
  - Executive_Summary.docx
"""
import os
import openpyxl
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from docx import Document
from docx.shared import Pt

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Data ──────────────────────────────────────────────────────────────────

TARGETS = {
    "Asia Pacific": 65000,
    "Europe": 60000,
    "Latin America": 55000,
    "Middle East": 50000,
    "North America": 55000,
}

ACTUALS = {
    "Asia Pacific":  {"revenue": 70510.11, "orders": 386, "customers": 246},
    "Europe":        {"revenue": 54490.62, "orders": 364, "customers": 246},
    "Latin America": {"revenue": 57100.57, "orders": 335, "customers": 231},
    "Middle East":   {"revenue": 57505.34, "orders": 341, "customers": 233},
    "North America": {"revenue": 51818.56, "orders": 358, "customers": 242},
}

SEGMENTS = {
    "Asia Pacific": [
        ("Enterprise", 25754.31, 108),
        ("Consumer", 16491.17, 101),
        ("SMB", 14248.70, 89),
        ("Government", 14015.93, 88),
    ],
    "Europe": [
        ("Enterprise", 16347.91, 103),
        ("Consumer", 16321.67, 100),
        ("SMB", 14196.64, 92),
        ("Government", 7624.40, 69),
    ],
    "Latin America": [
        ("Consumer", 20924.91, 104),
        ("Enterprise", 16106.22, 92),
        ("Government", 11244.69, 77),
        ("SMB", 8824.75, 62),
    ],
    "Middle East": [
        ("Government", 18668.73, 103),
        ("SMB", 16965.25, 94),
        ("Consumer", 11197.00, 76),
        ("Enterprise", 10674.36, 68),
    ],
    "North America": [
        ("Government", 16309.09, 105),
        ("Enterprise", 14532.33, 95),
        ("Consumer", 13054.90, 88),
        ("SMB", 7922.24, 70),
    ],
}


def create_excel():
    wb = openpyxl.Workbook()

    # ── Sheet 1: Regional Performance ──
    ws1 = wb.active
    ws1.title = "Regional Performance"
    headers = ["Region", "Target", "Actual", "Variance", "Variance_Pct", "Order_Count", "Customer_Count"]
    ws1.append(headers)

    # Style header
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="2C3E50", end_color="2C3E50", fill_type="solid")
    for col_idx, h in enumerate(headers, 1):
        cell = ws1.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for region in sorted(ACTUALS.keys()):
        target = TARGETS[region]
        actual = round(ACTUALS[region]["revenue"], 2)
        variance = round(actual - target, 2)
        variance_pct = round(variance / target * 100, 1)
        order_count = ACTUALS[region]["orders"]
        customer_count = ACTUALS[region]["customers"]
        ws1.append([region, target, actual, variance, variance_pct, order_count, customer_count])

    # Auto-width
    for col in ws1.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws1.column_dimensions[col[0].column_letter].width = max_len + 4

    # ── Sheet 2: Segment Breakdown ──
    ws2 = wb.create_sheet("Segment Breakdown")
    seg_headers = ["Region", "Segment", "Revenue", "Orders"]
    ws2.append(seg_headers)

    for col_idx, h in enumerate(seg_headers, 1):
        cell = ws2.cell(row=1, column=col_idx)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    for region in sorted(SEGMENTS.keys()):
        # Already sorted by revenue descending in the data
        for segment, revenue, orders in SEGMENTS[region]:
            ws2.append([region, segment, round(revenue, 2), orders])

    for col in ws2.columns:
        max_len = max(len(str(cell.value or "")) for cell in col)
        ws2.column_dimensions[col[0].column_letter].width = max_len + 4

    xlsx_path = os.path.join(OUTPUT_DIR, "Q4_2025_Sales_Report.xlsx")
    wb.save(xlsx_path)
    print(f"Created: {xlsx_path}")


def create_docx():
    doc = Document()

    # Title
    title = doc.add_heading("Q4 2025 Sales Reconciliation - Executive Summary", level=1)

    total_actual = sum(v["revenue"] for v in ACTUALS.values())
    total_target = sum(TARGETS.values())

    beat_regions = [r for r in sorted(ACTUALS.keys()) if ACTUALS[r]["revenue"] > TARGETS[r]]
    missed_regions = [r for r in sorted(ACTUALS.keys()) if ACTUALS[r]["revenue"] <= TARGETS[r]]

    beat_str = ", ".join(beat_regions)
    missed_str = ", ".join(missed_regions)

    summary_text = (
        f"In Q4 2025, the company generated a total revenue of ${total_actual:,.2f} against a combined "
        f"target of ${total_target:,.2f} across all five regions. Three regions exceeded their targets: "
        f"{beat_str}, demonstrating strong performance driven by robust enterprise and consumer segments. "
        f"However, two regions fell short of expectations: {missed_str}. Europe underperformed its "
        f"$60,000 target by approximately $5,509, while North America missed its $55,000 target by "
        f"roughly $3,181. Overall, the company achieved ${total_actual - total_target:,.2f} above the "
        f"aggregate target, indicating solid Q4 performance despite regional variances. "
        f"Management should investigate the underperformance in Europe and North America to identify "
        f"corrective actions for Q1 2026."
    )

    para = doc.add_paragraph(summary_text)
    for run in para.runs:
        run.font.size = Pt(11)

    docx_path = os.path.join(OUTPUT_DIR, "Executive_Summary.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")


if __name__ == "__main__":
    create_excel()
    create_docx()
