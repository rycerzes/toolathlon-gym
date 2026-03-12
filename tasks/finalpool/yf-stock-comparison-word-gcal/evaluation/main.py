"""Evaluation for yf-stock-comparison-word-gcal."""
import os
import argparse, os, sys
import psycopg2


def check_word_doc(agent_workspace):
    errors = []
    doc_path = os.path.join(agent_workspace, "Stock_Comparison_Report.docx")
    if not os.path.exists(doc_path):
        return ["Stock_Comparison_Report.docx not found in agent workspace"]
    try:
        from docx import Document
        doc = Document(doc_path)
        full_text = "\n".join(p.text for p in doc.paragraphs)
        # Also check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text += "\n" + cell.text
        full_text_lower = full_text.lower()
        for sym in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
            if sym not in full_text:
                errors.append(f"Stock symbol {sym} not found in Word doc")
        # Check for price data
        for price in ["300", "218", "293", "239", "150"]:
            if price not in full_text:
                errors.append(f"Price starting with {price} not found in Word doc")
                break  # only report one price issue
        if "stock comparison" not in full_text_lower:
            errors.append("Title 'Stock Comparison' not found in Word doc")
    except Exception as e:
        errors.append(f"Error reading Word doc: {e}")
    return errors


def check_gcal_event():
    errors = []
    try:
        conn = psycopg2.connect(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym",
                                user="eigent", password="camel")
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, start_datetime, end_datetime
            FROM gcal.events
            WHERE start_datetime::date = '2026-04-10'
            ORDER BY start_datetime
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        if not rows:
            errors.append("No GCal event found on 2026-04-10")
        else:
            summaries = [r[0].lower() if r[0] else "" for r in rows]
            found = any("portfolio" in s or "review" in s for s in summaries)
            if not found:
                errors.append(f"No portfolio review event on 2026-04-10 (found: {[r[0] for r in rows]})")
    except Exception as e:
        errors.append(f"Error checking GCal: {e}")
    return errors


def check_email():
    errors = []
    try:
        conn = psycopg2.connect(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym",
                                user="eigent", password="camel")
        cur = conn.cursor()
        cur.execute("""
            SELECT subject, to_addr, body_text FROM email.messages
            WHERE to_addr::text ILIKE '%analyst@investment.com%'
            ORDER BY id DESC LIMIT 5
        """)
        rows = cur.fetchall()
        cur.close(); conn.close()
        if not rows:
            errors.append("No email found to analyst@investment.com")
    except Exception as e:
        errors.append(f"Error checking email: {e}")
    return errors


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    agent_ws = args.agent_workspace or os.path.join(os.path.dirname(__file__), "..", "groundtruth_workspace")

    all_errors = []

    print("  Checking Word document...")
    errs = check_word_doc(agent_ws)
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]: print(f"    ERROR: {e}")
    else:
        print("    PASS")

    print("  Checking GCal event...")
    errs = check_gcal_event()
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]: print(f"    ERROR: {e}")
    else:
        print("    PASS")

    print("  Checking email...")
    errs = check_email()
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]: print(f"    ERROR: {e}")
    else:
        print("    PASS")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]: print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
