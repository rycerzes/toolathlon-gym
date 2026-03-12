"""Evaluation for canvas-assignment-feedback-gform-word."""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


EXPECTED_ASSIGNMENTS = [
    ("CMA 34904", 1454, 85.10, 32, 2.20),
    ("CMA 34905", 1363, 88.25, 45, 3.30),
    ("CMA 34906", 1254, 77.22, 70, 5.58),
    ("CMA 34907", 1234, 78.81, 72, 5.83),
    ("CMA 34908", 1213, 78.76, 85, 7.01),
    ("CMA 34909", 1213, 77.10, 134, 11.05),
    ("CMA 34910", 1184, 77.33, 118, 9.97),
    ("TMA 34899", 1826, 78.60, 173, 9.47),
    ("TMA 34900", 1601, 78.11, 212, 13.24),
    ("TMA 34901", 1398, 74.28, 322, 23.03),
    ("TMA 34902", 1307, 73.48, 257, 19.66),
    ("TMA 34903", 1137, 76.36, 149, 13.10),
]


def check_word(agent_workspace, groundtruth_workspace):
    """Check Word document output."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx available", False, "python-docx not installed")
        return

    agent_file = os.path.join(agent_workspace, "Assignment_Analysis.docx")
    check("Word file exists", os.path.isfile(agent_file), f"Expected {agent_file}")
    if not os.path.isfile(agent_file):
        return

    try:
        doc = Document(agent_file)
    except Exception as e:
        check("Word file readable", False, str(e))
        return

    # Check title
    full_text = "\n".join(p.text for p in doc.paragraphs)
    check("Document title present",
          "foundations of finance" in full_text.lower() and
          "assignment analysis" in full_text.lower(),
          f"Title not found in: {full_text[:200]}")

    # Find table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")
    if not doc.tables:
        return

    tbl = doc.tables[0]
    # Check headers
    if tbl.rows:
        headers = [c.text.strip() for c in tbl.rows[0].cells]
        check("Column 'Assignment_Name' present",
              any("assignment_name" in h.lower() for h in headers),
              f"Headers: {headers}")
        check("Column 'Total_Submissions' present",
              any("total_submissions" in h.lower() or "total" in h.lower() for h in headers),
              f"Headers: {headers}")
        check("Column 'Late_Submissions' present",
              any("late" in h.lower() for h in headers),
              f"Headers: {headers}")

    # Check rows - should have 12 data rows
    data_rows = list(tbl.rows)[1:]
    check("Table has 12 data rows", len(data_rows) >= 12,
          f"Found {len(data_rows)} rows")

    # Check specific assignment data
    row_names = [row.cells[0].text.strip() for row in data_rows if row.cells]
    check("CMA 34905 row present",
          any("cma 34905" in n.lower() for n in row_names),
          f"Row names: {row_names[:5]}")
    check("TMA 34901 row present",
          any("tma 34901" in n.lower() for n in row_names),
          f"Row names: {row_names[:5]}")


def check_gform():
    """Check Google Form creation."""
    print("\n=== Checking Google Form ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gform.forms")
    forms = cur.fetchall()
    check("At least one form created", len(forms) >= 1, f"Found {len(forms)} forms")

    form_id = None
    for fid, title in forms:
        if "assignment" in (title or "").lower() or "feedback" in (title or "").lower() or "finance" in (title or "").lower():
            form_id = fid
            break

    if not form_id and forms:
        form_id = forms[0][0]

    check("Form with relevant title found", form_id is not None,
          f"Forms: {forms}")

    if form_id:
        cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (form_id,))
        q_count = cur.fetchone()[0]
        check("Form has at least 5 questions", q_count >= 5,
              f"Found {q_count} questions")

    conn.close()


def check_emails():
    """Check that summary email was sent."""
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("SELECT subject, from_addr, to_addr, body_text FROM email.messages")
    all_emails = cur.fetchall()
    conn.close()

    def find_email_for_recipient(recipient):
        for subj, from_addr, to_addr, body in all_emails:
            if to_addr:
                recipients = []
                if isinstance(to_addr, list):
                    recipients = [str(r).strip().lower() for r in to_addr]
                elif isinstance(to_addr, str):
                    try:
                        parsed = json.loads(to_addr)
                        if isinstance(parsed, list):
                            recipients = [str(r).strip().lower() for r in parsed]
                        else:
                            recipients = [str(to_addr).strip().lower()]
                    except (json.JSONDecodeError, TypeError):
                        recipients = [str(to_addr).strip().lower()]
                if recipient.lower() in recipients:
                    return subj, from_addr, to_addr, body
        return None

    result = find_email_for_recipient("instructor@financeou.example.com")
    check("Summary email sent to instructor@financeou.example.com", result is not None,
          f"Total emails found: {len(all_emails)}")

    if result:
        subj, from_addr, to_addr, body = result
        check("Email subject contains 'Assignment Analysis'",
              "assignment" in (subj or "").lower() and "analysis" in (subj or "").lower(),
              f"Subject: {subj}")
        check("Email from analytics@university.example.com",
              "analytics@university.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        body_lower = (body or "").lower()
        check("Email mentions 12 assignments",
              "12" in body_lower or "twelve" in body_lower,
              "Expected mention of 12")
        check("Email mentions CMA 34905 or highest avg",
              "cma 34905" in body_lower or ("34905" in body_lower),
              "Expected CMA 34905 mentioned")
        check("Email mentions TMA 34901 or most late",
              "tma 34901" in body_lower or "34901" in body_lower,
              "Expected TMA 34901 mentioned")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("CANVAS ASSIGNMENT FEEDBACK GFORM WORD - EVALUATION")
    print("=" * 70)

    check_word(args.agent_workspace, gt_dir)
    check_gform()
    check_emails()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
