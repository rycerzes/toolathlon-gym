"""Evaluation for sf-salary-market-benchmark."""
import argparse
import os
import sys

import openpyxl


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Salary_Benchmark.xlsx")
    gt_file = os.path.join(gt_dir, "Salary_Benchmark.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_errors = []

    # Check Department Analysis sheet
    print("  Checking Department Analysis...")
    a_rows = load_sheet_rows(agent_wb, "Department Analysis")
    g_rows = load_sheet_rows(gt_wb, "Department Analysis")
    if a_rows is None:
        all_errors.append("Sheet 'Department Analysis' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Department Analysis' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing row: {g_row[0]}")
                continue
            if len(a_row) > 1 and len(g_row) > 1:
                if not num_close(a_row[1], g_row[1], 500):
                    errors.append(f"{key}.Avg_Salary: {a_row[1]} vs {g_row[1]}")
            if len(a_row) > 2 and len(g_row) > 2:
                if not num_close(a_row[2], g_row[2], 500):
                    errors.append(f"{key}.Market_Benchmark: {a_row[2]} vs {g_row[2]}")
            if len(a_row) > 3 and len(g_row) > 3:
                if not num_close(a_row[3], g_row[3], 0.05):
                    errors.append(f"{key}.Pay_Ratio: {a_row[3]} vs {g_row[3]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Role Details sheet
    print("  Checking Role Details...")
    a_rows = load_sheet_rows(agent_wb, "Role Details")
    g_rows = load_sheet_rows(gt_wb, "Role Details")
    if a_rows is None:
        all_errors.append("Sheet 'Role Details' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Role Details' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {}
        for r in a_data:
            if r and r[0] is not None and r[1] is not None:
                k = f"{str(r[0]).strip().lower()}|{str(r[1]).strip().lower()}"
                a_lookup[k] = r
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = f"{str(g_row[0]).strip().lower()}|{str(g_row[1]).strip().lower()}"
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing row: {g_row[0]}|{g_row[1]}")
                continue
            if len(a_row) > 3 and len(g_row) > 3:
                if not num_close(a_row[3], g_row[3], 500):
                    errors.append(f"{key}.Avg_Salary: {a_row[3]} vs {g_row[3]}")
            if len(a_row) > 4 and len(g_row) > 4:
                if not num_close(a_row[4], g_row[4], 500):
                    errors.append(f"{key}.Market_Benchmark: {a_row[4]} vs {g_row[4]}")
            if len(a_row) > 5 and len(g_row) > 5:
                if not num_close(a_row[5], g_row[5], 0.05):
                    errors.append(f"{key}.Pay_Ratio: {a_row[5]} vs {g_row[5]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Summary sheet
    print("  Checking Summary...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    g_rows = load_sheet_rows(gt_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found in agent output")
    elif g_rows is None:
        all_errors.append("Sheet 'Summary' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing row: {g_row[0]}")
                continue
            if len(a_row) > 1 and len(g_row) > 1:
                if key in ("total_employees", "total_roles_analyzed", "roles_above_market", "roles_below_market"):
                    if not num_close(a_row[1], g_row[1], 2):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif key in ("overall_pay_ratio", "most_above_market_ratio", "most_below_market_ratio"):
                    if not num_close(a_row[1], g_row[1], 0.05):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif key in ("overall_avg_salary", "overall_market_avg"):
                    if not num_close(a_row[1], g_row[1], 500):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "role" in key:
                    if not str_match(a_row[1], g_row[1]):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Word document exists
    print("  Checking Compensation_Report.docx...")
    word_file = os.path.join(args.agent_workspace, "Compensation_Report.docx")
    if not os.path.exists(word_file):
        all_errors.append("Compensation_Report.docx not found")
        print("    FAIL: file not found")
    else:
        try:
            from docx import Document
            doc = Document(word_file)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            if "executive summary" not in text and "summary" not in text:
                all_errors.append("Word doc missing executive summary section")
            if "recommend" not in text:
                all_errors.append("Word doc missing recommendations section")
            if len(doc.paragraphs) < 5:
                all_errors.append("Word doc too short")
            print("    PASS" if not any("Word" in e for e in all_errors) else "    ERRORS found")
        except Exception as e:
            all_errors.append(f"Word doc read error: {e}")
            print(f"    ERROR: {e}")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
