"""
Generate groundtruth Department_Review_2025.docx for evaluation comparison.
"""

import os
from docx import Document
from docx.shared import Pt


DEPARTMENTS = {
    "Engineering": {"headcount": 7096, "avg_salary": 58991.61, "avg_perf": 3.21},
    "Finance": {"headcount": 7148, "avg_salary": 57878.19, "avg_perf": 3.21},
    "HR": {"headcount": 7077, "avg_salary": 58920.45, "avg_perf": 3.20},
    "Operations": {"headcount": 7120, "avg_salary": 57808.74, "avg_perf": 3.18},
    "R&D": {"headcount": 7083, "avg_salary": 57905.93, "avg_perf": 3.20},
    "Sales": {"headcount": 7232, "avg_salary": 58864.79, "avg_perf": 3.19},
    "Support": {"headcount": 7244, "avg_salary": 58400.48, "avg_perf": 3.20},
}


def create_groundtruth_docx():
    doc = Document()

    # Title
    doc.add_heading("2025 Department Performance Review", level=0)

    # Department sections
    for dept_name, data in sorted(DEPARTMENTS.items()):
        doc.add_heading(dept_name, level=1)
        doc.add_paragraph(
            f"The {dept_name} department has a total headcount of {data['headcount']} "
            f"employees. The average salary across the department is ${data['avg_salary']:,.2f}. "
            f"The average performance rating is {data['avg_perf']:.2f} on a scale of 1 to 5."
        )

    # Conclusion
    doc.add_heading("Conclusion", level=1)
    total_employees = sum(d["headcount"] for d in DEPARTMENTS.values())
    avg_perf_all = sum(d["avg_perf"] * d["headcount"] for d in DEPARTMENTS.values()) / total_employees
    doc.add_paragraph(
        f"Across all seven departments, the company employs a total of {total_employees:,} "
        f"people. The overall weighted average performance rating is {avg_perf_all:.2f}, "
        f"which indicates that employees generally meet expectations. Continued focus on "
        f"professional development and goal alignment is recommended to drive performance "
        f"ratings closer to the target of 3.5."
    )

    output_path = os.path.join(os.path.dirname(__file__), "Department_Review_2025.docx")
    doc.save(output_path)
    print(f"Created {output_path}")


if __name__ == "__main__":
    create_groundtruth_docx()
    print("Groundtruth files generated successfully.")
