"""
Evaluation script for canvas-assignment-deadline-word-gcal task.

Checks:
1. Excel file Assignment_Deadlines_FFF2013J.xlsx - 2 sheets with correct data
2. Word document Assignment_Schedule_FFF2013J.docx exists with table
3. Google Calendar has assignment reminder events
4. Email sent to fff2013j.students@university.edu
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_by_name(wb, name):
    for sname in wb.sheetnames:
        if sname.strip().lower() == name.strip().lower():
            return [[cell.value for cell in row] for row in wb[sname].iter_rows()]
    return None


# ============================================================================
# Check 1: Excel file
# ============================================================================

def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Assignment_Deadlines_FFF2013J.xlsx ===")

    try:
        import openpyxl
    except ImportError:
        record("openpyxl available", False, "pip install openpyxl")
        return False

    agent_file = os.path.join(agent_workspace, "Assignment_Deadlines_FFF2013J.xlsx")
    gt_file = os.path.join(groundtruth_workspace, "Assignment_Deadlines_FFF2013J.xlsx")

    if not os.path.isfile(agent_file):
        record("Excel file exists", False, f"Not found: {agent_file}")
        return False
    record("Excel file exists", True)

    if not os.path.isfile(gt_file):
        record("Groundtruth file exists", False, f"Not found: {gt_file}")
        return False

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_ok = True

    # Check All Assignments sheet
    a_all = load_sheet_by_name(agent_wb, "All Assignments")
    g_all = load_sheet_by_name(gt_wb, "All Assignments")
    record("Sheet 'All Assignments' exists", a_all is not None)

    if a_all is not None and g_all is not None:
        a_data = [r for r in a_all[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_all[1:] if any(v is not None for v in r)]
        record("All Assignments row count matches",
               len(a_data) == len(g_data),
               f"Expected {len(g_data)}, got {len(a_data)}")

        # Build lookup by assignment name
        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            name = str(g_row[0]).strip()
            key = name.lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Assignment row: {name}", False, "Not found")
                all_ok = False
                continue
            record(f"Assignment row: {name}", True)

            # Points_Possible (col 1)
            if len(g_row) > 1 and len(a_row) > 1:
                record(f"{name}: Points_Possible correct",
                       num_close(a_row[1], g_row[1], 0.01),
                       f"got {a_row[1]}, expected {g_row[1]}")

    # Check Summary sheet
    a_summ = load_sheet_by_name(agent_wb, "Summary")
    g_summ = load_sheet_by_name(gt_wb, "Summary")
    record("Sheet 'Summary' exists", a_summ is not None)

    if a_summ is not None and g_summ is not None:
        a_data = [r for r in a_summ[1:] if any(v is not None for v in r)]
        g_data = [r for r in g_summ[1:] if any(v is not None for v in r)]

        a_lookup = {}
        for row in a_data:
            if row and row[0] is not None:
                a_lookup[str(row[0]).strip().lower()] = row

        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                record(f"Summary row: {g_row[0]}", False, "Not found")
                all_ok = False
                continue
            record(f"Summary row: {g_row[0]}", True)

            if key == "total_assignments":
                record("Total_Assignments = 13",
                       num_close(a_row[1], g_row[1], 0),
                       f"got {a_row[1]}, expected {g_row[1]}")
            elif key == "total_points_possible":
                record("Total_Points_Possible = 900",
                       num_close(a_row[1], g_row[1], 1.0),
                       f"got {a_row[1]}, expected {g_row[1]}")
            elif key == "avg_points_per_assignment":
                record("Avg_Points_Per_Assignment correct",
                       num_close(a_row[1], g_row[1], 1.0),
                       f"got {a_row[1]}, expected {g_row[1]}")

    return all_ok


# ============================================================================
# Check 2: Word document
# ============================================================================

def check_word(agent_workspace):
    print("\n=== Checking Assignment_Schedule_FFF2013J.docx ===")

    docx_path = os.path.join(agent_workspace, "Assignment_Schedule_FFF2013J.docx")
    if not os.path.isfile(docx_path):
        record("Word file exists", False, f"Not found: {docx_path}")
        return False
    record("Word file exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        headings_text = " ".join(p.text for p in doc.paragraphs
                                 if p.style.name.startswith("Heading")).lower()

        record("Word doc has content", len(all_text.strip()) >= 100,
               f"Content length: {len(all_text.strip())}")
        record("Word doc heading mentions Finance or FFF",
               any(term in (all_text + headings_text) for term in
                   ["finance", "fff", "foundations", "assignment"]),
               "Missing Finance/FFF content in doc")

        tables = doc.tables
        record("Word doc has at least 1 table", len(tables) >= 1,
               f"Found {len(tables)} tables")

        # Check table has assignment data
        if tables:
            table_text = " ".join(
                cell.text.lower()
                for row in tables[0].rows
                for cell in row.cells
            )
            record("Table has TMA assignments",
                   "tma" in table_text,
                   f"Table text: {table_text[:200]}")

        return True

    except ImportError:
        size = os.path.getsize(docx_path)
        record("Word file has content (>3KB)", size > 3000, f"Size: {size} bytes")
        return size > 3000
    except Exception as e:
        record("Word file readable", False, str(e))
        return False


# ============================================================================
# Check 3: Google Calendar
# ============================================================================

def check_gcal():
    print("\n=== Checking Google Calendar ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT summary, start_datetime, end_datetime
        FROM gcal.events
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_gcal] Found {len(events)} calendar events.")

    # Expect events with "Assignment Due:" in title
    assignment_events = [e for e in events
                         if e[0] and ("assignment" in e[0].lower() or
                                      "due" in e[0].lower() or
                                      "tma" in e[0].lower() or
                                      "cma" in e[0].lower())]
    record("Assignment reminder events created",
           len(assignment_events) >= 5,
           f"Found {len(assignment_events)} assignment events (expected >=5 for 13 assignments with due dates)")

    record("At least 1 calendar event created",
           len(events) >= 1,
           f"Found {len(events)}")

    return len(assignment_events) >= 5


# ============================================================================
# Check 4: Email
# ============================================================================

def check_emails():
    print("\n=== Checking Emails ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_emails] Found {len(all_emails)} total emails.")
    record("At least 1 email sent", len(all_emails) >= 1, f"Found {len(all_emails)}")

    found_email = False
    for subject, from_addr, to_addr, body_text in all_emails:
        to_str = str(to_addr or "").lower()
        subject_lower = (subject or "").lower()
        if ("fff2013j.students@university.edu" in to_str or
                "deadline" in subject_lower or "assignment" in subject_lower):
            found_email = True
            record("Email to fff2013j.students@university.edu found", True)

            record("Email subject mentions assignment or deadline",
                   "assignment" in subject_lower or "deadline" in subject_lower,
                   f"Subject: {subject}")

            body_lower = (body_text or "").lower()
            record("Email body lists assignments",
                   any(term in body_lower for term in ["tma", "cma", "assignment", "due date"]),
                   "Body missing assignment list")
            break

    if not found_email:
        record("Assignment deadline email found", False,
               f"Emails: {[(e[0], str(e[2])[:60]) for e in all_emails[:3]]}")

    return found_email


# ============================================================================
# Main
# ============================================================================

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    excel_ok = check_excel(args.agent_workspace, gt_dir)
    word_ok = check_word(args.agent_workspace)
    gcal_ok = check_gcal()
    email_ok = check_emails()

    all_passed = excel_ok and word_ok and gcal_ok and email_ok

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    print(f"  Overall: {'PASS' if all_passed else 'FAIL'}")

    if args.res_log_file:
        result = {
            "passed": PASS_COUNT,
            "failed": FAIL_COUNT,
            "success": all_passed,
        }
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if all_passed else 1)


if __name__ == "__main__":
    main()
