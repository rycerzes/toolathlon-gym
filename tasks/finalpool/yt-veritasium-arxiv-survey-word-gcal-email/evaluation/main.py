"""
Evaluation for yt-veritasium-arxiv-survey-word-gcal-email task.

Checks:
1. Science_Communication_Survey.docx exists
2. Word doc has >= 6 headings
3. Word doc text contains Veritasium and at least 3 science terms
4. Word doc text contains at least 3 paper titles or author names
5. GCal has 2 new events in April 2026
6. Email sent to seminar@science.edu
7. Email sent to collab@research.org
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Science_Communication_Survey.docx ===")
    docx_path = os.path.join(agent_workspace, "Science_Communication_Survey.docx")
    if not os.path.exists(docx_path):
        record("Science_Communication_Survey.docx exists", False, f"Not found at {docx_path}")
        return
    record("Science_Communication_Survey.docx exists", True)

    try:
        import docx
        doc = docx.Document(docx_path)
    except ImportError:
        # Fallback: check file size
        size = os.path.getsize(docx_path)
        record("Word doc is non-empty (>5KB)", size > 5000, f"Size: {size} bytes")
        return
    except Exception as e:
        record("Word doc readable", False, str(e))
        return

    # Count headings
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
    record("Word doc has >= 6 headings", len(headings) >= 6,
           f"Found {len(headings)} headings: {[h.text[:50] for h in headings[:8]]}")

    # Check text content
    all_text = " ".join(p.text for p in doc.paragraphs).lower()

    has_veritasium = "veritasium" in all_text
    record("Word doc mentions Veritasium", has_veritasium, "Veritasium not found in text")

    science_terms = ["quantum", "evolution", "cognitive", "mathematical", "fluid", "fermi", "brain",
                     "paradox", "decoherence", "neuroplasticity", "biomimetic", "game theory"]
    found_terms = [t for t in science_terms if t in all_text]
    record("Word doc contains >= 3 science topic terms", len(found_terms) >= 3,
           f"Found: {found_terms}")

    # Check for paper author/title references
    key_refs = ["sean carroll", "martin nowak", "daniel kahneman", "timothy gowers",
                "john dabiri", "anders sandberg", "michael merzenich",
                "many-worlds", "game theory and evolution", "cognitive biases",
                "paradoxes in mathematics", "fluid dynamics in nature", "fermi paradox",
                "neuroplasticity"]
    found_refs = [r for r in key_refs if r in all_text]
    record("Word doc references >= 3 paper authors/titles", len(found_refs) >= 3,
           f"Found refs: {found_refs}")

    # Check for executive summary
    has_exec_summary = "executive summary" in all_text or "summary" in all_text
    record("Word doc contains Executive Summary section", has_exec_summary,
           "No Executive Summary found")

    # Check for references section
    has_references = "references" in all_text or "arxiv:" in all_text or "arxiv.org" in all_text
    record("Word doc contains References section", has_references,
           "No References section found")

    # Check word count is substantial
    word_count = len(all_text.split())
    record("Word doc has substantial content (>= 500 words)", word_count >= 500,
           f"Word count: {word_count}")


def check_gcal():
    print("\n=== Check 2: GCal April 2026 Events ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()
    cur.close()
    conn.close()

    # Expect at least 2 total events in April (Lab Meeting preloaded + at least 1 new)
    record("GCal has events in April 2026 (at least preloaded Lab Meeting)",
           len(events) >= 1,
           f"Found {len(events)} events: {[e[0] for e in events]}")

    # Check if agent added any seminar/survey/review events
    new_events = [e for e in events if "lab meeting" not in (e[0] or "").lower()]
    summaries = " ".join(e[0] or "" for e in events).lower()
    has_seminar = "seminar" in summaries or "survey" in summaries or "science" in summaries or "review" in summaries
    record("GCal has >= 1 seminar or review event added by agent", len(new_events) >= 1 and has_seminar,
           f"New events: {[e[0] for e in new_events]}")


def check_emails_sent():
    print("\n=== Check 3: Emails Sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    try:
        # Check messages in Sent/SENT folders
        cur.execute("""
            SELECT m.to_addr FROM email.messages m
            JOIN email.folders f ON m.folder_id = f.id
            WHERE UPPER(f.name) = 'SENT'
        """)
        sent_rows = cur.fetchall()
        # Also check via sent_log join
        cur.execute("""
            SELECT m.to_addr FROM email.sent_log sl
            JOIN email.messages m ON sl.message_id = m.id
        """)
        sent_rows += cur.fetchall()
        sent_text = " ".join(str(row[0]) for row in sent_rows).lower()

        record("Email sent to seminar@science.edu",
               "seminar@science.edu" in sent_text,
               f"Sent entries: {len(sent_rows)}")
        record("Email sent to collab@research.org",
               "collab@research.org" in sent_text,
               f"Sent entries: {len(sent_rows)}")
    except Exception as e:
        record("Email sent check", False, str(e))
    finally:
        cur.close()
        conn.close()


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_gcal()
    check_emails_sent()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
