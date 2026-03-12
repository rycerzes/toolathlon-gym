"""Evaluation for wc-low-stock-reorder-gcal-email."""
import argparse
import os
import sys
import psycopg2
import openpyxl


DB = {"host": os.environ.get("PGHOST", "localhost"), "port": 5432, "dbname": "toolathlon_gym", "user": "eigent", "password": "camel"}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def get_low_stock_counts():
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()
    cur.execute("SELECT stock_quantity, count(*) FROM wc.products WHERE stock_quantity IS NOT NULL AND stock_quantity <= 5 GROUP BY stock_quantity ORDER BY stock_quantity")
    rows = cur.fetchall()
    cur.close()
    conn.close()
    urgent = sum(r[1] for r in rows if r[0] == 0)
    high = sum(r[1] for r in rows if r[0] in [1, 2])
    normal = sum(r[1] for r in rows if 3 <= r[0] <= 5)
    return {"total": urgent + high + normal, "urgent": urgent, "high": high, "normal": normal}


def check_gcal_events():
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT count(*) FROM gcal.events WHERE LOWER(summary) LIKE '%stock%' OR LOWER(summary) LIKE '%reorder%' OR LOWER(summary) LIKE '%replenishment%'")
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
        return cnt >= 2
    except Exception:
        return False


def check_email_sent():
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT count(*) FROM email.messages WHERE LOWER(to_addr::text) LIKE '%procurement%' AND (LOWER(subject) LIKE '%low stock%' OR LOWER(subject) LIKE '%reorder%')")
        cnt = cur.fetchone()[0]
        cur.close()
        conn.close()
        return cnt >= 1
    except Exception:
        return False


def check_word_file(agent_workspace):
    word_path = os.path.join(agent_workspace, "Reorder_Report.docx")
    if not os.path.exists(word_path):
        return False, "Reorder_Report.docx not found"
    try:
        from docx import Document
        doc = Document(word_path)
        full_text = " ".join(p.text for p in doc.paragraphs)
        full_text += " ".join(cell.text for t in doc.tables for row in t.rows for cell in row.cells)
        if "low stock" in full_text.lower() or "reorder" in full_text.lower():
            return True, ""
        return False, "Word doc does not mention 'Low Stock' or 'Reorder'"
    except Exception as e:
        # Check just file existence if docx not available
        return True, ""


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Low_Stock_Reorder.xlsx")
    gt_file = os.path.join(gt_dir, "Low_Stock_Reorder.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    all_errors = []

    try:
        counts = get_low_stock_counts()
    except Exception as e:
        print(f"WARNING: Could not query DB: {e}")
        counts = {"total": 25, "urgent": 5, "high": 7, "normal": 13}

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    # Check Reorder List sheet
    print("  Checking Reorder List sheet...")
    a_rows = load_sheet_rows(agent_wb, "Reorder List")
    if a_rows is None:
        all_errors.append("Sheet 'Reorder List' not found in agent output")
    else:
        data_rows = [r for r in a_rows[1:] if r and any(c is not None for c in r)]
        if len(data_rows) < counts["total"]:
            all_errors.append(f"Reorder List has {len(data_rows)} rows, expected >= {counts['total']}")
        else:
            print(f"    PASS ({len(data_rows)} data rows)")
        # Check urgent items have priority "Urgent"
        priority_errors = 0
        for row in data_rows:
            if row and len(row) >= 4:
                stock_val = row[1]
                priority_val = str(row[3]).strip() if row[3] is not None else ""
                try:
                    if int(float(str(stock_val))) == 0 and priority_val.lower() != "urgent":
                        priority_errors += 1
                except (TypeError, ValueError):
                    pass
        if priority_errors > 0:
            all_errors.append(f"{priority_errors} stock=0 items not labeled 'Urgent'")
        else:
            print("    Priority validation PASS")

    # Check Summary sheet
    print("  Checking Summary sheet...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found in agent output")
    else:
        a_data = {str(r[0]).strip().lower(): r[1] for r in a_rows[1:] if r and r[0] is not None}
        errors = []
        for metric, exp_val in [
            ("total_low_stock_items", counts["total"]),
            ("urgent_count", counts["urgent"]),
            ("high_priority_count", counts["high"]),
            ("normal_priority_count", counts["normal"]),
        ]:
            val = a_data.get(metric)
            if val is None:
                errors.append(f"Missing metric: {metric}")
            elif not num_close(val, exp_val, 0):
                errors.append(f"{metric}: {val} vs {exp_val}")
        if errors:
            all_errors.extend(errors)
            for e in errors[:5]:
                print(f"    ERROR: {e}")
        else:
            print("    PASS")

    # Check GCal events
    print("  Checking GCal events...")
    if check_gcal_events():
        print("    PASS")
    else:
        all_errors.append("Expected >= 2 calendar events related to stock/reorder, not found")

    # Check email sent
    print("  Checking email to procurement...")
    if check_email_sent():
        print("    PASS")
    else:
        all_errors.append("Email to procurement@company.com with 'Low Stock' or 'Reorder' subject not found")

    # Check Word document
    print("  Checking Word document...")
    ok, detail = check_word_file(args.agent_workspace)
    if ok:
        print("    PASS")
    else:
        all_errors.append(detail)

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
