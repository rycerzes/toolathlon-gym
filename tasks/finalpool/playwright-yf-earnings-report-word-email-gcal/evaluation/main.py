"""Evaluation for playwright-yf-earnings-report-word-email-gcal."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")
    agent_ws = args.agent_workspace or task_root

    all_errors = []

    # --- Check 1: Excel file ---
    import openpyxl

    agent_file = os.path.join(agent_ws, "Earnings_Data_Appendix.xlsx")
    gt_file = os.path.join(gt_dir, "Earnings_Data_Appendix.xlsx")

    print("Checking Excel file...")
    if not os.path.exists(agent_file):
        all_errors.append("Earnings_Data_Appendix.xlsx not found in agent workspace")
    else:
        agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
        gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

        # Check Estimates vs Actuals sheet
        print("  Checking Estimates vs Actuals sheet...")
        a_rows = load_sheet_rows(agent_wb, "Estimates vs Actuals")
        g_rows = load_sheet_rows(gt_wb, "Estimates vs Actuals")
        if a_rows is None:
            all_errors.append("Sheet 'Estimates vs Actuals' not found")
        else:
            a_data = a_rows[1:] if len(a_rows) > 1 else []
            if len(a_data) < 5:
                all_errors.append(f"Estimates vs Actuals: {len(a_data)} rows, expected 5")
            else:
                a_lookup = {}
                for row in a_data:
                    if row and row[0]:
                        a_lookup[str(row[0]).strip().upper()] = row
                for g_row in g_rows[1:]:
                    if not g_row or not g_row[0]:
                        continue
                    ticker = str(g_row[0]).strip().upper()
                    a_row = a_lookup.get(ticker)
                    if a_row is None:
                        all_errors.append(f"Missing ticker: {ticker}")
                        continue
                    # Current_Price (col 3)
                    if not num_close(a_row[3], g_row[3], 5):
                        all_errors.append(f"{ticker} Price: {a_row[3]} vs {g_row[3]}")
                    # Trailing_EPS (col 4)
                    if not num_close(a_row[4], g_row[4], 0.5):
                        all_errors.append(f"{ticker} Trailing_EPS: {a_row[4]} vs {g_row[4]}")
                    # Consensus_EPS (col 5)
                    if not num_close(a_row[5], g_row[5], 0.5):
                        all_errors.append(f"{ticker} Consensus_EPS: {a_row[5]} vs {g_row[5]}")
            print("    Done.")

        # Check Market Summary sheet
        print("  Checking Market Summary sheet...")
        a_rows2 = load_sheet_rows(agent_wb, "Market Summary")
        if a_rows2 is None:
            all_errors.append("Sheet 'Market Summary' not found")
        else:
            a_data2 = a_rows2[1:] if len(a_rows2) > 1 else []
            if len(a_data2) < 5:
                all_errors.append(f"Market Summary: {len(a_data2)} rows, expected 5")
            print("    Done.")

    # --- Check 2: Word document ---
    print("Checking Word document...")
    doc_path = os.path.join(agent_ws, "Earnings_Preview_Report.docx")
    if not os.path.exists(doc_path):
        all_errors.append("Earnings_Preview_Report.docx not found")
    else:
        try:
            from docx import Document
            doc = Document(doc_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            for sym in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
                if sym not in full_text:
                    all_errors.append(f"Stock symbol {sym} not found in Word doc")
            if "earnings" not in full_text.lower():
                all_errors.append("Word doc does not mention 'earnings'")
        except Exception as e:
            all_errors.append(f"Error reading Word doc: {e}")

    # --- Check 3: Email sent ---
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(*) FROM email.messages
            WHERE to_addr::text ILIKE '%equity-team@company.com%'
            AND subject ILIKE '%earnings%'
        """)
        count = cur.fetchone()[0]
        if count == 0:
            all_errors.append("No email sent to equity-team@company.com about earnings")
        else:
            print(f"    Email found ({count})")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking email: {e}")

    # --- Check 4: GCal events ---
    print("Checking GCal events...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT summary, start_datetime::date
            FROM gcal.events
            WHERE summary ILIKE '%earnings%watch%'
            ORDER BY start_datetime
        """)
        rows = cur.fetchall()
        cur.close()
        conn.close()
        if len(rows) < 5:
            all_errors.append(f"GCal: found {len(rows)} earnings watch events, expected 5")
        else:
            # Check for each ticker
            summaries = " ".join(r[0] for r in rows).upper()
            for ticker in ["GOOGL", "AMZN", "JPM", "JNJ", "XOM"]:
                if ticker not in summaries:
                    all_errors.append(f"GCal: no earnings watch event for {ticker}")
    except Exception as e:
        all_errors.append(f"Error checking GCal: {e}")

    # --- Final result ---
    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
