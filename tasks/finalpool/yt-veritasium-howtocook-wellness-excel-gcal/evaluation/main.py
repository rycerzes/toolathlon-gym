"""
Evaluation for yt-veritasium-howtocook-wellness-excel-gcal task.

Checks:
1. Wellness_Plan.xlsx exists with "Videos" sheet having >= 4 rows
2. "Recipes" sheet has >= 8 rows
3. "Weekly_Meal_Plan" sheet has exactly 7 rows
4. Wellness_Guide.docx exists
5. Word doc has >= 3 heading-level sections
6. Word doc contains health/nutrition/recipe/wellness keywords
7. GCal has >= 3 new Wellness Check-in events in April 2026
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def check_excel(agent_workspace, groundtruth_workspace="."):
    print("\n=== Check 1-3: Wellness_Plan.xlsx ===")
    xlsx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".xlsx") and ("wellness" in fname.lower() or "plan" in fname.lower()):
            xlsx_path = os.path.join(agent_workspace, fname)
            break
    if not xlsx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".xlsx"):
                xlsx_path = os.path.join(agent_workspace, fname)
                break

    record("Wellness_Plan.xlsx exists", xlsx_path is not None,
           f"No matching xlsx in {agent_workspace}")

    if not xlsx_path:
        for chk in ["Videos sheet >= 4 rows", "Recipes sheet >= 8 rows", "Weekly_Meal_Plan sheet = 7 rows"]:
            record(chk, False, "xlsx not found")
        return

    try:
        import openpyxl
        wb = openpyxl.load_workbook(xlsx_path)

        # Videos sheet
        videos_sheet = None
        for name in wb.sheetnames:
            if "video" in name.lower():
                videos_sheet = wb[name]
                break
        if not videos_sheet and wb.sheetnames:
            videos_sheet = wb[wb.sheetnames[0]]

        if videos_sheet:
            data_rows = [r for r in videos_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Videos sheet has >= 4 data rows", len(data_rows) >= 4,
                   f"Found {len(data_rows)} rows")
        else:
            record("Videos sheet has >= 4 data rows", False, "No Videos sheet")

        # Recipes sheet
        recipe_sheet = None
        for name in wb.sheetnames:
            if "recipe" in name.lower():
                recipe_sheet = wb[name]
                break
        record("Recipes sheet exists", recipe_sheet is not None, f"Sheets: {wb.sheetnames}")

        if recipe_sheet:
            data_rows = [r for r in recipe_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Recipes sheet has >= 8 data rows", len(data_rows) >= 8,
                   f"Found {len(data_rows)} rows")
        else:
            record("Recipes sheet has >= 8 data rows", False, "Sheet not found")

        # Weekly_Meal_Plan sheet
        meal_sheet = None
        for name in wb.sheetnames:
            if "meal" in name.lower() or "weekly" in name.lower() or "plan" in name.lower():
                meal_sheet = wb[name]
                break
        record("Weekly_Meal_Plan sheet exists", meal_sheet is not None, f"Sheets: {wb.sheetnames}")

        if meal_sheet:
            data_rows = [r for r in meal_sheet.iter_rows(min_row=2, values_only=True)
                         if any(c is not None for c in r)]
            record("Weekly_Meal_Plan sheet has exactly 7 rows", len(data_rows) == 7,
                   f"Found {len(data_rows)} rows (expected 7)")
        else:
            record("Weekly_Meal_Plan sheet has exactly 7 rows", False, "Sheet not found")

        # --- Groundtruth XLSX value comparison ---
        gt_path = os.path.join(groundtruth_workspace, "Wellness_Plan.xlsx")
        if os.path.isfile(gt_path):
            gt_wb = openpyxl.load_workbook(gt_path, data_only=True)
            for gt_sname in gt_wb.sheetnames:
                gt_ws = gt_wb[gt_sname]
                a_ws = None
                for asn in wb.sheetnames:
                    if asn.strip().lower() == gt_sname.strip().lower():
                        a_ws = wb[asn]
                        break
                if a_ws is None:
                    record(f"GT sheet '{gt_sname}' exists in agent xlsx", False, f"Available: {wb.sheetnames}")
                    continue
                gt_rows = [r for r in gt_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                a_rows = [r for r in a_ws.iter_rows(min_row=2, values_only=True) if any(c is not None for c in r)]
                record(f"GT '{gt_sname}' row count", len(a_rows) == len(gt_rows),
                       f"Expected {len(gt_rows)}, got {len(a_rows)}")
                for ri in range(min(3, len(gt_rows))):
                    if ri >= len(a_rows):
                        break
                    ok = True
                    for ci in range(min(len(gt_rows[ri]), len(a_rows[ri]))):
                        gv, av = gt_rows[ri][ci], a_rows[ri][ci]
                        if gv is None:
                            continue
                        if isinstance(gv, (int, float)):
                            if not num_close(av, gv, max(abs(gv) * 0.1, 1.0)):
                                ok = False
                                break
                        else:
                            if not str_match(av, gv):
                                ok = False
                                break
                    record(f"GT '{gt_sname}' row {ri+1} values", ok,
                           f"gt={gt_rows[ri][:4]}, agent={a_rows[ri][:4] if ri < len(a_rows) else 'missing'}")
            gt_wb.close()

    except Exception as e:
        for chk in ["Videos sheet >= 4 rows", "Recipes sheet >= 8 rows", "Weekly_Meal_Plan = 7 rows"]:
            record(chk, False, str(e))


def check_word(agent_workspace):
    print("\n=== Check 4-6: Wellness_Guide.docx ===")
    docx_path = None
    for fname in os.listdir(agent_workspace):
        if fname.lower().endswith(".docx") and ("wellness" in fname.lower() or "guide" in fname.lower()):
            docx_path = os.path.join(agent_workspace, fname)
            break
    if not docx_path:
        for fname in os.listdir(agent_workspace):
            if fname.lower().endswith(".docx"):
                docx_path = os.path.join(agent_workspace, fname)
                break

    record("Wellness_Guide.docx exists", docx_path is not None,
           f"No wellness/guide docx in {agent_workspace}")

    if not docx_path:
        record("Word doc has >= 3 headings", False, "docx not found")
        record("Word doc contains wellness keywords", False, "docx not found")
        return

    try:
        from docx import Document
        doc = Document(docx_path)
        headings = [p for p in doc.paragraphs if p.style.name.lower().startswith("heading")]
        record("Word doc has >= 3 headings", len(headings) >= 3,
               f"Found {len(headings)} headings")

        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        keywords = ["health", "nutrition", "recipe", "wellness", "diet", "meal", "biology", "science"]
        found = [k for k in keywords if k in full_text]
        record("Word doc contains health/wellness/recipe keywords", len(found) >= 3,
               f"Found keywords: {found}")
    except Exception as e:
        record("Word doc has >= 3 headings", False, str(e))
        record("Word doc contains wellness keywords", False, str(e))


def check_gcal():
    print("\n=== Check 7: GCal Wellness Check-in Events in April 2026 ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT summary, start_datetime FROM gcal.events
        WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
        AND summary ILIKE '%wellness%'
        ORDER BY start_datetime
    """)
    events = cur.fetchall()

    if not events:
        # Broader check for any new events in April 2026 (beyond the 2 preinjected ones)
        cur.execute("""
            SELECT summary, start_datetime FROM gcal.events
            WHERE start_datetime >= '2026-04-01' AND start_datetime < '2026-05-01'
            AND (summary ILIKE '%check%' OR summary ILIKE '%wellness%' OR summary ILIKE '%checkin%')
            ORDER BY start_datetime
        """)
        events = cur.fetchall()

    cur.close()
    conn.close()

    record("GCal has >= 3 Wellness events in April 2026",
           len(events) >= 3, f"Found {len(events)} wellness events in April 2026")

    if events:
        wed_events = [e for e in events if e[1] and e[1].weekday() == 2]
        record("Wellness events are on Wednesdays", len(wed_events) >= 3,
               f"Wednesday events: {len(wed_events)}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_excel(args.agent_workspace, args.groundtruth_workspace)
    check_word(args.agent_workspace)
    check_gcal()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
