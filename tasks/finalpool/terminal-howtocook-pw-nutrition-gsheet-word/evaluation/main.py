"""Evaluation for terminal-howtocook-pw-nutrition-gsheet-word.
Checks:
1. Google Sheets "Nutrition Dashboard" with Recipe Comparison and Daily Plan sheets
2. Wellness_Diet_Plan.docx with required sections
3. nutrition_calculator.py script exists
"""
import argparse
import json
import os
import sys
import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        print(f"  [FAIL] {name}: {str(detail)[:200]}")


def check_gsheet():
    print("\n=== Check 1: Google Sheets Nutrition Dashboard ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()
    dashboard = None
    for ss_id, title in spreadsheets:
        if title and "nutrition" in title.lower():
            dashboard = (ss_id, title)
            break

    check("Nutrition Dashboard spreadsheet exists", dashboard is not None,
          f"Spreadsheets found: {[s[1] for s in spreadsheets]}")

    if dashboard:
        ss_id = dashboard[0]
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (ss_id,))
        sheet_list = cur.fetchall()
        sheet_titles = [s[1].lower() for s in sheet_list]
        check("Has at least 2 sheets", len(sheet_list) >= 2, f"Found {len(sheet_list)}: {[s[1] for s in sheet_list]}")

        has_recipe = any("recipe" in t or "comparison" in t for t in sheet_titles)
        has_daily = any("daily" in t or "plan" in t for t in sheet_titles)
        check("Has Recipe Comparison sheet", has_recipe, f"Sheets: {sheet_titles}")
        check("Has Daily Plan sheet", has_daily, f"Sheets: {sheet_titles}")

        # Check Recipe Comparison sheet has data
        recipe_sheet = next((s for s in sheet_list if "recipe" in s[1].lower() or "comparison" in s[1].lower()), None)
        if recipe_sheet:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, recipe_sheet[0]))
            row_count = cur.fetchone()[0]
            check("Recipe sheet has at least 9 recipe rows", row_count >= 9,
                  f"Found {row_count} data rows")

        # Check Daily Plan sheet has data
        daily_sheet = next((s for s in sheet_list if "daily" in s[1].lower() or "plan" in s[1].lower()), None)
        if daily_sheet:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, daily_sheet[0]))
            row_count = cur.fetchone()[0]
            check("Daily Plan has at least 3 meal rows", row_count >= 3,
                  f"Found {row_count} data rows")

        # Check for calorie data in cells
        cur.execute("""
            SELECT value FROM gsheet.cells
            WHERE spreadsheet_id = %s AND row_index = 0
            ORDER BY col_index
        """, (ss_id,))
        header_cells = cur.fetchall()
        header_text = " ".join(str(c[0]) for c in header_cells).lower()
        check("Headers mention calories", "calori" in header_text or "kcal" in header_text,
              f"Headers: {header_text[:100]}")

    cur.close()
    conn.close()


def check_word(workspace):
    print("\n=== Check 2: Wellness_Diet_Plan.docx ===")
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if not os.path.exists(path):
        check("Word document exists", False, f"Not found at {path}")
        return
    check("Word document exists", True)

    try:
        from docx import Document
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()
        check("Doc mentions wellness or diet plan", "wellness" in full_text or "diet plan" in full_text,
              f"Text: {full_text[:100]}")
        check("Doc mentions nutritional standards", "standard" in full_text or "recommended" in full_text,
              f"Text: {full_text[:100]}")
        check("Doc mentions breakfast", "breakfast" in full_text, f"Text: {full_text[:100]}")
        check("Doc mentions lunch", "lunch" in full_text, f"Text: {full_text[:100]}")
        check("Doc mentions dinner", "dinner" in full_text, f"Text: {full_text[:100]}")
        check("Doc mentions calories", "calori" in full_text, f"Text: {full_text[:100]}")
        check("Doc has gap analysis section", "gap" in full_text or "supplement" in full_text,
              f"Text: {full_text[:200]}")
    except Exception as e:
        check("Word readable", False, str(e))


def check_script(workspace):
    print("\n=== Check 3: nutrition_calculator.py ===")
    path = os.path.join(workspace, "nutrition_calculator.py")
    check("nutrition_calculator.py exists", os.path.exists(path))


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in output."""
    print("\n=== Reverse Validation ===")

    # Word doc should not contain placeholder text
    path = os.path.join(workspace, "Wellness_Diet_Plan.docx")
    if os.path.isfile(path):
        try:
            from docx import Document
            doc = Document(path)
            full_text = " ".join(p.text for p in doc.paragraphs).lower()
            check("Word doc has no placeholder text",
                  "[insert" not in full_text and "todo" not in full_text and "xxx" not in full_text,
                  "Found placeholder text in document")
        except Exception:
            pass

    # GSheet: no duplicate spreadsheets with same nutrition dashboard name
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT title, COUNT(*) FROM gsheet.spreadsheets
            WHERE LOWER(title) LIKE '%%nutrition%%'
            GROUP BY title HAVING COUNT(*) > 1
        """)
        dupes = cur.fetchall()
        check("No duplicate Nutrition Dashboard spreadsheets", len(dupes) == 0,
              f"Found duplicates: {dupes}")
        cur.close()
        conn.close()
    except Exception:
        pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_script(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {"total_passed": PASS_COUNT, "total_checks": total, "accuracy": accuracy}
    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
