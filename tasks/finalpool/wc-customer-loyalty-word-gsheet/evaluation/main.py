"""Evaluation for wc-customer-loyalty-word-gsheet.

Blocking check: Customer_Loyalty.docx (Word document structure and content).
Non-blocking check: Google Sheet data (DB check).
"""
import argparse
import os
import sys

try:
    from docx import Document
except ImportError:
    Document = None

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace, gt_workspace):
    """Check Word document structure and content against groundtruth."""
    print("\n=== Checking Word Document ===")

    agent_path = os.path.join(agent_workspace, "Customer_Loyalty.docx")
    gt_path = os.path.join(gt_workspace, "Customer_Loyalty.docx")

    if not os.path.isfile(agent_path):
        record("Word file exists", False, f"Not found: {agent_path}")
        return False
    record("Word file exists", True)

    if Document is None:
        record("python-docx installed", False, "Cannot import docx")
        return False

    doc = Document(agent_path)
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Check title
    record("Has title 'customer loyalty analysis report'",
           "customer loyalty analysis report" in full_text,
           "Title not found")

    # Check date
    record("Mentions 2026-03-06",
           "2026-03-06" in full_text,
           "Date not found")

    # Check tier sections
    record("Has 'tier summary' section",
           "tier summary" in full_text,
           "'Tier Summary' not found")
    record("Has 'gold tier' section",
           "gold tier" in full_text,
           "'Gold Tier' not found")
    record("Has 'silver tier' section",
           "silver tier" in full_text,
           "'Silver Tier' not found")
    record("Has 'bronze tier' section",
           "bronze tier" in full_text,
           "'Bronze Tier' not found")

    # Load groundtruth doc to verify customer counts
    if os.path.isfile(gt_path):
        gt_doc = Document(gt_path)
        gt_text = "\n".join(p.text for p in gt_doc.paragraphs).lower()

        # Check for key customers in groundtruth appearing in agent output
        # Top spender should be mentioned
        record("Mentions 'william gonzalez' (top spender)",
               "william gonzalez" in full_text or "gonzalez" in full_text,
               "Top spender not found")

        # Check total active customers count (43)
        record("Mentions total active customers",
               "43" in full_text,
               "Total active customers count 43 not found")

        # Check tier counts are mentioned
        record("Mentions gold count (14)",
               "14" in full_text,
               "Gold tier count not found")

    return True


def check_gsheet_nonblocking():
    """Non-blocking: check Google Sheet DB for created data."""
    print("\n=== Non-blocking: Google Sheet DB Check ===")
    try:
        import psycopg2
        conn = psycopg2.connect(
            host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym",
            user="eigent", password="camel"
        )
        cur = conn.cursor()
        cur.execute("SELECT COUNT(*) FROM gsheet.spreadsheets")
        count = cur.fetchone()[0]
        if count > 0:
            print(f"  [INFO] Found {count} spreadsheet(s) in gsheet schema")
            cur.execute("SELECT COUNT(*) FROM gsheet.cells")
            cell_count = cur.fetchone()[0]
            print(f"  [INFO] Found {cell_count} cell(s) in gsheet schema")
        else:
            print("  [INFO] No spreadsheets found in gsheet schema (non-blocking)")
        cur.close()
        conn.close()
    except Exception as e:
        print(f"  [INFO] GSheet check skipped: {e} (non-blocking)")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    word_ok = check_word(args.agent_workspace, gt_dir)
    check_gsheet_nonblocking()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")

    if FAIL_COUNT == 0:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)
    else:
        print(f"\n=== RESULT: FAIL ({FAIL_COUNT} failures) ===")
        sys.exit(1)


if __name__ == "__main__":
    main()
