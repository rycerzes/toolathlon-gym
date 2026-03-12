"""Evaluation for yf-stock-comparison-word."""
import argparse
import os
import sys

import openpyxl
import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {detail[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=1.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def get_expected_data():
    """Get expected stock data from DB."""
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    symbols = ['GOOGL', 'AMZN', 'JPM']
    summary = {}
    for sym in symbols:
        cur.execute("""
            SELECT date, close FROM yf.stock_prices
            WHERE symbol = %s AND date >= '2025-03-06' AND date <= '2026-03-05'
            ORDER BY date
        """, (sym,))
        rows = cur.fetchall()
        if rows:
            start_p = round(float(rows[0][1]), 2)
            end_p = round(float(rows[-1][1]), 2)
            ret = round(((end_p - start_p) / start_p) * 100, 2)
            summary[sym] = {"start": start_p, "end": end_p, "return": ret, "count": len(rows)}

    # Get total trading days
    cur.execute("""
        SELECT COUNT(DISTINCT date) FROM yf.stock_prices
        WHERE symbol = 'GOOGL' AND date >= '2025-03-06' AND date <= '2026-03-05'
    """)
    total_days = cur.fetchone()[0]

    conn.close()
    return summary, total_days


def check_excel(agent_workspace, groundtruth_workspace):
    """Check Excel output."""
    print("\n=== Checking Excel Output ===")

    agent_file = os.path.join(agent_workspace, "Stock_Comparison.xlsx")
    gt_file = os.path.join(groundtruth_workspace, "Stock_Comparison.xlsx")

    check("Excel file exists", os.path.isfile(agent_file), f"Expected {agent_file}")
    if not os.path.isfile(agent_file):
        return False

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    def get_sheet(wb, name):
        for s in wb.sheetnames:
            if s.strip().lower() == name.strip().lower():
                return wb[s]
        return None

    # Check Daily Prices sheet
    agent_dp = get_sheet(agent_wb, "Daily Prices")
    check("Sheet 'Daily Prices' exists", agent_dp is not None,
          f"Found: {agent_wb.sheetnames}")

    summary, total_days = get_expected_data()

    if agent_dp:
        dp_rows = list(agent_dp.iter_rows(min_row=2, values_only=True))
        check("Daily Prices row count", abs(len(dp_rows) - total_days) <= 2,
              f"Expected ~{total_days}, got {len(dp_rows)}")

    # Check Summary sheet
    agent_sum = get_sheet(agent_wb, "Summary")
    gt_sum = get_sheet(gt_wb, "Summary")
    check("Sheet 'Summary' exists", agent_sum is not None,
          f"Found: {agent_wb.sheetnames}")

    if agent_sum and gt_sum:
        gt_rows = list(gt_sum.iter_rows(min_row=2, values_only=True))
        agent_rows = list(agent_sum.iter_rows(min_row=2, values_only=True))

        check("Summary has 3 stock rows", len(agent_rows) == 3,
              f"Got {len(agent_rows)}")

        for gt_row in gt_rows:
            sym, start_p, end_p, ret_pct = gt_row
            matched = None
            for ar in agent_rows:
                if ar and str_match(ar[0], sym):
                    matched = ar
                    break
            if matched:
                check(f"{sym} Start_Price",
                      num_close(matched[1], start_p, 1.0),
                      f"Expected {start_p}, got {matched[1]}")
                check(f"{sym} End_Price",
                      num_close(matched[2], end_p, 1.0),
                      f"Expected {end_p}, got {matched[2]}")
                check(f"{sym} Return_Pct",
                      num_close(matched[3], ret_pct, 0.5),
                      f"Expected {ret_pct}, got {matched[3]}")
            else:
                check(f"{sym} found in Summary", False)

    return True


def check_word_doc(agent_workspace):
    """Check the Word document."""
    print("\n=== Checking Word Document ===")
    try:
        from docx import Document
    except ImportError:
        check("python-docx installed", False, "pip install python-docx")
        return False

    doc_path = os.path.join(agent_workspace, "Stock_Analysis.docx")
    check("Word file exists", os.path.isfile(doc_path), f"Expected {doc_path}")
    if not os.path.isfile(doc_path):
        return False

    doc = Document(doc_path)

    # Check heading
    has_heading = False
    for p in doc.paragraphs:
        text = p.text.lower()
        if "stock" in text and ("performance" in text or "analysis" in text):
            has_heading = True
            break
    check("Document has stock analysis heading", has_heading)

    # Check table
    check("Document has at least one table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        table = doc.tables[0]
        rows = [r for r in table.rows]
        check("Table has header + 3 data rows", len(rows) >= 4,
              f"Got {len(rows)} rows")

        # Check symbols present
        all_text = " ".join(cell.text for row in table.rows for cell in row.cells).lower()
        for sym in ['googl', 'amzn', 'jpm']:
            check(f"Table mentions {sym.upper()}", sym in all_text)

    # Check conclusion paragraph
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    has_conclusion = "best" in full_text or "highest" in full_text or "worst" in full_text or "lowest" in full_text
    check("Document has conclusion about best/worst", has_conclusion)

    summary, _ = get_expected_data()
    best = max(summary.items(), key=lambda x: x[1]["return"])[0]
    worst = min(summary.items(), key=lambda x: x[1]["return"])[0]
    check(f"Conclusion mentions best performer ({best})",
          best.lower() in full_text)
    check(f"Conclusion mentions worst performer ({worst})",
          worst.lower() in full_text)

    return True


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    print("=" * 70)
    print("YF STOCK COMPARISON WORD - EVALUATION")
    print("=" * 70)

    check_excel(args.agent_workspace, gt_dir)
    check_word_doc(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
