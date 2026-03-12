"""
Evaluation for yt-veritasium-science-quiz-word-gcal task.

Checks:
1. Veritasium_Science_Quiz.docx exists
2. Word doc has >= 4 headings (level 1)
3. Word doc text contains scientific keywords
4. GCal has >= 3 new "Science Study Session" events in March-April 2026
5. GSheet has spreadsheet with "Videos" sheet containing >= 6 rows
6. Email sent to studygroup@university.edu
"""
import os
import sys
import json
from argparse import ArgumentParser

import psycopg2

PASS_COUNT = 0
FAIL_COUNT = 0

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent",
    "password": "camel",
}


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1-3: Veritasium_Science_Quiz.docx ===")
    docx_path = os.path.join(agent_workspace, "Veritasium_Science_Quiz.docx")
    if not os.path.exists(docx_path):
        record("Veritasium_Science_Quiz.docx exists", False, f"Not found at {docx_path}")
        record("Word doc has >= 4 level-1 headings", False, "File missing")
        record("Word doc contains scientific keywords", False, "File missing")
        return
    record("Veritasium_Science_Quiz.docx exists", True)

    try:
        import docx
        doc = docx.Document(docx_path)
    except ImportError:
        # Fallback: check file exists and size > 0
        size = os.path.getsize(docx_path)
        record("Word doc readable (docx library not available, checking size)", size > 1000, f"Size: {size}")
        return
    except Exception as e:
        record("Word doc readable", False, str(e))
        return

    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading 1")]
    record("Word doc has >= 4 level-1 headings", len(headings) >= 4,
           f"Found {len(headings)} Heading 1 paragraphs")

    all_text = " ".join(p.text.lower() for p in doc.paragraphs)
    science_keywords = ["physics", "math", "science", "engineering", "biology",
                        "chemistry", "quantum", "gravity", "speed", "light",
                        "energy", "wave", "probability", "paradox", "experiment"]
    found = [kw for kw in science_keywords if kw in all_text]
    record("Word doc contains scientific keywords", len(found) >= 1,
           f"Found keywords: {found[:5]}")


def check_gcal():
    print("\n=== Check 4: GCal Science Study Sessions ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COUNT(*) FROM gcal.events
                WHERE summary ILIKE '%Science Study%' OR summary ILIKE '%Study Session%'
                  AND start_datetime >= '2026-03-14'
                  AND start_datetime < '2026-04-15'
            """)
            count = cur.fetchone()[0]
        conn.close()
        record("GCal has >= 3 Science Study Session events (Mar-Apr 2026)",
               count >= 3, f"Found {count} events")
    except Exception as e:
        record("GCal check", False, str(e))


def check_gsheet():
    print("\n=== Check 5: GSheet Veritasium Study Tracker ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            # Find spreadsheet with "Veritasium" or "Study" in title
            cur.execute("""
                SELECT s.id, s.title FROM gsheet.spreadsheets s
                WHERE s.title ILIKE '%Veritasium%' OR s.title ILIKE '%Study%'
            """)
            spreadsheets = cur.fetchall()
            if not spreadsheets:
                record("GSheet 'Veritasium Study Tracker' spreadsheet exists", False,
                       "No matching spreadsheet found")
                record("Videos sheet has >= 6 data rows", False, "Spreadsheet missing")
                conn.close()
                return
            record("GSheet 'Veritasium Study Tracker' spreadsheet exists", True,
                   f"Found: {[r[1] for r in spreadsheets]}")

            ss_id = spreadsheets[0][0]
            # Find Videos sheet
            cur.execute("""
                SELECT sh.id, sh.title FROM gsheet.sheets sh
                WHERE sh.spreadsheet_id = %s
                  AND sh.title ILIKE '%%video%%'
            """, (ss_id,))
            sheets = cur.fetchall()
            if not sheets:
                record("Videos sheet has >= 6 data rows", False, "Videos sheet not found")
                conn.close()
                return
            sh_id = sheets[0][0]
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index > 0
            """, (ss_id, sh_id))
            row_count = cur.fetchone()[0]
            record("Videos sheet has >= 6 data rows", row_count >= 6,
                   f"Found {row_count} data rows")
        conn.close()
    except Exception as e:
        record("GSheet check", False, str(e))


def check_email():
    print("\n=== Check 6: Email sent to studygroup@university.edu ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        with conn.cursor() as cur:
            cur.execute("""
                SELECT COUNT(*) FROM email.messages
                WHERE to_addr::text ILIKE '%studygroup@university.edu%'
            """)
            count = cur.fetchone()[0]
            if count == 0:
                try:
                    cur.execute("""
                        SELECT COUNT(*) FROM email.sent_log
                        WHERE to_addr ILIKE '%studygroup@university.edu%'
                    """)
                    count = cur.fetchone()[0]
                except Exception:
                    pass
        conn.close()
        record("Email sent to studygroup@university.edu", count > 0,
               f"Found {count} messages")
    except Exception as e:
        record("Email check", False, str(e))


def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    print(f"Running evaluation for yt-veritasium-science-quiz-word-gcal")
    print(f"Agent workspace: {agent_workspace}")

    check_word_doc(agent_workspace)
    check_gcal()
    check_gsheet()
    check_email()

    all_passed = FAIL_COUNT == 0
    summary = f"Passed: {PASS_COUNT}, Failed: {FAIL_COUNT}"
    print(f"\n{'='*40}")
    print(f"Result: {'PASS' if all_passed else 'FAIL'} - {summary}")

    if res_log_file:
        with open(res_log_file, "w") as f:
            json.dump({"passed": PASS_COUNT, "failed": FAIL_COUNT, "all_passed": all_passed}, f)

    return all_passed, summary


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()
    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)


if __name__ == "__main__":
    main()
