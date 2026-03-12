"""
Evaluation for yt-transcript-word-gform task.

Checks:
1. Afrobeat_Playlist.docx exists with 3 headings (Introduction, Song List, Summary)
2. Word doc mentions Afrobeat content from transcript
3. GForm "Afrobeat Mix Feedback" exists with 3 questions
4. Form has correct question types and options
5. Email sent to music@company.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Check 1: Afrobeat_Playlist.docx ===")
    docx_path = os.path.join(agent_workspace, "Afrobeat_Playlist.docx")
    if not os.path.exists(docx_path):
        record("Afrobeat_Playlist.docx exists", False, f"Not found at {docx_path}")
        return
    record("Afrobeat_Playlist.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith('Heading')]
    heading_text = " ".join(headings).lower()
    has_intro = "introduction" in heading_text or "intro" in heading_text
    has_songs = "song" in heading_text or "playlist" in heading_text or "track" in heading_text
    has_summary = "summary" in heading_text or "conclusion" in heading_text
    record("Has Introduction heading", has_intro, f"Headings: {headings}")
    record("Has Song List heading", has_songs, f"Headings: {headings}")
    record("Has Summary heading", has_summary, f"Headings: {headings}")

    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    record("Doc mentions Afrobeat content", "afrobeat" in full_text or "afrobeats" in full_text,
           "Afrobeat not mentioned in document")
    record("Doc has at least 3 paragraphs of content", len([p for p in doc.paragraphs if len(p.text.strip()) > 30]) >= 3,
           "Less than 3 substantial paragraphs")


def check_gform():
    print("\n=== Check 2: Google Form 'Afrobeat Mix Feedback' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gform.forms WHERE title ILIKE %s",
                ("%Afrobeat Mix Feedback%",))
    forms = cur.fetchall()
    record("Form 'Afrobeat Mix Feedback' exists", len(forms) >= 1,
           f"Forms found: {[f[1] for f in forms]}")

    if forms:
        form_id = forms[0][0]
        cur.execute("SELECT COUNT(*) FROM gform.questions WHERE form_id = %s", (form_id,))
        q_count = cur.fetchone()[0]
        record("Form has exactly 3 questions", q_count == 3, f"Found {q_count} questions")

        cur.execute("""
            SELECT title, question_type, config
            FROM gform.questions WHERE form_id = %s
            ORDER BY position
        """, (form_id,))
        questions = cur.fetchall()

        if len(questions) >= 1:
            q1_title = (questions[0][0] or "").lower()
            q1_type = (questions[0][1] or "").upper()
            record("Q1 is 'How would you rate this mix?' (multiple choice)",
                   ("rate" in q1_title or "rating" in q1_title) and q1_type in ("RADIO", "MULTIPLE_CHOICE"),
                   f"Q1 title: {questions[0][0]}, type: {questions[0][1]}")

            # Check Q1 has Excellent/Good/Average/Poor
            q1_config = questions[0][2] or {}
            if isinstance(q1_config, str):
                import json
                q1_config = json.loads(q1_config)
            options = q1_config.get("options", [])
            options_lower = [str(o).lower() for o in options]
            has_excellent = any("excellent" in o for o in options_lower)
            has_poor = any("poor" in o for o in options_lower)
            record("Q1 has Excellent and Poor options", has_excellent and has_poor,
                   f"Options: {options}")

        if len(questions) >= 2:
            q2_title = (questions[1][0] or "").lower()
            q2_type = (questions[1][1] or "").upper()
            record("Q2 is 'Which song was your favorite?' (text)",
                   "favorite" in q2_title or "favourite" in q2_title or "song" in q2_title,
                   f"Q2: {questions[1][0]}, type: {questions[1][1]}")

        if len(questions) >= 3:
            q3_title = (questions[2][0] or "").lower()
            q3_config = questions[2][2] or {}
            if isinstance(q3_config, str):
                import json
                q3_config = json.loads(q3_config)
            options3 = q3_config.get("options", [])
            options3_lower = [str(o).lower() for o in options3]
            has_yes = any("yes" in o for o in options3_lower)
            has_maybe = any("maybe" in o for o in options3_lower)
            record("Q3 is 'Would you recommend?' with Yes/No/Maybe options",
                   ("recommend" in q3_title) and has_yes and has_maybe,
                   f"Q3: {questions[2][0]}, options: {options3}")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT m.to_addr, m.subject FROM email.messages m
        JOIN email.sent_log sl ON sl.message_id = m.id
        WHERE m.to_addr::text ILIKE %s
        ORDER BY sl.sent_at DESC LIMIT 5
    """, ("%music%",))
    emails = cur.fetchall()

    if not emails:
        cur.execute("""
            SELECT to_addr, subject FROM email.messages
            WHERE to_addr::text ILIKE %s
            ORDER BY date DESC LIMIT 5
        """, ("%music%",))
        emails = cur.fetchall()

    cur.close()
    conn.close()

    record("Email sent to music@company.com", len(emails) >= 1,
           f"Found: {emails}")
    if emails:
        subject = str(emails[0][1]).lower() if emails[0][1] else ""
        record("Email subject mentions 'Afrobeat' or 'Analysis'",
               "afrobeat" in subject or "analysis" in subject or "complete" in subject,
               f"Subject: {emails[0][1]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gform()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
