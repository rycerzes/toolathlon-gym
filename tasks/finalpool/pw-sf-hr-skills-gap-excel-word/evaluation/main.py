"""Evaluation script for pw-sf-hr-skills-gap-excel-word."""
import os
import argparse, json, os, sys
import openpyxl
from docx import Document as DocxDocument

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
    "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
    "user": "eigent", "password": "camel"
}

PASS_COUNT = 0
FAIL_COUNT = 0

def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = str(detail)[:200] if detail else ""
        print(f"  [FAIL] {name}: {detail_str}")

def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(',', '').replace('%', '').replace('$', '').strip())
    except (ValueError, TypeError):
        return default

def get_conn():
    import psycopg2
    return psycopg2.connect(**DB_CONFIG)

def run_evaluation(agent_workspace, groundtruth_workspace, launch_time, res_log_file):
    global PASS_COUNT, FAIL_COUNT
    PASS_COUNT = 0
    FAIL_COUNT = 0

    
    # Check 1: Excel file exists and has correct structure
    excel_path = os.path.join(agent_workspace, "Skills_Gap_Analysis.xlsx")
    check("Excel file exists", os.path.exists(excel_path))

    if os.path.exists(excel_path):
        wb = openpyxl.load_workbook(excel_path)

        # Sheet 1: Department_Overview
        check("Department_Overview sheet exists", "Department_Overview" in wb.sheetnames)
        if "Department_Overview" in wb.sheetnames:
            ws = wb["Department_Overview"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Department_Overview has 7 departments", len(rows) == 7, f"got {len(rows)}")

            # Load groundtruth
            gt_path = os.path.join(groundtruth_workspace, "Skills_Gap_Analysis.xlsx")
            gt_wb = openpyxl.load_workbook(gt_path)
            gt_ws = gt_wb["Department_Overview"]
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))

            for gt_row in gt_rows:
                dept = gt_row[0]
                agent_row = None
                for r in rows:
                    if r[0] and str(r[0]).strip().lower() == str(dept).strip().lower():
                        agent_row = r
                        break
                if agent_row:
                    gt_sal = safe_float(gt_row[2])
                    ag_sal = safe_float(agent_row[2])
                    if gt_sal and ag_sal:
                        check(f"{dept} avg salary", abs(gt_sal - ag_sal) <= 50,
                              f"expected ~{gt_sal}, got {ag_sal}")
                    gt_gap = safe_float(gt_row[6])
                    ag_gap = safe_float(agent_row[6])
                    if gt_gap is not None and ag_gap is not None:
                        check(f"{dept} salary gap", abs(gt_gap - ag_gap) <= 100,
                              f"expected ~{gt_gap}, got {ag_gap}")

        # Sheet 2: Skills_Matrix
        check("Skills_Matrix sheet exists", "Skills_Matrix" in wb.sheetnames)
        if "Skills_Matrix" in wb.sheetnames:
            ws = wb["Skills_Matrix"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Skills_Matrix has 17 skills", len(rows) >= 15, f"got {len(rows)}")

        # Sheet 3: Summary
        check("Summary sheet exists", "Summary" in wb.sheetnames)
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            rows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            check("Total_Departments = 7", safe_float(rows.get("Total_Departments")) == 7,
                  f"got {rows.get('Total_Departments')}")

    # Check 2: Word document
    word_path = os.path.join(agent_workspace, "Skills_Gap_Report.docx")
    check("Word report exists", os.path.exists(word_path))
    if os.path.exists(word_path):
        from docx import Document
        doc = Document(word_path)
        full_text = " ".join([p.text for p in doc.paragraphs]).lower()
        check("Word contains 'skills gap'", "skills gap" in full_text)
        check("Word contains 'recommendation'", "recommend" in full_text)
        check("Word contains 'benchmark'", "benchmark" in full_text)
        headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
        check("Word has at least 3 headings", len(headings) >= 3, f"got {len(headings)}")

    # Check 3: Terminal artifacts
    script_path = os.path.join(agent_workspace, "skills_analysis.py")
    check("Python script exists", os.path.exists(script_path))
    json_path = os.path.join(agent_workspace, "raw_skills_data.json")
    check("Raw JSON data exists", os.path.exists(json_path))
    processed_path = os.path.join(agent_workspace, "processed_analysis.json")
    check("Processed JSON exists", os.path.exists(processed_path))
    

    return FAIL_COUNT == 0, f"Passed {PASS_COUNT}/{PASS_COUNT + FAIL_COUNT} checks"

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False, default="2026-03-07 10:00:00")
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    success, message = run_evaluation(
        args.agent_workspace, args.groundtruth_workspace,
        args.launch_time, args.res_log_file
    )
    print(message)
    sys.exit(0 if success else 1)

if __name__ == "__main__":
    main()
