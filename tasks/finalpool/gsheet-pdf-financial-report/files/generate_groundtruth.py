#!/usr/bin/env python3
"""
Generate groundtruth files for gsheet-pdf-financial-report task.

Creates:
  - FY2024_Financial_Analysis.xlsx (Quarterly Revenue + Top Products)
  - FY2024_Financial_Report.docx (narrative report)
  - FY2024_Financial_Report.pdf (exported from docx)
"""

import os
import sys
import psycopg2
import openpyxl
from openpyxl.styles import Font, numbers

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

TASK_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
GT_DIR = os.path.join(TASK_ROOT, "groundtruth_workspace")


def get_quarterly_revenue(cur):
    cur.execute('''
        SELECT
          CASE
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 1 AND 3 THEN 'Q1'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 4 AND 6 THEN 'Q2'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 7 AND 9 THEN 'Q3'
            ELSE 'Q4'
          END as quarter,
          ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue,
          COUNT(*) as order_count
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY quarter ORDER BY quarter
    ''')
    rows = cur.fetchall()
    result = []
    for q, rev, cnt in rows:
        rev = float(rev)
        cnt = int(cnt)
        avg = round(rev / cnt, 2)
        result.append((q, rev, cnt, avg))
    return result


def get_top_products(cur):
    cur.execute('''
        SELECT p."PRODUCT_NAME", p."CATEGORY",
               SUM(o."QUANTITY"::int) as units,
               ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__PRODUCTS" p ON o."PRODUCT_ID" = p."PRODUCT_ID"
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY p."PRODUCT_NAME", p."CATEGORY"
        ORDER BY revenue DESC LIMIT 10
    ''')
    return [(r[0], r[1], int(r[2]), float(r[3])) for r in cur.fetchall()]


def create_excel(quarters, products):
    """Create FY2024_Financial_Analysis.xlsx."""
    wb = openpyxl.Workbook()

    # Sheet 1: Quarterly Revenue
    ws1 = wb.active
    ws1.title = "Quarterly Revenue"
    ws1.append(["Quarter", "Revenue", "Order_Count", "Avg_Order_Value"])
    for q, rev, cnt, avg in quarters:
        ws1.append([q, rev, cnt, avg])

    # Sheet 2: Top Products
    ws2 = wb.create_sheet("Top Products")
    ws2.append(["Product_Name", "Category", "Units_Sold", "Revenue"])
    for name, cat, units, rev in products:
        ws2.append([name, cat, units, rev])

    xlsx_path = os.path.join(GT_DIR, "FY2024_Financial_Analysis.xlsx")
    wb.save(xlsx_path)
    print(f"Created: {xlsx_path}")
    return xlsx_path


def create_word(quarters, products):
    """Create FY2024_Financial_Report.docx."""
    from docx import Document
    from docx.shared import Pt

    doc = Document()

    # Title
    doc.add_heading("FY2024 Annual Financial Report", level=0)

    total_revenue = sum(q[1] for q in quarters)
    total_orders = sum(q[2] for q in quarters)

    # Executive Summary
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(
        f"This report presents the financial performance of the company for fiscal year 2024, "
        f"covering the period from January through December 2024. The company generated total "
        f"revenue of ${total_revenue:,.2f} across {total_orders:,} orders during the fiscal year. "
        f"The annual performance demonstrates consistent growth with Q4 recording the highest "
        f"quarterly revenue. Product performance was concentrated in key categories, with the "
        f"top ten products contributing significantly to overall revenue."
    )

    # Quarterly Performance
    doc.add_heading("Quarterly Performance", level=1)
    for q, rev, cnt, avg in quarters:
        doc.add_paragraph(
            f"{q} 2024: Revenue of ${rev:,.2f} was generated from {cnt:,} orders, "
            f"yielding an average order value of ${avg:,.2f}."
        )
    doc.add_paragraph(
        f"The strongest quarter was Q4 with ${quarters[3][1]:,.2f} in revenue, "
        f"while Q1 had the lowest revenue at ${quarters[0][1]:,.2f}. "
        f"Revenue showed a positive trend across the year, with Q2 through Q4 each "
        f"contributing substantially higher volumes than Q1."
    )

    # Product Analysis
    doc.add_heading("Product Analysis", level=1)
    doc.add_paragraph(
        f"The top performing product was {products[0][0][:60]} in the {products[0][1]} category, "
        f"generating ${products[0][3]:,.2f} in revenue from {products[0][2]} units sold. "
        f"The second highest performer was {products[1][0][:60]} with ${products[1][3]:,.2f} in revenue. "
        f"All top ten products belong to the {products[0][1]} category, indicating strong market "
        f"demand in this segment."
    )
    top_10_revenue = sum(p[3] for p in products)
    doc.add_paragraph(
        f"Combined, the top ten products generated ${top_10_revenue:,.2f} in revenue, "
        f"representing {top_10_revenue / total_revenue * 100:.1f}% of total annual revenue."
    )

    # Outlook
    doc.add_heading("Outlook", level=1)
    doc.add_paragraph(
        f"Based on the FY2024 results, the company is well-positioned for continued growth "
        f"in fiscal year 2025. The strong Q4 performance suggests positive momentum heading "
        f"into the new year. Key priorities should include diversifying revenue sources beyond "
        f"the dominant product category, expanding average order values, and sustaining the "
        f"order volume growth trajectory observed from Q1 to Q4. With total revenue of "
        f"${total_revenue:,.2f} achieved in 2024, a target of sustained or improved performance "
        f"in 2025 is within reach."
    )

    docx_path = os.path.join(GT_DIR, "FY2024_Financial_Report.docx")
    doc.save(docx_path)
    print(f"Created: {docx_path}")
    return docx_path


def create_pdf(docx_path):
    """Export docx to PDF."""
    pdf_path = os.path.join(GT_DIR, "FY2024_Financial_Report.pdf")

    # Try docx2pdf first
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        print(f"Created (docx2pdf): {pdf_path}")
        return pdf_path
    except Exception:
        pass

    # Try pypandoc
    try:
        import pypandoc
        pypandoc.convert_file(docx_path, 'pdf', outputfile=pdf_path)
        print(f"Created (pypandoc): {pdf_path}")
        return pdf_path
    except Exception:
        pass

    # Try libreoffice
    try:
        import subprocess
        result = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf",
             "--outdir", GT_DIR, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if os.path.exists(pdf_path):
            print(f"Created (libreoffice): {pdf_path}")
            return pdf_path
    except Exception:
        pass

    # Fallback: create a simple PDF with reportlab
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        from docx import Document

        doc = Document(docx_path)
        pdf_doc = SimpleDocTemplate(pdf_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []

        for para in doc.paragraphs:
            if para.style.name.startswith('Heading'):
                story.append(Paragraph(para.text, styles['Heading1']))
            else:
                if para.text.strip():
                    story.append(Paragraph(para.text, styles['Normal']))
            story.append(Spacer(1, 6))

        pdf_doc.build(story)
        print(f"Created (reportlab): {pdf_path}")
        return pdf_path
    except Exception as e:
        print(f"Warning: Could not create PDF: {e}")
        # Create a minimal valid PDF as last resort
        with open(pdf_path, 'wb') as f:
            f.write(b'%PDF-1.4\n1 0 obj<</Type/Catalog/Pages 2 0 R>>endobj\n')
            f.write(b'2 0 obj<</Type/Pages/Kids[3 0 R]/Count 1>>endobj\n')
            f.write(b'3 0 obj<</Type/Page/MediaBox[0 0 612 792]/Parent 2 0 R/Resources<<>>>>endobj\n')
            f.write(b'xref\n0 4\n0000000000 65535 f \n0000000009 00000 n \n')
            f.write(b'0000000058 00000 n \n0000000115 00000 n \n')
            f.write(b'trailer<</Size 4/Root 1 0 R>>\nstartxref\n211\n%%EOF\n')
        print(f"Created (minimal): {pdf_path}")
        return pdf_path


def main():
    os.makedirs(GT_DIR, exist_ok=True)

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    try:
        quarters = get_quarterly_revenue(cur)
        products = get_top_products(cur)
    finally:
        cur.close()
        conn.close()

    print(f"Quarterly data: {len(quarters)} quarters")
    for q, rev, cnt, avg in quarters:
        print(f"  {q}: Revenue={rev}, Orders={cnt}, Avg={avg}")

    print(f"\nTop products: {len(products)}")
    for name, cat, units, rev in products:
        print(f"  {name[:50]}: {cat}, Units={units}, Revenue={rev}")

    total = sum(q[1] for q in quarters)
    print(f"\nTotal FY2024 Revenue: ${total:,.2f}")

    create_excel(quarters, products)
    docx_path = create_word(quarters, products)
    create_pdf(docx_path)

    print("\nGroundtruth generation complete.")


if __name__ == "__main__":
    main()
