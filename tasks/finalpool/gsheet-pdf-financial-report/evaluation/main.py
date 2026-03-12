"""
Evaluation script for gsheet-pdf-financial-report task.

Dynamically queries PostgreSQL to compute expected FY2024 values,
then checks agent output files for correctness.
"""

import argparse
import os
import sys
from pathlib import Path

import psycopg2
import openpyxl

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def get_quarterly_revenue():
    """Query PostgreSQL for FY2024 quarterly revenue."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute('''
        SELECT
          CASE
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 1 AND 3 THEN 'Q1'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 4 AND 6 THEN 'Q2'
            WHEN EXTRACT(MONTH FROM o."ORDER_DATE"::date) BETWEEN 7 AND 9 THEN 'Q3'
            ELSE 'Q4'
          END as quarter,
          ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue,
          COUNT(*) as order_count
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY quarter ORDER BY quarter
    ''')
    rows = cur.fetchall()
    result = []
    for q, rev, cnt in rows:
        rev = float(rev)
        cnt = int(cnt)
        avg = round(rev / cnt, 2)
        result.append((q, rev, cnt, avg))
    conn.close()
    return result


def get_top_products():
    """Query PostgreSQL for top 10 FY2024 products by revenue."""
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute('''
        SELECT p."PRODUCT_NAME", p."CATEGORY",
               SUM(o."QUANTITY"::int) as units,
               ROUND(SUM(o."TOTAL_AMOUNT"::float)::numeric, 2) as revenue
        FROM sf_data."SALES_DW__PUBLIC__ORDERS" o
        JOIN sf_data."SALES_DW__PUBLIC__PRODUCTS" p ON o."PRODUCT_ID" = p."PRODUCT_ID"
        WHERE o."ORDER_DATE" >= '2024-01-01' AND o."ORDER_DATE" < '2025-01-01'
        GROUP BY p."PRODUCT_NAME", p."CATEGORY"
        ORDER BY revenue DESC LIMIT 10
    ''')
    rows = [(r[0], r[1], int(r[2]), float(r[3])) for r in cur.fetchall()]
    conn.close()
    return rows


def check_excel(workspace):
    """Check FY2024_Financial_Analysis.xlsx for correctness."""
    print("\n--- Check 1: Excel File ---")
    xlsx_path = Path(workspace) / "FY2024_Financial_Analysis.xlsx"
    if not xlsx_path.exists():
        print(f"  [FAIL] FY2024_Financial_Analysis.xlsx not found in {workspace}")
        return False

    wb = openpyxl.load_workbook(xlsx_path, data_only=True)

    # Check sheet names
    if "Quarterly Revenue" not in wb.sheetnames:
        print(f"  [FAIL] Missing 'Quarterly Revenue' sheet. Found: {wb.sheetnames}")
        return False
    if "Top Products" not in wb.sheetnames:
        print(f"  [FAIL] Missing 'Top Products' sheet. Found: {wb.sheetnames}")
        return False

    # --- Check Quarterly Revenue ---
    ws1 = wb["Quarterly Revenue"]
    rows1 = list(ws1.iter_rows(values_only=True))
    if len(rows1) < 2:
        print("  [FAIL] Quarterly Revenue sheet has no data rows")
        return False

    header1 = [str(h).strip() if h else "" for h in rows1[0]]
    expected_cols = ["Quarter", "Revenue", "Order_Count", "Avg_Order_Value"]
    for col in expected_cols:
        if col not in header1:
            print(f"  [FAIL] Quarterly Revenue missing column '{col}'. Found: {header1}")
            return False

    idx = {col: header1.index(col) for col in expected_cols}
    data_rows = rows1[1:]

    if len(data_rows) != 4:
        print(f"  [FAIL] Quarterly Revenue: expected 4 data rows, got {len(data_rows)}")
        return False

    expected_quarters = get_quarterly_revenue()
    for i, (exp_q, exp_rev, exp_cnt, exp_avg) in enumerate(expected_quarters):
        row = data_rows[i]
        q_val = str(row[idx["Quarter"]]).strip() if row[idx["Quarter"]] else ""
        rev_val = row[idx["Revenue"]]
        cnt_val = row[idx["Order_Count"]]
        avg_val = row[idx["Avg_Order_Value"]]

        if q_val != exp_q:
            print(f"  [FAIL] Quarterly Revenue row {i+1}: expected quarter '{exp_q}', got '{q_val}'")
            return False

        if rev_val is None or abs(float(rev_val) - exp_rev) > 5.0:
            print(f"  [FAIL] Quarterly Revenue '{exp_q}': expected revenue {exp_rev}, got {rev_val}")
            return False

        if cnt_val is None or abs(int(cnt_val) - exp_cnt) > 2:
            print(f"  [FAIL] Quarterly Revenue '{exp_q}': expected order count {exp_cnt}, got {cnt_val}")
            return False

        if avg_val is None or abs(float(avg_val) - exp_avg) > 1.0:
            print(f"  [FAIL] Quarterly Revenue '{exp_q}': expected avg {exp_avg}, got {avg_val}")
            return False

    print("  [PASS] Quarterly Revenue data correct")

    # --- Check Top Products ---
    ws2 = wb["Top Products"]
    rows2 = list(ws2.iter_rows(values_only=True))
    if len(rows2) < 2:
        print("  [FAIL] Top Products sheet has no data rows")
        return False

    header2 = [str(h).strip() if h else "" for h in rows2[0]]
    expected_cols2 = ["Product_Name", "Category", "Units_Sold", "Revenue"]
    for col in expected_cols2:
        if col not in header2:
            print(f"  [FAIL] Top Products missing column '{col}'. Found: {header2}")
            return False

    idx2 = {col: header2.index(col) for col in expected_cols2}
    data_rows2 = rows2[1:]

    if len(data_rows2) != 10:
        print(f"  [FAIL] Top Products: expected 10 data rows, got {len(data_rows2)}")
        return False

    expected_products = get_top_products()

    # Check top product matches (by revenue ordering)
    for i, (exp_name, exp_cat, exp_units, exp_rev) in enumerate(expected_products):
        row = data_rows2[i]
        name_val = str(row[idx2["Product_Name"]]).strip() if row[idx2["Product_Name"]] else ""
        rev_val = row[idx2["Revenue"]]
        units_val = row[idx2["Units_Sold"]]

        # Name match: check first 30 chars to handle truncation
        if name_val[:30].lower() != exp_name[:30].lower():
            print(f"  [FAIL] Top Products row {i+1}: expected product starting with '{exp_name[:30]}', got '{name_val[:30]}'")
            return False

        if rev_val is None or abs(float(rev_val) - exp_rev) > 5.0:
            print(f"  [FAIL] Top Products '{exp_name[:40]}': expected revenue {exp_rev}, got {rev_val}")
            return False

        if units_val is None or abs(int(units_val) - exp_units) > 2:
            print(f"  [FAIL] Top Products '{exp_name[:40]}': expected units {exp_units}, got {units_val}")
            return False

    print("  [PASS] Top Products data correct")
    wb.close()
    return True


def check_gsheet():
    """Check that FY2024 Dashboard spreadsheet exists in gsheet schema with quarterly data."""
    print("\n--- Check 2: Google Sheet ---")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    try:
        # Find spreadsheet with "FY2024" or "dashboard" in title
        cur.execute(
            "SELECT id, title FROM gsheet.spreadsheets WHERE LOWER(title) LIKE '%fy2024%' OR LOWER(title) LIKE '%dashboard%'"
        )
        spreadsheets = cur.fetchall()
        if not spreadsheets:
            print("  [FAIL] No spreadsheet found with 'FY2024' or 'dashboard' in the title")
            return False

        ss_id = spreadsheets[0][0]
        ss_title = spreadsheets[0][1]
        print(f"  Found spreadsheet: '{ss_title}' (id={ss_id})")

        # Check that cells exist
        cur.execute(
            "SELECT COUNT(*) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,)
        )
        cell_count = cur.fetchone()[0]
        if cell_count < 16:  # At least header (4 cols) + 4 data rows x 4 cols = 20 cells minimum
            print(f"  [FAIL] Spreadsheet has only {cell_count} cells, expected at least 16")
            return False

        # Check that quarter labels are present
        cur.execute(
            """SELECT value FROM gsheet.cells
               WHERE spreadsheet_id = %s AND LOWER(value) IN ('q1', 'q2', 'q3', 'q4')""",
            (ss_id,),
        )
        found_quarters = cur.fetchall()
        if len(found_quarters) < 4:
            print(f"  [FAIL] Spreadsheet has only {len(found_quarters)} quarter labels, expected 4")
            return False

        print(f"  [PASS] Google Sheet has {cell_count} cells with all 4 quarter labels")
        return True

    finally:
        cur.close()
        conn.close()


def check_word(workspace):
    """Check FY2024_Financial_Report.docx for required content."""
    print("\n--- Check 3: Word Document ---")
    docx_path = Path(workspace) / "FY2024_Financial_Report.docx"
    if not docx_path.exists():
        print(f"  [FAIL] FY2024_Financial_Report.docx not found in {workspace}")
        return False

    try:
        from docx import Document
        doc = Document(str(docx_path))

        # Collect all text
        all_text = " ".join([p.text for p in doc.paragraphs])
        all_text_lower = all_text.lower()

        # Check title heading
        if not any("fy2024" in p.text.lower() and "annual" in p.text.lower() for p in doc.paragraphs):
            # Relaxed: check if title appears anywhere
            if "fy2024 annual financial report" not in all_text_lower:
                print("  [FAIL] Title 'FY2024 Annual Financial Report' not found")
                return False

        # Check sections exist
        required_sections = ["executive summary", "quarterly performance", "product analysis", "outlook"]
        for section in required_sections:
            if section not in all_text_lower:
                print(f"  [FAIL] Section '{section}' not found in document")
                return False

        # Check document references 2024
        if "2024" not in all_text:
            print("  [FAIL] Document does not reference '2024'")
            return False

        # Check document contains some revenue figures (at least one number > 10000)
        import re
        numbers = re.findall(r'[\d,]+\.?\d*', all_text)
        has_revenue = False
        for n in numbers:
            try:
                val = float(n.replace(",", ""))
                if val > 10000:
                    has_revenue = True
                    break
            except ValueError:
                continue
        if not has_revenue:
            print("  [FAIL] Document does not appear to contain any revenue figures")
            return False

        # Check document is not too short
        if len(all_text.strip()) < 200:
            print("  [FAIL] Document is too short (less than 200 characters)")
            return False

        print(f"  [PASS] Word document has all sections and revenue data ({len(all_text)} chars)")
        return True

    except ImportError:
        # Fallback: just check file size
        file_size = docx_path.stat().st_size
        if file_size < 1000:
            print(f"  [FAIL] Word document too small ({file_size} bytes)")
            return False
        print(f"  [PASS] Word document exists ({file_size} bytes), python-docx not available for deep check")
        return True


def check_pdf(workspace):
    """Check FY2024_Financial_Report.pdf exists and has reasonable size."""
    print("\n--- Check 4: PDF File ---")
    pdf_path = Path(workspace) / "FY2024_Financial_Report.pdf"
    if not pdf_path.exists():
        print(f"  [FAIL] FY2024_Financial_Report.pdf not found in {workspace}")
        return False

    size = pdf_path.stat().st_size
    if size < 5000:
        print(f"  [FAIL] PDF file too small ({size} bytes), likely invalid")
        return False

    print(f"  [PASS] PDF exists, size={size} bytes")
    return True


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--res_log_file", required=False)
    parser.add_argument("--launch_time", required=False, help="Launch time")
    args = parser.parse_args()

    workspace = args.agent_workspace

    print("Fetching expected data from database...")
    try:
        quarters = get_quarterly_revenue()
        products = get_top_products()
        total_revenue = sum(q[1] for q in quarters)
        print(f"  Quarters: {len(quarters)}, Top Products: {len(products)}")
        print(f"  Total FY2024 Revenue: ${total_revenue:,.2f}")
    except Exception as e:
        print(f"Error querying database: {e}")
        sys.exit(1)

    all_passed = True

    # Check 1: Excel
    try:
        if not check_excel(workspace):
            all_passed = False
    except Exception as e:
        print(f"  [FAIL] Excel check error: {e}")
        all_passed = False

    # Check 2: Google Sheet
    try:
        if not check_gsheet():
            all_passed = False
    except Exception as e:
        print(f"  [FAIL] Google Sheet check error: {e}")
        all_passed = False

    # Check 3: Word Document
    try:
        if not check_word(workspace):
            all_passed = False
    except Exception as e:
        print(f"  [FAIL] Word check error: {e}")
        all_passed = False

    # Check 4: PDF
    try:
        if not check_pdf(workspace):
            all_passed = False
    except Exception as e:
        print(f"  [FAIL] PDF check error: {e}")
        all_passed = False

    if all_passed:
        print("\nPass all tests!")
    else:
        print("\nSome checks failed.")
        sys.exit(1)
