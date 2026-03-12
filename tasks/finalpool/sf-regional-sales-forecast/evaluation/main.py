"""Evaluation for sf-regional-sales-forecast."""
import argparse
import os
import sys

import openpyxl
import psycopg2


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def str_match(a, b):
    if a is None or b is None:
        return a is None and b is None
    return str(a).strip().lower() == str(b).strip().lower()


def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")

    agent_file = os.path.join(args.agent_workspace, "Sales_Forecast.xlsx")
    gt_file = os.path.join(gt_dir, "Sales_Forecast.xlsx")

    if not os.path.exists(agent_file):
        print(f"FAIL: Agent output not found: {agent_file}")
        sys.exit(1)
    if not os.path.exists(gt_file):
        print(f"FAIL: Groundtruth not found: {gt_file}")
        sys.exit(1)

    agent_wb = openpyxl.load_workbook(agent_file, data_only=True)
    gt_wb = openpyxl.load_workbook(gt_file, data_only=True)

    all_errors = []

    # Check Historical sheet
    print("  Checking Historical...")
    a_rows = load_sheet_rows(agent_wb, "Historical")
    g_rows = load_sheet_rows(gt_wb, "Historical")
    if a_rows is None:
        all_errors.append("Sheet 'Historical' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Historical' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {}
        for r in a_data:
            if r and r[0] is not None and r[1] is not None:
                k = f"{str(r[0]).strip().lower()}|{str(r[1]).strip()}"
                a_lookup[k] = r
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = f"{str(g_row[0]).strip().lower()}|{str(g_row[1]).strip()}"
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing: {g_row[0]}|{g_row[1]}")
                continue
            # Revenue
            if len(a_row) > 3 and len(g_row) > 3:
                if not num_close(a_row[3], g_row[3], 100):
                    errors.append(f"{key}.Revenue: {a_row[3]} vs {g_row[3]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Forecast sheet
    print("  Checking Forecast...")
    a_rows = load_sheet_rows(agent_wb, "Forecast")
    g_rows = load_sheet_rows(gt_wb, "Forecast")
    if a_rows is None:
        all_errors.append("Sheet 'Forecast' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Forecast' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing region: {g_row[0]}")
                continue
            # Trailing_3M_Avg
            if len(a_row) > 1 and len(g_row) > 1:
                if not num_close(a_row[1], g_row[1], 200):
                    errors.append(f"{key}.Trailing_Avg: {a_row[1]} vs {g_row[1]}")
            # GDP_Growth_Pct
            if len(a_row) > 2 and len(g_row) > 2:
                if not num_close(a_row[2], g_row[2], 0.1):
                    errors.append(f"{key}.GDP: {a_row[2]} vs {g_row[2]}")
            # Next_Quarter_Forecast
            if len(a_row) > 5 and len(g_row) > 5:
                if not num_close(a_row[5], g_row[5], 1000):
                    errors.append(f"{key}.Q_Forecast: {a_row[5]} vs {g_row[5]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Summary sheet
    print("  Checking Summary...")
    a_rows = load_sheet_rows(agent_wb, "Summary")
    g_rows = load_sheet_rows(gt_wb, "Summary")
    if a_rows is None:
        all_errors.append("Sheet 'Summary' not found")
    elif g_rows is None:
        all_errors.append("Sheet 'Summary' not found in groundtruth")
    else:
        a_data = a_rows[1:] if len(a_rows) > 1 else []
        g_data = g_rows[1:] if len(g_rows) > 1 else []
        a_lookup = {str(r[0]).strip().lower(): r for r in a_data if r and r[0] is not None}
        errors = []
        for g_row in g_data:
            if not g_row or g_row[0] is None:
                continue
            key = str(g_row[0]).strip().lower()
            a_row = a_lookup.get(key)
            if a_row is None:
                errors.append(f"Missing: {g_row[0]}")
                continue
            if len(a_row) > 1 and len(g_row) > 1:
                if "region" in key:
                    if not str_match(a_row[1], g_row[1]):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "count" in key:
                    if not num_close(a_row[1], g_row[1], 1):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                elif "pct" in key:
                    if not num_close(a_row[1], g_row[1], 1.0):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
                else:
                    if not num_close(a_row[1], g_row[1], 2000):
                        errors.append(f"{key}: {a_row[1]} vs {g_row[1]}")
        if errors:
            all_errors.extend(errors)
            print(f"    ERRORS: {len(errors)}")
            for e in errors[:5]:
                print(f"      {e}")
        else:
            print("    PASS")

    # Check Word document
    print("  Checking Forecast_Report.docx...")
    word_file = os.path.join(args.agent_workspace, "Forecast_Report.docx")
    if not os.path.exists(word_file):
        all_errors.append("Forecast_Report.docx not found")
        print("    FAIL: file not found")
    else:
        try:
            from docx import Document
            doc = Document(word_file)
            text = " ".join(p.text for p in doc.paragraphs).lower()
            checks = [
                ("methodology" in text or "formula" in text, "Missing methodology section"),
                ("forecast" in text, "Missing forecast content"),
                (len(doc.paragraphs) >= 5, "Document too short"),
            ]
            for cond, msg in checks:
                if not cond:
                    all_errors.append(msg)
            print("    PASS" if all(c for c, _ in checks) else "    ERRORS found")
        except Exception as e:
            all_errors.append(f"Word doc error: {e}")
            print(f"    ERROR: {e}")

    # Check Google Sheet
    print("  Checking Google Sheet...")
    try:
        db_config = {
            "host": os.environ.get("PGHOST", "localhost"), "port": 5432,
            "dbname": os.environ.get("PGDATABASE", "toolathlon_gym"),
            "user": "eigent", "password": "camel",
        }
        conn = psycopg2.connect(**db_config)
        cur = conn.cursor()
        cur.execute("SELECT id, title FROM gsheet.spreadsheets WHERE title LIKE '%Forecast%' OR title LIKE '%forecast%'")
        sheets = cur.fetchall()
        if len(sheets) < 1:
            all_errors.append("No Google Sheet with 'Forecast' in title found")
            print("    FAIL: no forecast spreadsheet")
        else:
            ss_id = sheets[0][0]
            cur.execute("SELECT COUNT(*) FROM gsheet.cells WHERE spreadsheet_id = %s", (ss_id,))
            cell_count = cur.fetchone()[0]
            if cell_count < 10:
                all_errors.append(f"Google Sheet has only {cell_count} cells")
                print(f"    FAIL: only {cell_count} cells")
            else:
                print("    PASS")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"GSheet check error: {e}")
        print(f"    ERROR: {e}")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
