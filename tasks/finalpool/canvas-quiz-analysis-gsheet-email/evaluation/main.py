"""Evaluation for canvas-quiz-analysis-gsheet-email."""
import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE LOWER(title) LIKE '%quiz%' AND LOWER(title) LIKE '%performance%'
        """)
        sheets = cur.fetchall()
        check("Google Sheet 'Quiz Performance Tracker' exists",
              len(sheets) >= 1,
              f"Found {len(sheets)} matching spreadsheets")

        if sheets:
            ss_id = sheets[0][0]
            # Check that sheet named 'Quiz Scores' exists
            cur.execute("""
                SELECT id, title FROM gsheet.sheets
                WHERE spreadsheet_id = %s AND LOWER(title) LIKE '%%quiz%%score%%'
            """, (ss_id,))
            quiz_sheets = cur.fetchall()
            check("Sheet 'Quiz Scores' exists in spreadsheet",
                  len(quiz_sheets) >= 1,
                  f"Found {len(quiz_sheets)} matching sheets")

            # Check cells contain quiz data
            cur.execute("""
                SELECT c.value FROM gsheet.cells c
                WHERE c.spreadsheet_id = %s
            """, (ss_id,))
            cells = cur.fetchall()
            all_values = " ".join(str(c[0]) for c in cells if c[0])
            check("GSheet contains course names",
                  "biochemistry" in all_values.lower() or "creative computing" in all_values.lower(),
                  f"Sample values: {all_values[:200]}")
            check("GSheet contains quiz titles (CMA)",
                  "cma" in all_values.lower(),
                  f"Sample values: {all_values[:200]}")

            # Check number of data rows (should have 76 quizzes)
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells c
                WHERE c.spreadsheet_id = %s AND c.row_index > 1
            """, (ss_id,))
            row_count_result = cur.fetchone()
            row_count = row_count_result[0] if row_count_result else 0
            check("GSheet has substantial quiz data rows (>= 10)",
                  row_count >= 10,
                  f"Found {row_count} data rows")

        cur.close()
        conn.close()
    except Exception as e:
        check("GSheet check", False, str(e))


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Quiz_Performance_Summary.docx")
    if not os.path.isfile(docx_path):
        check("Quiz_Performance_Summary.docx exists", False, f"Not found: {docx_path}")
        return
    check("Quiz_Performance_Summary.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        check("Word doc has meaningful content (>= 100 chars)",
              len(all_text.strip()) >= 100,
              f"Content length: {len(all_text)}")
        check("Word doc contains quiz-related content",
              "quiz" in all_text or "score" in all_text or "performance" in all_text,
              f"Sample: {all_text[:200]}")
        check("Word doc contains recommendations",
              "recommend" in all_text or "support" in all_text or "intervention" in all_text,
              f"Sample: {all_text[:200]}")
    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_email():
    print("\n=== Checking Email ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT id, subject, to_addr, body_text
            FROM email.messages
            WHERE to_addr::text ILIKE '%%academic_coordinator@university.edu%%'
               OR subject ILIKE '%%quiz%%performance%%'
               OR subject ILIKE '%%quiz%%analysis%%'
        """)
        emails = cur.fetchall()
        check("Email sent to academic_coordinator@university.edu", len(emails) >= 1,
              "No matching email found")
        if emails:
            email = emails[0]
            subject = str(email[1]).lower() if email[1] else ""
            check("Email subject contains 'quiz'",
                  "quiz" in subject,
                  f"Subject: {email[1]}")
            body = str(email[3]) if email[3] else ""
            check("Email body has content", len(body) > 30,
                  f"Body length: {len(body)}")
        cur.close()
        conn.close()
    except Exception as e:
        check("Email check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_gsheet()
    check_word(args.agent_workspace)
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
