"""Evaluation for terminal-wc-yf-commodity-gsheet-word."""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432,
          dbname=os.environ.get("PGDATABASE", "toolathlon_gym"),
          user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

def get_expected_categories():
    """Dynamically query WC product categories from the DB."""
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("SELECT DISTINCT name FROM wc.categories ORDER BY name")
        cats = [r[0] for r in cur.fetchall()]
        cur.close()
        conn.close()
        return cats if cats else ["Audio", "Cameras", "Electronics", "Home Appliances",
                                  "TV & Home Theater", "Watches"]
    except Exception:
        return ["Audio", "Cameras", "Electronics", "Home Appliances",
                "TV & Home Theater", "Watches"]


EXPECTED_CATEGORIES = get_expected_categories()
SENSITIVITY_MAP = {"Watches": 40, "Electronics": 15, "Audio": 10,
                   "Cameras": 10, "TV & Home Theater": 5, "Home Appliances": 8}


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        d = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{d}")


def safe_float(val, default=None):
    try:
        if val is None:
            return default
        return float(str(val).replace(",", "").replace("$", "").replace("%", "").strip())
    except (ValueError, TypeError):
        return default


def check_gsheet():
    """Check Google Sheets spreadsheet."""
    print("\n=== Checking Google Sheets ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    spreadsheets = cur.fetchall()

    target_ss = None
    for ss_id, title in spreadsheets:
        if "commodity" in (title or "").lower() or "pricing" in (title or "").lower() \
                or "dashboard" in (title or "").lower():
            target_ss = ss_id
            break

    check("Commodity Pricing Dashboard spreadsheet exists", target_ss is not None,
          f"Spreadsheets: {[s[1] for s in spreadsheets]}")

    if target_ss:
        cur.execute("SELECT id, title FROM gsheet.sheets WHERE spreadsheet_id = %s", (target_ss,))
        sheets = cur.fetchall()
        sheet_names = {s[1].lower().replace(" ", "_"): s for s in sheets}

        # Check Product_Categories sheet
        pc_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "product" in name or "categor" in name:
                pc_sheet = sid
                break
        check("Product_Categories sheet exists", pc_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

        if pc_sheet:
            cur.execute("""
                SELECT value FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
                AND col_index = 1
            """, (target_ss, pc_sheet))
            cat_values = [r[0] for r in cur.fetchall() if r[0]]
            cats_found = sum(1 for c in cat_values
                           if any(ec.lower() in c.lower() for ec in EXPECTED_CATEGORIES))
            check("Product_Categories has 6 categories", cats_found >= 5,
                  f"Found {cats_found} matching categories from {cat_values}")

        # Check Gold_Price_Trend sheet
        gp_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "gold" in name or "price" in name or "trend" in name:
                gp_sheet = sid
                break
        check("Gold_Price_Trend sheet exists", gp_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

        if gp_sheet:
            cur.execute("""
                SELECT COUNT(DISTINCT row_index) FROM gsheet.cells
                WHERE spreadsheet_id = %s AND sheet_id = %s AND row_index >= 2
            """, (target_ss, gp_sheet))
            row_count = cur.fetchone()[0]
            check("Gold_Price_Trend has ~20 data rows", row_count >= 18,
                  f"Found {row_count} rows")

        # Check Pricing_Impact sheet
        pi_sheet = None
        for name, (sid, stitle) in sheet_names.items():
            if "impact" in name or "pricing" in name:
                pi_sheet = sid
                break
        check("Pricing_Impact sheet exists", pi_sheet is not None,
              f"Sheets: {[s[1] for s in sheets]}")

    conn.close()


def check_word(ws_path):
    """Check Pricing_Strategy_Report.docx."""
    print("\n=== Checking Word Document ===")
    path = os.path.join(ws_path, "Pricing_Strategy_Report.docx")
    if not os.path.isfile(path):
        check("Word document exists", False, f"Not found: {path}")
        return
    check("Word document exists", True)

    from docx import Document
    doc = Document(path)
    full_text = "\n".join(p.text for p in doc.paragraphs).lower()

    check("Document mentions gold", "gold" in full_text)
    check("Document mentions watches", "watch" in full_text)
    check("Document mentions sensitivity or exposure",
          "sensitiv" in full_text or "exposure" in full_text)
    check("Document mentions trend or moving average",
          "trend" in full_text or "moving average" in full_text)
    check("Document mentions recommendations",
          "recommend" in full_text or "action" in full_text)
    check("Document length >= 500 chars", len(full_text) >= 500,
          f"Length: {len(full_text)}")


def check_xlsx_content(workspace):
    """Check Commodity_Pricing_Dashboard.xlsx has valid content."""
    print("\n=== Checking XLSX Content ===")
    try:
        import openpyxl
    except ImportError:
        check("openpyxl available", False, "Cannot import openpyxl")
        return False

    xlsx_path = os.path.join(workspace, "Commodity_Pricing_Dashboard.xlsx")
    if not os.path.isfile(xlsx_path):
        check("Commodity_Pricing_Dashboard.xlsx exists", False, f"Not found: {xlsx_path}")
        return False
    check("Commodity_Pricing_Dashboard.xlsx exists", True)

    try:
        wb = openpyxl.load_workbook(xlsx_path, data_only=True)
        check("XLSX has at least one sheet", len(wb.worksheets) >= 1,
              f"Found {len(wb.worksheets)} sheets")
        all_ok = True
        for ws in wb.worksheets:
            rows = list(ws.iter_rows(values_only=True))
            has_data = len(rows) >= 2
            check(f"XLSX sheet '{ws.title}' has data rows", has_data,
                  f"Only {len(rows)} rows")
            if not has_data:
                all_ok = False
        wb.close()
        return all_ok
    except Exception as e:
        check("XLSX readable", False, str(e))
        return False


def check_reverse_validation(workspace):
    """Verify things that should NOT exist in the output."""
    print("\n=== Reverse Validation ===")
    # Google Sheet: no duplicate spreadsheets with same commodity/pricing name
    try:
        conn = psycopg2.connect(**DB)
        cur = conn.cursor()
        cur.execute("""
            SELECT title, COUNT(*) FROM gsheet.spreadsheets
            WHERE lower(title) LIKE '%%commodity%%' OR lower(title) LIKE '%%pricing%%'
            GROUP BY title HAVING COUNT(*) > 1
        """)
        dupes = cur.fetchall()
        check("No duplicate pricing spreadsheets", len(dupes) == 0,
              f"Duplicates: {dupes}")
        cur.close()
        conn.close()
    except Exception:
        pass

    # Word: document should not contain placeholder text
    path = os.path.join(workspace, "Pricing_Strategy_Report.docx")
    if os.path.isfile(path):
        try:
            from docx import Document
            doc = Document(path)
            full_text = " ".join(p.text for p in doc.paragraphs).lower()
            check("No placeholder text in Word doc",
                  "[insert" not in full_text and "todo" not in full_text and "xxx" not in full_text,
                  "Found placeholder text")
        except Exception:
            pass


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("TERMINAL-WC-YF-COMMODITY-GSHEET-WORD - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_xlsx_content(args.agent_workspace)
    check_reverse_validation(args.agent_workspace)

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
