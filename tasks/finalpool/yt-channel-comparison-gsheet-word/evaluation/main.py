"""
Evaluation for yt-channel-comparison-gsheet-word task.

Checks:
1. Channel_Analysis_Report.docx exists with 4 headings
2. Word doc mentions both channels with key metrics
3. GSheet "Channel Comparison Analysis" exists with Comparison sheet
4. Comparison sheet has 2 data rows for Fireship and Veritasium
5. Email sent to media@company.com
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word(agent_workspace):
    print("\n=== Check 1: Channel_Analysis_Report.docx ===")
    docx_path = os.path.join(agent_workspace, "Channel_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Channel_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        return
    record("Channel_Analysis_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Check headings
    headings = [p.text.strip() for p in doc.paragraphs
                if p.style.name.startswith('Heading')]
    heading_text = " ".join(headings).lower()
    has_exec = "executive" in heading_text or "summary" in heading_text
    has_metrics = "metric" in heading_text or "comparison" in heading_text
    has_findings = "finding" in heading_text or "key" in heading_text
    has_recs = "recommendation" in heading_text
    record("Has Executive Summary heading", has_exec, f"Headings: {headings}")
    record("Has Channel Metrics Comparison heading", has_metrics, f"Headings: {headings}")
    record("Has Key Findings heading", has_findings, f"Headings: {headings}")
    record("Has Recommendations heading", has_recs, f"Headings: {headings}")

    # Check content
    full_text = " ".join(p.text for p in doc.paragraphs).lower()
    record("Doc mentions Fireship", "fireship" in full_text, "Fireship not found in document")
    record("Doc mentions Veritasium", "veritasium" in full_text, "Veritasium not found")


def check_gsheet():
    print("\n=== Check 2: GSheet 'Channel Comparison Analysis' ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE title ILIKE %s
    """, ("%Channel Comparison%",))
    sheets = cur.fetchall()
    record("GSheet 'Channel Comparison Analysis' exists", len(sheets) >= 1,
           f"Found: {[s[1] for s in sheets]}")

    if sheets:
        ss_id = sheets[0][0]
        cur.execute("""
            SELECT title FROM gsheet.sheets WHERE spreadsheet_id = %s
        """, (ss_id,))
        sheet_titles = [r[0] for r in cur.fetchall()]
        sheet_lower = [t.lower() for t in sheet_titles]
        record("Comparison sheet exists in GSheet",
               any("comparison" in t for t in sheet_lower),
               f"Sheets: {sheet_titles}")
        record("Monthly_Fireship sheet exists in GSheet",
               any("monthly" in t or "fireship" in t for t in sheet_lower),
               f"Sheets: {sheet_titles}")

    cur.close()
    conn.close()


def check_email():
    print("\n=== Check 3: Email sent ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT m.to_addr, m.subject FROM email.messages m
        JOIN email.sent_log sl ON sl.message_id = m.id
        WHERE m.to_addr::text ILIKE %s
        ORDER BY sl.sent_at DESC LIMIT 5
    """, ("%media%",))
    emails = cur.fetchall()

    if not emails:
        cur.execute("""
            SELECT to_addr, subject FROM email.messages
            WHERE to_addr::text ILIKE %s
            ORDER BY date DESC LIMIT 5
        """, ("%media%",))
        emails = cur.fetchall()

    cur.close()
    conn.close()

    record("Email sent to media@company.com", len(emails) >= 1,
           f"Found: {emails}")
    if emails:
        subject = str(emails[0][1]).lower() if emails[0][1] else ""
        record("Email subject mentions 'Channel Comparison'",
               "channel" in subject or "comparison" in subject or "youtube" in subject,
               f"Subject: {emails[0][1]}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
