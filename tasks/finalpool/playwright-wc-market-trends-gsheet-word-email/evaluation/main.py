"""Evaluation for playwright-wc-market-trends-gsheet-word-email."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}


def num_close(a, b, tol=1.0):
    if a is None and b is None:
        return True
    if a is None or b is None:
        return False
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return str(a).strip().lower() == str(b).strip().lower()


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    gt_dir = args.groundtruth_workspace or os.path.join(task_root, "groundtruth_workspace")
    agent_ws = args.agent_workspace or task_root

    all_errors = []

    # --- Check 1: Google Sheet (check via DB) ---
    print("Checking Google Sheet...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check spreadsheet exists
        cur.execute("""
            SELECT id, title FROM gsheet.spreadsheets
            WHERE title ILIKE '%market%trend%'
        """)
        rows = cur.fetchall()
        if not rows:
            all_errors.append("Google Sheet 'Market Trend Analysis 2026' not found")
        else:
            ss_id = rows[0][0]
            # Check sheets exist
            cur.execute("""
                SELECT title FROM gsheet.sheets
                WHERE spreadsheet_id = %s
                ORDER BY "index"
            """, (ss_id,))
            sheet_names = [r[0].strip().lower() for r in cur.fetchall()]

            if not any("industry" in s for s in sheet_names):
                all_errors.append("Sheet 'Industry Trends' not found in GSheet")
            if not any("internal" in s for s in sheet_names):
                all_errors.append("Sheet 'Internal Performance' not found in GSheet")
            if not any("gap" in s for s in sheet_names):
                all_errors.append("Sheet 'Gap Analysis' not found in GSheet")
            else:
                print(f"    Found sheets: {sheet_names}")

            # Check cells for key data
            cur.execute("""
                SELECT COUNT(DISTINCT value) FROM gsheet.cells c
                JOIN gsheet.sheets s ON c.sheet_id = s.id
                WHERE s.spreadsheet_id = %s
                AND LOWER(value) IN ('audio', 'cameras', 'electronics', 'home appliances', 'tv & home theater', 'watches')
            """, (ss_id,))
            cat_count = cur.fetchone()[0]
            if cat_count < 6:
                all_errors.append(f"GSheet: found {cat_count}/6 categories in cells")

        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking GSheet: {e}")

    # --- Check 2: Word document ---
    print("Checking Word document...")
    doc_path = os.path.join(agent_ws, "Market_Insights_Report.docx")
    if not os.path.exists(doc_path):
        all_errors.append("Market_Insights_Report.docx not found")
    else:
        try:
            from docx import Document
            doc = Document(doc_path)
            full_text = "\n".join(p.text for p in doc.paragraphs)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        full_text += "\n" + cell.text
            full_lower = full_text.lower()

            # Check for key content
            categories = ["audio", "cameras", "electronics", "home appliances", "watches"]
            found_cats = sum(1 for c in categories if c in full_lower)
            if found_cats < 4:
                all_errors.append(f"Word doc mentions only {found_cats}/5 categories")

            if "trend" not in full_lower and "market" not in full_lower:
                all_errors.append("Word doc missing market/trend references")

            if "recommendation" not in full_lower and "strategic" not in full_lower:
                all_errors.append("Word doc missing recommendations section")

            if "emerging" not in full_lower:
                all_errors.append("Word doc missing emerging trends section")

        except Exception as e:
            all_errors.append(f"Error reading Word doc: {e}")

    # --- Check 3: Email sent ---
    print("Checking email...")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()
        cur.execute("""
            SELECT COUNT(*) FROM email.messages
            WHERE to_addr::text ILIKE '%product-team@company.com%'
            AND subject ILIKE '%market%trend%'
        """)
        count = cur.fetchone()[0]
        if count == 0:
            all_errors.append("No email sent to product-team@company.com about market trends")
        else:
            print(f"    Email found ({count})")
        cur.close()
        conn.close()
    except Exception as e:
        all_errors.append(f"Error checking email: {e}")

    # --- Check 4: XLSX content ---
    print("Checking XLSX content...")
    xlsx_path = os.path.join(agent_ws, "Market_Trend_Analysis.xlsx")
    if not os.path.exists(xlsx_path):
        all_errors.append("Market_Trend_Analysis.xlsx not found")
    else:
        try:
            import openpyxl
            wb = openpyxl.load_workbook(xlsx_path, data_only=True)
            if len(wb.worksheets) < 1:
                all_errors.append("XLSX has no sheets")
            for ws in wb.worksheets:
                rows = list(ws.iter_rows(values_only=True))
                if len(rows) < 2:
                    all_errors.append(f"XLSX sheet '{ws.title}' has only {len(rows)} rows (need >= 2)")
            wb.close()
            print(f"    XLSX OK ({len(wb.worksheets)} sheets)")
        except Exception as e:
            all_errors.append(f"Error reading XLSX: {e}")

    # --- Final result ---
    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
