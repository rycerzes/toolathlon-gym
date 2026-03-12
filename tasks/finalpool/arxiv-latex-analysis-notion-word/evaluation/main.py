"""
Evaluation for arxiv-latex-analysis-notion-word task.

Checks:
1. Paper_Analysis_Report.docx exists and has at least 4 sections
2. Word doc contains all 3 paper title keywords
3. Word doc mentions Scaling Laws and RLHF/InstructGPT and OPT
4. Word doc has a Comparative Analysis section
5. Notion has 3 pages about the papers
6. Email sent to research_lead@university.edu
"""
import json
import os
import sys
from argparse import ArgumentParser

import psycopg2
from docx import Document

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0


def record(name, passed, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if passed:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_word_doc(agent_workspace):
    print("\n=== Check 1: Word Document Paper_Analysis_Report.docx ===")

    docx_path = os.path.join(agent_workspace, "Paper_Analysis_Report.docx")
    if not os.path.exists(docx_path):
        record("Paper_Analysis_Report.docx exists", False, f"Not found at {docx_path}")
        return
    record("Paper_Analysis_Report.docx exists", True)

    try:
        doc = Document(docx_path)
    except Exception as e:
        record("Word doc readable", False, str(e))
        return
    record("Word doc readable", True)

    # Get all text
    all_text = "\n".join(p.text for p in doc.paragraphs).lower()

    # Count headings/sections
    headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading") or
                (p.text.strip() and len(p.text.strip()) < 100 and
                 any(kw in p.text.lower() for kw in ["scaling", "instruct", "opt:", "analysis", "paper"]))]

    record("Word doc has at least 4 sections", len(headings) >= 4,
           f"Found {len(headings)} section-like headings")

    # Check paper keywords
    has_scaling = "scaling laws" in all_text or "scaling" in all_text
    has_rlhf = "rlhf" in all_text or "instructgpt" in all_text or "follow instructions" in all_text or "human feedback" in all_text
    has_opt = "opt" in all_text and ("open pre-trained" in all_text or "open-source" in all_text or "175b" in all_text)

    record("Mentions Scaling Laws paper", has_scaling, "No scaling laws content found")
    record("Mentions RLHF/InstructGPT paper", has_rlhf, "No RLHF/InstructGPT content found")
    record("Mentions OPT paper", has_opt, "No OPT content found")

    has_comparative = "comparative" in all_text or "comparison" in all_text or "connect" in all_text or "relate" in all_text
    record("Has comparative analysis section", has_comparative,
           "No comparative analysis content found")


def check_notion():
    print("\n=== Check 2: Notion Pages ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()

    cur.execute("SELECT properties FROM notion.pages")
    pages = cur.fetchall()
    cur.close()
    conn.close()

    page_titles = []
    for (props,) in pages:
        if isinstance(props, dict):
            title_prop = props.get("title", {})
            if isinstance(title_prop, dict):
                for item in title_prop.get("title", []):
                    if isinstance(item, dict):
                        page_titles.append(item.get("text", {}).get("content", "").lower())
        page_titles.append(str(props).lower())

    all_page_text = " ".join(page_titles)
    has_scaling = "scaling" in all_page_text
    has_instruct = "instruct" in all_page_text or "human feedback" in all_page_text or "rlhf" in all_page_text
    has_opt = "opt" in all_page_text and ("pre-trained" in all_page_text or "transformer" in all_page_text)

    paper_pages_found = sum([has_scaling, has_instruct, has_opt])
    record("Notion has pages for all 3 papers", paper_pages_found >= 3,
           f"Found: scaling={has_scaling}, instruct={has_instruct}, opt={has_opt}. Total pages: {len(pages)}")

    record("At least 3 Notion pages created", len(pages) >= 3,
           f"Found {len(pages)} pages total")


def check_email():
    print("\n=== Check 3: Email to research_lead@university.edu ===")

    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    messages = cur.fetchall()
    cur.close()
    conn.close()

    matching = None
    for subject, from_addr, to_addr, body_text in messages:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                to_str = " ".join(str(r).lower() for r in parsed) if isinstance(parsed, list) else str(to_addr).lower()
            except Exception:
                to_str = str(to_addr).lower()
        if "research_lead@university.edu" in to_str:
            matching = (subject, from_addr, to_addr, body_text)
            break

    record("Email sent to research_lead@university.edu", matching is not None,
           f"Messages found: {len(messages)}")

    if matching:
        subject, _, _, body_text = matching
        subject_lower = (subject or "").lower()
        body_lower = (body_text or "").lower()
        has_paper_ref = (
            "paper" in subject_lower or "analysis" in subject_lower or "report" in subject_lower or
            "scaling" in body_lower or "rlhf" in body_lower or "opt" in body_lower
        )
        record("Email mentions paper analysis", has_paper_ref,
               f"Subject: {subject}, body preview: {body_text[:100] if body_text else ''}")


def main():
    parser = ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word_doc(args.agent_workspace)
    check_notion()
    check_email()

    total = PASS_COUNT + FAIL_COUNT
    if total == 0:
        print("\nFAIL: No checks were performed.")
        sys.exit(1)

    accuracy = PASS_COUNT / total * 100
    print(f"\nOverall: {PASS_COUNT}/{total} checks passed ({accuracy:.1f}%)")

    result = {
        "total_passed": PASS_COUNT,
        "total_checks": total,
        "accuracy": accuracy,
    }

    if args.res_log_file:
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    if accuracy >= 70:
        print("PASS")
        sys.exit(0)
    else:
        print("FAIL")
        sys.exit(1)


if __name__ == "__main__":
    main()
