"""Evaluation for wc-price-competitor-audit."""
import argparse
import json
import os
import sys

def num_close(a, b, rel_tol=0.15, abs_tol=0.5):
    return abs(float(a) - float(b)) <= max(abs_tol, abs(float(b)) * rel_tol)



def load_sheet_rows(wb, sheet_name):
    for name in wb.sheetnames:
        if name.strip().lower() == sheet_name.strip().lower():
            return [[cell.value for cell in row] for row in wb[name].iter_rows()]
    return None


def check_excel(agent_workspace, gt_data):
    errors = []
    import openpyxl
    path = os.path.join(agent_workspace, "Price_Comparison.xlsx")
    if not os.path.exists(path):
        return ["Price_Comparison.xlsx not found"]
    try:
        wb = openpyxl.load_workbook(path, data_only=True)

        # Check Price Match sheet
        rows = load_sheet_rows(wb, "Price Match")
        if rows is None:
            errors.append("Sheet 'Price Match' not found")
        else:
            data_rows = [r for r in rows[1:] if r and r[0] is not None]
            expected = gt_data["summary"]["total"]
            if len(data_rows) < expected - 2:
                errors.append(f"Price Match has {len(data_rows)} rows, expected ~{expected}")

            # Check specific product prices with tolerance
            gt_by_name = {}
            for m in gt_data["matches"]:
                key = m["name"][:30].lower()
                gt_by_name[key] = m

            matched_count = 0
            for r in data_rows:
                if not r[0]:
                    continue
                name_key = str(r[0])[:30].lower()
                for gt_key, gt_val in gt_by_name.items():
                    if gt_key[:20] in name_key or name_key[:20] in gt_key:
                        matched_count += 1
                        # Check our price
                        if r[1] is not None:
                            try:
                                our_p = float(r[1])
                                if abs(our_p - gt_val["our_price"]) > gt_val["our_price"] * 0.05:
                                    errors.append(f"Our price for '{str(r[0])[:40]}' = {our_p}, expected ~{gt_val['our_price']}")
                            except (ValueError, TypeError):
                                pass
                        # Check status
                        if r[5] is not None:
                            status = str(r[5]).strip()
                            if status.lower() != gt_val["status"].lower():
                                errors.append(f"Status for '{str(r[0])[:40]}' = {status}, expected {gt_val['status']}")
                        break

            if matched_count < expected - 3:
                errors.append(f"Only {matched_count} products matched groundtruth names")

        # Check Summary sheet
        rows2 = load_sheet_rows(wb, "Summary")
        if rows2 is None:
            errors.append("Sheet 'Summary' not found")
        else:
            data_rows2 = [r for r in rows2[1:] if r and r[0] is not None]
            summary_dict = {}
            for r in data_rows2:
                if r[0]:
                    summary_dict[str(r[0]).strip().lower().replace(" ", "_")] = r[1]

            gt_summary = gt_data["summary"]
            check_keys = {
                "overpriced": gt_summary["overpriced"],
                "underpriced": gt_summary["underpriced"],
                "competitive": gt_summary["competitive"],
            }
            for key, expected_val in check_keys.items():
                found = False
                for sk, sv in summary_dict.items():
                    if key in sk.replace("_", "").lower():
                        found = True
                        if sv is not None:
                            try:
                                val = int(float(sv))
                                if abs(val - expected_val) > 1:
                                    errors.append(f"Summary {key} = {val}, expected {expected_val}")
                            except (ValueError, TypeError):
                                pass
                        break

    except Exception as e:
        errors.append(f"Error reading Excel: {e}")
    return errors


def check_word(agent_workspace):
    errors = []
    from docx import Document
    path = os.path.join(agent_workspace, "Competitive_Analysis.docx")
    if not os.path.exists(path):
        return ["Competitive_Analysis.docx not found"]
    try:
        doc = Document(path)
        full_text = " ".join(p.text for p in doc.paragraphs).lower()

        if "competitive" not in full_text and "analysis" not in full_text:
            errors.append("Word doc missing competitive analysis content")

        if "overpriced" not in full_text:
            errors.append("Word doc does not mention overpriced products")

        if "recommend" not in full_text:
            errors.append("Word doc missing recommendations section")

        # Check for tables
        if len(doc.tables) < 1:
            errors.append("Word doc should contain at least one table of overpriced products")

    except Exception as e:
        errors.append(f"Error reading Word doc: {e}")
    return errors


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    task_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    agent_ws = args.agent_workspace or os.path.join(task_root, "groundtruth_workspace")

    gt_data_path = os.path.join(task_root, "files", "groundtruth_data.json")
    with open(gt_data_path) as f:
        gt_data = json.load(f)

    all_errors = []

    print("  Checking Excel file...")
    errs = check_excel(agent_ws, gt_data)
    if errs:
        all_errors.extend(errs)
        for e in errs[:5]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    print("  Checking Word document...")
    errs = check_word(agent_ws)
    if errs:
        all_errors.extend(errs)
        for e in errs[:3]:
            print(f"    ERROR: {e}")
    else:
        print("    PASS")

    if all_errors:
        print(f"\n=== RESULT: FAIL ({len(all_errors)} errors) ===")
        for e in all_errors[:10]:
            print(f"  {e}")
        sys.exit(1)
    else:
        print("\n=== RESULT: PASS ===")
        sys.exit(0)


if __name__ == "__main__":
    main()
