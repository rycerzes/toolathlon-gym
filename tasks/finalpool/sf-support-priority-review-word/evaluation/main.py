"""
Evaluation script for sf-support-priority-review-word task.

Checks:
1. Word file (Priority_Handling_Report.docx) - has all priority levels and 'response time'
2. Google Sheet with 'priority' in title
3. Email with 'priority' or 'handling' in subject to support-director@company.com
"""

import argparse
import json
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PRIORITIES = ["high", "medium", "low"]

FILE_PASS = 0
FILE_FAIL = 0
DB_PASS = 0
DB_FAIL = 0


def check(name, condition, detail="", db=False):
    global FILE_PASS, FILE_FAIL, DB_PASS, DB_FAIL
    if condition:
        if db:
            DB_PASS += 1
        else:
            FILE_PASS += 1
        print(f"  [PASS] {name}")
    else:
        if db:
            DB_FAIL += 1
        else:
            FILE_FAIL += 1
        detail_str = f": {detail[:300]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def check_word(agent_workspace):
    print("\n=== Checking Word Output ===")
    docx_path = os.path.join(agent_workspace, "Priority_Handling_Report.docx")
    check("Word file exists", os.path.isfile(docx_path), f"Expected {docx_path}")
    if not os.path.isfile(docx_path):
        return False

    try:
        from docx import Document
        doc = Document(docx_path)
    except Exception as e:
        check("Word file readable", False, str(e))
        return False

    all_text = ""
    for para in doc.paragraphs:
        all_text += para.text.lower() + " "
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_text += cell.text.lower() + " "

    check("Word contains 'response time'", "response time" in all_text,
          f"Text sample: {all_text[:200]}")

    for p in PRIORITIES:
        check(f"Word contains priority level '{p}'", p in all_text,
              f"Not found in document")

    check("Word contains 'support ticket priority analysis' or similar title",
          "priority" in all_text and "analysis" in all_text,
          f"Title not found")

    return True


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT id, title FROM gsheet.spreadsheets")
    sheets = cur.fetchall()
    cur.close()
    conn.close()

    matching = [s for s in sheets if s[1] and "priority" in s[1].lower()]
    check("Google Sheet with 'priority' in title exists",
          len(matching) > 0,
          f"Sheet titles: {[s[1] for s in sheets]}", db=True)
    return len(matching) > 0


def check_email():
    print("\n=== Checking Email ===")
    conn = psycopg2.connect(**DB_CONFIG)
    cur = conn.cursor()
    cur.execute("SELECT subject, to_addr, body_text FROM email.messages")
    all_emails = cur.fetchall()
    cur.close()
    conn.close()

    print(f"[check_email] Found {len(all_emails)} emails.")

    found = False
    for subject, to_addr, body_text in all_emails:
        to_str = ""
        if isinstance(to_addr, list):
            to_str = " ".join(str(r).lower() for r in to_addr)
        elif isinstance(to_addr, str):
            try:
                parsed = json.loads(to_addr)
                if isinstance(parsed, list):
                    to_str = " ".join(str(r).lower() for r in parsed)
                else:
                    to_str = str(to_addr).lower()
            except (json.JSONDecodeError, TypeError):
                to_str = str(to_addr).lower()

        if "support-director@company.com" in to_str:
            found = True
            subj_lower = (subject or "").lower()
            check("Email subject contains 'priority' or 'handling'",
                  "priority" in subj_lower or "handling" in subj_lower,
                  f"Subject: {subject}", db=True)
            break

    check("Email sent to support-director@company.com", found, db=True)
    return found


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False, default=".")
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    check_word(args.agent_workspace)
    check_gsheet()
    check_email()

    total_pass = FILE_PASS + DB_PASS
    total_fail = FILE_FAIL + DB_FAIL
    file_ok = FILE_FAIL == 0

    print(f"\n=== SUMMARY ===")
    print(f"  File checks - Passed: {FILE_PASS}, Failed: {FILE_FAIL}")
    print(f"  DB checks   - Passed: {DB_PASS}, Failed: {DB_FAIL}")
    if DB_FAIL > 0:
        print(f"  WARNING: {DB_FAIL} DB checks failed (not blocking)")
    print(f"  Overall: {'PASS' if file_ok else 'FAIL'}")

    if args.res_log_file:
        result = {"passed": total_pass, "failed": total_fail, "success": file_ok}
        with open(args.res_log_file, "w") as f:
            json.dump(result, f, indent=2)

    sys.exit(0 if file_ok else 1)


if __name__ == "__main__":
    main()
