"""Evaluation for canvas-grade-equity-excel-word-gcal."""
import argparse
import os
import sys

import psycopg2

DB_CONFIG = {
    "host": os.environ.get("PGHOST", "localhost"),
    "port": 5432,
    "dbname": "toolathlon_gym",
    "user": "eigent",
    "password": "camel",
}

PASS_COUNT = 0
FAIL_COUNT = 0

# Known course data from actual DB
COURSE_NAMES = [
    "Applied Analytics & Algorithms",
    "Biochemistry & Bioinformatics",
    "Data-Driven Design",
    "Environmental Economics & Ethics",
    "Foundations of Finance",
    "Global Governance & Geopolitics",
]

ACTION_REQUIRED_COURSES = ["Biochemistry & Bioinformatics"]


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        msg = f": {str(detail)[:300]}" if detail else ""
        print(f"  [FAIL] {name}{msg}")


def check_excel(agent_workspace, groundtruth_workspace):
    print("\n=== Checking Excel ===")
    xlsx_path = os.path.join(agent_workspace, "Grade_Equity_Analysis.xlsx")
    if not os.path.isfile(xlsx_path):
        check("Grade_Equity_Analysis.xlsx exists", False, f"Not found: {xlsx_path}")
        return
    check("Grade_Equity_Analysis.xlsx exists", True)

    try:
        from openpyxl import load_workbook

        wb = load_workbook(xlsx_path)

        # Check sheet names
        check("Has 'Course Comparison' sheet", "Course Comparison" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Grade Distribution' sheet", "Grade Distribution" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")
        check("Has 'Summary' sheet", "Summary" in wb.sheetnames,
              f"Sheets: {wb.sheetnames}")

        # Load groundtruth for comparison
        gt_path = os.path.join(groundtruth_workspace, "Grade_Equity_Analysis.xlsx")
        gt_wb = load_workbook(gt_path)

        # Check Course Comparison sheet
        if "Course Comparison" in wb.sheetnames:
            ws = wb["Course Comparison"]
            gt_ws = gt_wb["Course Comparison"]

            rows = list(ws.iter_rows(min_row=2, values_only=True))
            gt_rows = list(gt_ws.iter_rows(min_row=2, values_only=True))

            check(f"Course Comparison has {len(gt_rows)} data rows",
                  len(rows) == len(gt_rows),
                  f"Found {len(rows)} rows, expected {len(gt_rows)}")

            # Check each course exists
            agent_names = [str(r[0]).strip() if r[0] else "" for r in rows]
            for name in COURSE_NAMES:
                check(f"Course '{name}' in Course Comparison",
                      name in agent_names,
                      f"Found: {agent_names}")

            # Check numeric values with tolerance
            gt_dict = {str(r[0]).strip(): r for r in gt_rows}
            for row in rows:
                name = str(row[0]).strip() if row[0] else ""
                if name in gt_dict:
                    gt_row = gt_dict[name]
                    # Check mean scores (columns 1,2) with tolerance 0.5
                    for col_idx, col_name in [(1, "Fall_2013_Mean"), (2, "Fall_2014_Mean"),
                                               (3, "Score_Difference")]:
                        if row[col_idx] is not None and gt_row[col_idx] is not None:
                            diff = abs(float(row[col_idx]) - float(gt_row[col_idx]))
                            check(f"{name} {col_name} within tolerance",
                                  diff <= 1.0,
                                  f"Agent={row[col_idx]}, GT={gt_row[col_idx]}, diff={diff:.2f}")

                    # Check equity status
                    if row[7] is not None:
                        check(f"{name} Equity_Status matches",
                              str(row[7]).strip() == str(gt_row[7]).strip(),
                              f"Agent='{row[7]}', GT='{gt_row[7]}'")

        # Check Grade Distribution sheet
        if "Grade Distribution" in wb.sheetnames:
            ws = wb["Grade Distribution"]
            rows = list(ws.iter_rows(min_row=2, values_only=True))
            check("Grade Distribution has 12 data rows (6 courses x 2 years)",
                  len(rows) == 12,
                  f"Found {len(rows)} rows")

        # Check Summary sheet
        if "Summary" in wb.sheetnames:
            ws = wb["Summary"]
            gt_ws = gt_wb["Summary"]
            rows = {str(r[0]).strip(): r[1] for r in ws.iter_rows(min_row=2, values_only=True) if r[0]}
            gt_rows = {str(r[0]).strip(): r[1] for r in gt_ws.iter_rows(min_row=2, values_only=True) if r[0]}

            check("Summary has Total_Courses_Compared",
                  "Total_Courses_Compared" in rows,
                  f"Keys: {list(rows.keys())}")

            if "Total_Courses_Compared" in rows:
                check("Total_Courses_Compared = 6",
                      int(rows["Total_Courses_Compared"]) == 6,
                      f"Got {rows['Total_Courses_Compared']}")

            if "Courses_Action_Required" in rows:
                check("Courses_Action_Required = 1",
                      int(rows["Courses_Action_Required"]) == 1,
                      f"Got {rows['Courses_Action_Required']}")

            for key in ["Overall_Avg_2013", "Overall_Avg_2014"]:
                if key in rows and key in gt_rows:
                    diff = abs(float(rows[key]) - float(gt_rows[key]))
                    check(f"Summary {key} within tolerance",
                          diff <= 1.0,
                          f"Agent={rows[key]}, GT={gt_rows[key]}")

    except ImportError:
        check("openpyxl available", False, "openpyxl not installed")
    except Exception as e:
        check("Excel parsing", False, str(e))


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    docx_path = os.path.join(agent_workspace, "Equity_Report.docx")
    if not os.path.isfile(docx_path):
        check("Equity_Report.docx exists", False, f"Not found: {docx_path}")
        return
    check("Equity_Report.docx exists", True)

    try:
        from docx import Document
        doc = Document(docx_path)
        all_text = " ".join(p.text for p in doc.paragraphs).lower()
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_text += " " + cell.text.lower()

        check("Word doc has meaningful content (>= 200 chars)",
              len(all_text.strip()) >= 200,
              f"Content length: {len(all_text)}")

        # Check required sections
        for section in ["executive summary", "methodology", "course-by-course",
                        "recommendation", "appendix"]:
            check(f"Word doc contains '{section}' section",
                  section in all_text,
                  f"Not found in document text")

        # Check course names mentioned
        found_courses = sum(1 for name in COURSE_NAMES if name.lower() in all_text)
        check("Word doc mentions at least 4 course names",
              found_courses >= 4,
              f"Found {found_courses} of {len(COURSE_NAMES)} course names")

        # Check for action required course
        check("Word doc mentions 'Biochemistry & Bioinformatics'",
              "biochemistry" in all_text,
              "Action required course not mentioned")

        # Check for equity-related content
        check("Word doc discusses equity/grading",
              "equity" in all_text or "grade" in all_text or "score" in all_text,
              f"Sample: {all_text[:300]}")

    except ImportError:
        check("Word doc has content", os.path.getsize(docx_path) > 1000,
              f"Size: {os.path.getsize(docx_path)}")
    except Exception as e:
        check("Word doc readable", False, str(e))


def check_calendar():
    print("\n=== Checking Google Calendar ===")
    try:
        conn = psycopg2.connect(**DB_CONFIG)
        cur = conn.cursor()

        # Check for grade equity review meetings
        cur.execute("""
            SELECT summary, description, start_datetime, end_datetime
            FROM gcal.events
            WHERE LOWER(summary) LIKE '%%equity%%'
               OR LOWER(summary) LIKE '%%grade%%review%%'
               OR LOWER(summary) LIKE '%%biochemistry%%'
        """)
        events = cur.fetchall()
        check("At least 1 grade equity review meeting scheduled",
              len(events) >= 1,
              f"Found {len(events)} matching events")

        if events:
            # Check events are in the week of March 16-20, 2026
            in_target_week = [e for e in events
                              if e[2] and '2026-03-1' in str(e[2]) or '2026-03-20' in str(e[2])]
            check("Meeting(s) scheduled in week of March 16-20, 2026",
                  len(in_target_week) >= 1,
                  f"{len(in_target_week)} events in target week out of {len(events)} total. "
                  f"Dates: {[str(e[2]) for e in events]}")

            # Check event mentions Biochemistry (the action required course)
            event_texts = " ".join(
                (str(e[0]) + " " + str(e[1] or "")).lower() for e in events
            )
            check("Meeting mentions 'Biochemistry' or relevant course",
                  "biochemistry" in event_texts or "bioinformatics" in event_texts,
                  f"Event text: {event_texts[:300]}")

            # Check duration is approximately 45 minutes
            for e in events:
                if e[2] and e[3]:
                    from datetime import datetime
                    start = e[2] if isinstance(e[2], datetime) else datetime.fromisoformat(str(e[2]))
                    end = e[3] if isinstance(e[3], datetime) else datetime.fromisoformat(str(e[3]))
                    duration_min = (end - start).total_seconds() / 60
                    check(f"Meeting '{e[0]}' duration ~45 minutes",
                          30 <= duration_min <= 60,
                          f"Duration: {duration_min} minutes")
                    break  # Check at least one

        cur.close()
        conn.close()
    except Exception as e:
        check("Calendar check", False, str(e))


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=True)
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    gt = args.groundtruth_workspace or os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        "groundtruth_workspace"
    )

    check_excel(args.agent_workspace, gt)
    check_word(args.agent_workspace)
    check_calendar()

    total = PASS_COUNT + FAIL_COUNT
    print(f"\n=== Results: {PASS_COUNT}/{total} passed ===")
    if FAIL_COUNT > 0:
        print(f"{FAIL_COUNT} checks failed")
        sys.exit(1)
    else:
        print("All checks passed!")
        sys.exit(0)


if __name__ == "__main__":
    main()
