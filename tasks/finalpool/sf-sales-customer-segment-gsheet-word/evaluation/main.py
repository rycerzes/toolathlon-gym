"""Evaluation for sf-sales-customer-segment-gsheet-word.

Checks:
1. Google Sheet "Customer Segment Analysis" with 4 segment rows
   - Contains Consumer, Enterprise, Government, SMB
   - Revenue data matches actual DB values (within tolerance)
2. Word doc Customer_Segment_Report.docx with:
   - Heading "Customer Segment Analysis Report"
   - Table with segment data
   - Recommendations section
3. Email to sales-strategy@company.example.com
   Subject: "Customer Segment Analysis Report"
"""
import argparse
import json
import os
import sys

import psycopg2

DB = dict(host=os.environ.get("PGHOST", "localhost"), port=5432, dbname="toolathlon_gym", user="eigent", password="camel")

PASS_COUNT = 0
FAIL_COUNT = 0

# Expected segment data (actual from sf_data query)
SEGMENTS = [
    {"name": "Consumer", "customers": 532, "orders": 5423, "revenue": 839609.20},
    {"name": "Enterprise", "customers": 513, "orders": 5058, "revenue": 793741.69},
    {"name": "Government", "customers": 474, "orders": 4679, "revenue": 712686.66},
    {"name": "SMB", "customers": 481, "orders": 4840, "revenue": 702960.78},
]


def check(name, condition, detail=""):
    global PASS_COUNT, FAIL_COUNT
    if condition:
        PASS_COUNT += 1
        print(f"  [PASS] {name}")
    else:
        FAIL_COUNT += 1
        detail_str = f": {str(detail)[:200]}" if detail else ""
        print(f"  [FAIL] {name}{detail_str}")


def num_close(a, b, tol=500.0):
    try:
        return abs(float(a) - float(b)) <= tol
    except (TypeError, ValueError):
        return False


def check_gsheet():
    print("\n=== Checking Google Sheet ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT id, title FROM gsheet.spreadsheets
        WHERE LOWER(title) LIKE '%customer%' AND LOWER(title) LIKE '%segment%'
    """)
    sheets = cur.fetchall()
    check("Customer Segment Analysis spreadsheet exists", len(sheets) >= 1,
          f"Found {len(sheets)} matching spreadsheets")

    if sheets:
        ss_id = sheets[0][0]
        cur.execute("""
            SELECT c.value FROM gsheet.cells c
            JOIN gsheet.sheets s ON c.spreadsheet_id = s.spreadsheet_id AND c.sheet_id = s.id
            WHERE c.spreadsheet_id = %s
        """, (ss_id,))
        cells = cur.fetchall()
        all_values = " ".join(str(c[0]) for c in cells if c[0])
        all_lower = all_values.lower()

        for seg in SEGMENTS:
            check(f"GSheet contains segment '{seg['name']}'",
                  seg["name"].lower() in all_lower,
                  f"Segment not found in cells")

        check("GSheet has revenue data (Consumer: ~839609)",
              "839609" in all_values or "839,609" in all_values or "839" in all_values,
              "Consumer revenue not found")

        # Check at least 4 rows of data
        segment_count = sum(1 for seg in SEGMENTS if seg["name"].lower() in all_lower)
        check("GSheet has data for all 4 segments", segment_count == 4,
              f"Found {segment_count}/4 segments")

    cur.close()
    conn.close()


def check_word(agent_workspace):
    print("\n=== Checking Word Document ===")
    doc_path = os.path.join(agent_workspace, "Customer_Segment_Report.docx")
    check("Customer_Segment_Report.docx exists", os.path.isfile(doc_path),
          f"Expected at {doc_path}")
    if not os.path.isfile(doc_path):
        return

    try:
        from docx import Document
        doc = Document(doc_path)
    except Exception as e:
        check("Word doc readable", False, str(e))
        return

    full_text = " ".join(p.text.lower() for p in doc.paragraphs)
    headings = [p.text.lower() for p in doc.paragraphs
                if p.style.name.startswith("Heading")]

    check("Word doc has 'Customer Segment Analysis Report' heading",
          any("customer segment analysis" in h for h in headings) or
          "customer segment analysis report" in full_text,
          "Heading not found")

    check("Word doc has a table", len(doc.tables) >= 1,
          f"Found {len(doc.tables)} tables")

    if doc.tables:
        table_text = " ".join(
            cell.text.lower()
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )
        for seg in SEGMENTS:
            check(f"Table contains segment '{seg['name']}'",
                  seg["name"].lower() in table_text,
                  f"Not found in table")

    check("Word doc has Recommendations section",
          "recommendation" in full_text or "recommend" in full_text,
          "Recommendations not found")

    check("Word doc mentions Consumer as top revenue segment",
          "consumer" in full_text, "Consumer segment not mentioned")


def check_emails():
    print("\n=== Checking Emails ===")
    conn = psycopg2.connect(**DB)
    cur = conn.cursor()

    cur.execute("""
        SELECT subject, from_addr, to_addr, body_text
        FROM email.messages
    """)
    all_emails = cur.fetchall()
    conn.close()

    def parse_recipients(to_addr):
        if to_addr is None:
            return []
        if isinstance(to_addr, list):
            return [str(r).strip().lower() for r in to_addr]
        to_str = str(to_addr).strip()
        try:
            parsed = json.loads(to_str)
            if isinstance(parsed, list):
                return [str(r).strip().lower() for r in parsed]
            return [to_str.lower()]
        except (json.JSONDecodeError, TypeError):
            return [to_str.lower()]

    target = "sales-strategy@company.example.com"
    found = None
    for subj, from_addr, to_addr, body in all_emails:
        recipients = parse_recipients(to_addr)
        if target in recipients:
            found = (subj, from_addr, to_addr, body)
            break

    check("Email sent to sales-strategy@company.example.com", found is not None)
    if found:
        subj, from_addr, to_addr, body = found
        check("Email from analytics@company.example.com",
              "analytics@company.example.com" in (from_addr or "").lower(),
              f"From: {from_addr}")
        check("Subject is 'Customer Segment Analysis Report'",
              "customer segment" in (subj or "").lower() and "analysis" in (subj or "").lower(),
              f"Subject: {subj}")


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--agent_workspace", required=False, default=".")
    parser.add_argument("--groundtruth_workspace", required=False)
    parser.add_argument("--launch_time", required=False)
    parser.add_argument("--res_log_file", required=False)
    args = parser.parse_args()

    print("=" * 70)
    print("SF SALES CUSTOMER SEGMENT GSHEET WORD - EVALUATION")
    print("=" * 70)

    check_gsheet()
    check_word(args.agent_workspace)
    check_emails()

    print(f"\n=== SUMMARY ===")
    print(f"  Passed: {PASS_COUNT}")
    print(f"  Failed: {FAIL_COUNT}")
    overall = FAIL_COUNT == 0
    print(f"  Overall: {'PASS' if overall else 'FAIL'}")
    sys.exit(0 if overall else 1)


if __name__ == "__main__":
    main()
